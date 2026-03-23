#!/usr/bin/env python3
"""
Brokr Contract Generator
Generates DOCX contracts from JSON data
Usage: python3 generar_contrato.py <tipo> <datos.json> <output.docx>
"""

import sys
import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def numero_a_letras(n):
    """Convert number to Spanish words for legal contracts"""
    unidades = ['','UNO','DOS','TRES','CUATRO','CINCO','SEIS','SIETE','OCHO','NUEVE',
                'DIEZ','ONCE','DOCE','TRECE','CATORCE','QUINCE','DIECISÉIS',
                'DIECISIETE','DIECIOCHO','DIECINUEVE']
    decenas = ['','DIEZ','VEINTE','TREINTA','CUARENTA','CINCUENTA',
               'SESENTA','SETENTA','OCHENTA','NOVENTA']
    centenas = ['','CIENTO','DOSCIENTOS','TRESCIENTOS','CUATROCIENTOS','QUINIENTOS',
                'SEISCIENTOS','SETECIENTOS','OCHOCIENTOS','NOVECIENTOS']

    def convertir_grupo(n):
        if n == 0: return ''
        if n == 100: return 'CIEN'
        if n < 20: return unidades[n]
        if n < 100:
            d, u = divmod(n, 10)
            return decenas[d] + (' Y ' + unidades[u] if u else '')
        c, r = divmod(n, 100)
        return centenas[c] + (' ' + convertir_grupo(r) if r else '')

    n = int(n)
    if n == 0: return 'CERO'
    if n < 0: return 'MENOS ' + numero_a_letras(-n)

    partes = []
    if n >= 1000000:
        m, r = divmod(n, 1000000)
        partes.append(('UN MILLÓN' if m == 1 else convertir_grupo(m) + ' MILLONES'))
        n = r
    if n >= 1000:
        m, r = divmod(n, 1000)
        partes.append(('MIL' if m == 1 else convertir_grupo(m) + ' MIL'))
        n = r
    if n > 0:
        partes.append(convertir_grupo(n))

    return ' '.join(p for p in partes if p)

def fmt_monto(cantidad_str):
    """Format amount and return (formatted_number, words)"""
    try:
        # Strip currency symbols and commas
        clean = re.sub(r'[,$\s]', '', str(cantidad_str))
        n = float(clean)
        entero = int(n)
        cents = round((n - entero) * 100)
        formatted = f"${entero:,.2f}"
        words = numero_a_letras(entero)
        if cents > 0:
            words += f" PESOS {cents:02d}/100 M.N."
        else:
            words += " PESOS 00/100 M.N."
        return formatted, words
    except:
        return cantidad_str, cantidad_str

def setup_doc():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    return doc

def p(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      size=10, space_before=0, space_after=6, indent=False):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if indent:
        para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(text)
    run.bold = run.bold or bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return para

def heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11 if level == 1 else 10)
    return para

def clausula(doc, numero, titulo, texto):
    """Add a clause with title and body"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)
    r = para.add_run(f'"{titulo}"')
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(10)

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after  = Pt(6)
    body.paragraph_format.left_indent  = Cm(0.5)
    r2 = body.add_run(f'{numero}- {texto}')
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)
    return body

def firma_line(doc, label, nombre):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(30)
    para.paragraph_format.space_after  = Pt(2)
    r = para.add_run('_' * 35)
    r.font.name = 'Arial'
    r.font.size = Pt(10)

    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_before = Pt(0)
    para2.paragraph_format.space_after  = Pt(2)
    r2 = para2.add_run(label)
    r2.bold = True
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)

    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para3.paragraph_format.space_before = Pt(0)
    para3.paragraph_format.space_after  = Pt(20)
    r3 = para3.add_run(nombre.upper())
    r3.font.name = 'Arial'
    r3.font.size = Pt(10)

# ─────────────────────────────────────────────
# CONTRATO DE ARRENDAMIENTO
# ─────────────────────────────────────────────
def generar_arrendamiento(datos, output_path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    doc = setup_doc()
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT

    # ── HELPERS ──
    def d(key, default=''):
        return str(datos.get(key, default) or default).strip()

    def u(key, default=''):
        return d(key, default).upper()

    def cp(key):
        v = d(key)
        return f"C.P. {v}" if v else ''

    def dir_completa(calle, num_ext, num_int, colonia, cp_key, mpio):
        parts = []
        c = u(calle); ne = u(num_ext); ni = u(num_int)
        col = u(colonia); cp_v = d(cp_key); mun = u(mpio)
        if c: parts.append(c)
        if ne: parts.append(ne)
        if ni: parts.append(ni)
        if col: parts.append(col)
        if cp_v: parts.append(f"C.P. {cp_v}")
        return ', '.join(parts), mun

    # ── GÉNERO ──
    sa = d('sexo_arrendador','M').upper()
    sat = d('sexo_arrendatario','M').upper()
    sos = d('sexo_os','M').upper()

    def g(sexo, m, f): return f if sexo == 'F' else m

    el_arr    = g(sa, 'EL', 'LA')
    prop_arr  = g(sa, 'PROPIETARIO', 'PROPIETARIA')
    mex_arr   = g(sa, 'MEXICANO', 'MEXICANA')
    leg_prop  = g(sa, 'LEGÍTIMO PROPIETARIO', 'LEGÍTIMA PROPIETARIA')
    el_arrt   = g(sat, 'EL', 'LA')
    mex_arrt  = g(sat, 'MEXICANO', 'MEXICANA')
    el_os_g   = g(sos, 'EL', 'LA')
    mex_os    = g(sos, 'MEXICANO', 'MEXICANA')

    # ── DATOS PRINCIPALES ──
    nombre_arr  = u('nombre_arrendador')
    nombre_arrt = u('nombre_arrendatario')
    nombre_os   = u('nombre_obligado_solidario')
    destino     = u('destino_uso')
    plazo       = u('plazo_contrato')
    MESES_PY = ['ENERO','FEBRERO','MARZO','ABRIL','MAYO','JUNIO',
                'JULIO','AGOSTO','SEPTIEMBRE','OCTUBRE','NOVIEMBRE','DICIEMBRE']

    def _fecha_escrita(v):
        """Convert ISO date or already-written date to 'DD DE MES DE YYYY'."""
        if not v: return ''
        v = v.strip()
        # Already written format e.g. '16 DE ENERO DE 2026'
        if 'DE' in v.upper(): return v.upper()
        # ISO format YYYY-MM-DD
        try:
            parts = v.split('-')
            if len(parts) == 3:
                y,m,d_n = int(parts[0]),int(parts[1]),int(parts[2])
                return f"{d_n} DE {MESES_PY[m-1]} DE {y}"
        except: pass
        return v.upper()

    fecha_firma_raw = d('fecha_contrato')
    # fecha_contrato is built as "DD DE MES DE YYYY" in JS - use as-is
    fecha_firma = fecha_firma_raw.upper() if fecha_firma_raw else ''
    fecha_ini   = _fecha_escrita(d('fecha_inicio'))
    fecha_fin_s = _fecha_escrita(d('fecha_fin'))
    mpio_inm    = u('municipio_estado_inmueble')
    estado_inm  = mpio_inm.split(',')[-1].strip() if ',' in mpio_inm else mpio_inm
    renta_num   = d('renta_mensual')
    # Letras: use JS-provided value, fallback to Python calculation
    def _letra(num_str):
        """Convert amount string to Spanish words for MXN."""
        try:
            n = int(float(str(num_str).replace(',','').replace('$','').strip()))
            return numero_a_letras(n) + ' PESOS 00/100 M.N.'
        except:
            return ''

    renta_letra_raw = u('renta_letra')
    renta_letra = renta_letra_raw if renta_letra_raw else _letra(d('renta_mensual'))

    dep_num     = d('deposito_garantia')
    dep_letra_raw = u('deposito_letra')
    dep_letra   = dep_letra_raw if dep_letra_raw else _letra(dep_num)

    forma_pago  = u('forma_pago')
    dia_pago    = d('dia_pago')
    pena_num    = d('pena_dia_num')
    pena_letra_raw = u('pena_dia_letra')
    pena_letra  = pena_letra_raw if pena_letra_raw else _letra(pena_num)

    try:
        renta_f = f"${float(str(renta_num).replace(',','').replace('$','')):,.2f}"
    except: renta_f = renta_num
    try:
        dep_f = f"${float(str(dep_num).replace(',','').replace('$','')):,.2f}"
    except: dep_f = dep_num
    try:
        pena_f = f"${float(str(pena_num).replace(',','').replace('$','')):,.2f}"
    except: pena_f = pena_num

    # Fecha de incremento INPC = día siguiente al término
    try:
        from datetime import datetime as _dt, timedelta as _td
        fin_iso = datos.get('fecha_fin_iso','') or datos.get('fecha_fin','')
        # Try ISO first, then parse written date
        fin_dt = None
        if fin_iso and '-' in fin_iso:
            fin_dt = _dt.strptime(fin_iso.split('T')[0], '%Y-%m-%d')
        if fin_dt:
            sig_dia = fin_dt + _td(days=1)
            fecha_inpc = f"{sig_dia.day} DE {MESES_PY[sig_dia.month-1]} DE {sig_dia.year}"
        else:
            fecha_inpc = fecha_fin_s
    except:
        fecha_inpc = fecha_fin_s

    # Direcciones
    dir_inm, mpio_inm2 = dir_completa('calle_inmueble','num_ext_inmueble','num_int_inmueble',
                                       'colonia_inmueble','cp_inmueble','municipio_estado_inmueble')
    dir_arr, mpio_arr  = dir_completa('calle_arrendador','num_ext_arrendador','num_int_arrendador',
                                       'colonia_arrendador','cp_arrendador','municipio_estado_arrendador')
    dir_os,  mpio_os   = dir_completa('calle_os','num_ext_os','num_int_os',
                                       'colonia_os','cp_os','municipio_estado_os')
    dir_inm_os, mpio_inm_os = dir_completa('calle_inm_os','num_ext_inm_os','num_int_inm_os',
                                            'colonia_inm_os','cp_inm_os','municipio_estado_inm_os')

    # ── HELPER: add paragraph ──
    def par(text, bold=False, underline=False, align=J, size=10,
            space_before=0, space_after=6, indent_left=0, indent_first=0):
        para = doc.add_paragraph()
        para.alignment = align
        pf = para.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        if indent_left:  pf.left_indent  = Cm(indent_left)
        if indent_first: pf.first_line_indent = Cm(indent_first)
        run = para.add_run(text)
        run.bold = bold
        run.underline = underline
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        return para

    def titulo(text):
        """Cláusula title: bold + underline, left-aligned"""
        par(text, bold=True, underline=True, align=L, size=10, space_before=10, space_after=2)

    def clausula(num_nombre, texto):
        """Número de cláusula en bold seguido del texto"""
        para = doc.add_paragraph()
        para.alignment = J
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(6)
        r1 = para.add_run(f"{num_nombre}.- ")
        r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(10)
        r2 = para.add_run(texto)
        r2.bold = False; r2.font.name = 'Arial'; r2.font.size = Pt(10)

    def sub(texto):
        par(texto, align=J, size=10, indent_left=1)

    # ══════════════════════════════════════════════
    # DOCUMENTO
    # ══════════════════════════════════════════════

    # Fecha de firma al inicio
    par(fecha_firma, bold=True, align=C, size=11, space_before=0, space_after=10)

    # Encabezado
    encab = (
        f"CONTRATO DE ARRENDAMIENTO QUE CELEBRAN POR UNA PARTE {el_arr} C. {nombre_arr}, "
        f"{prop_arr} DEL INMUEBLE UBICADO EN {dir_inm}, DE ESTA CIUDAD DE {mpio_inm}, "
        f"A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ \"LA PARTE ARRENDADORA\", Y POR LA OTRA PARTE "
        f"{el_arrt} C. {nombre_arrt}, A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ "
        f"\"LA PARTE ARRENDATARIA\", Y {el_os_g} C. {nombre_os}, A QUIEN EN LO SUCESIVO SE LE "
        f"DENOMINARÁ \"EL OBLIGADO SOLIDARIO\" SUJETÁNDOSE LAS PARTES A LAS SIGUIENTES "
        f"DECLARACIONES Y CLAUSULAS:"
    )
    par(encab, bold=True, size=11, space_before=6, space_after=10)

    # DECLARACIONES
    par('D E C L A R A C I O N E S :', bold=True, align=C, size=11, space_before=10, space_after=6)

    par(f'1.- DECLARA LA PARTE ARRENDADORA BAJO PROTESTA DE DECIR VERDAD:', bold=True, size=10, space_before=6, space_after=4)

    sub(f"1. SER {mex_arr}, MAYOR DE EDAD, ASÍ COMO TENER EL CARÁCTER Y CAPACIDAD LEGAL PARA OBLIGARSE "
        f"EN LOS TÉRMINOS DEL PRESENTE INSTRUMENTO, QUIEN SE IDENTIFICA PERSONALMENTE PARA LA FIRMA DEL "
        f"PRESENTE INSTRUMENTO MEDIANTE CREDENCIAL DE ELECTOR; EMITIDA POR EL INSTITUTO NACIONAL ELECTORAL, "
        f"QUE EN ORIGINAL EXHIBE, Y LA CUAL SE ANEXA AL PRESENTE INSTRUMENTO EN COPIA SIMPLE.")

    sub(f"ASÍ MISMO DECLARA TENER LA VOLUNTAD DE DAR EN ARRENDAMIENTO Y SER {leg_prop} Y USUFRUCTUARIO DEL "
        f"INMUEBLE UBICADO EN {dir_inm}, DE ESTA CIUDAD DE {mpio_inm}; MISMO QUE NO PRESENTA NI SUFRE VICIOS "
        f"OCULTOS O DEFECTOS Y POR LO TANTO SE ENCUENTRA EN PERFECTAS CONDICIONES DE USO Y CONSERVACIÓN "
        f"PARA SER UTILIZADO COMO {destino}.")

    sub(f"2. TENER SU DOMICILIO PARA OIR Y RECIBIR TODO TIPO DE NOTIFICACIONES, DE CARACTER LEGAL, "
        f"ADMINISTRATIVO, FISCAL ETC., DERIVADOS DEL PRESENTE ACUERDO DE VOLUNTADES EL UBICADO EN "
        f"{dir_arr}, DE ESTA CIUDAD DE {mpio_arr}.")

    par('2.- DECLARA LA PARTE ARRENDATARIA BAJO PROTESTA DE DECIR VERDAD:', bold=True, size=10, space_before=6, space_after=4)

    sub(f"1. SER {mex_arrt}, MAYOR DE EDAD, ASÍ COMO TENER EL CARÁCTER, LA VOLUNTAD Y CAPACIDAD LEGAL "
        f"PARA OBLIGARSE EN LOS TÉRMINOS DEL PRESENTE INSTRUMENTO; Y QUIEN SE IDENTIFICA PERSONALMENTE "
        f"PARA LA FIRMA DEL PRESENTE CONTRATO MEDIANTE CREDENCIAL PARA VOTAR, QUE EN ORIGINAL EXHIBE, "
        f"Y LA CUAL SE ANEXA AL PRESENTE INSTRUMENTO EN COPIA SIMPLE.")

    sub(f"2. QUE HA CONSTATADO PERSONALMENTE LAS CONDICIONES FÍSICAS Y MATERIALES Y JURÍDICAS EN LAS "
        f"CUALES SE ENCUENTRA EL INMUEBLE OBJETO DEL PRESENTE CONTRATO, LAS CUALES ENCUENTRA A SU ENTERA "
        f"SATISFACCIÓN, MISMO QUE DESEA RECIBIR EN ARRENDAMIENTO A CAMBIO DEL PAGO DE LA RENTA QUE SE "
        f"ESTIPULA. POR LO QUE CONSIDERA QUE ÉSTE REÚNE LOS REQUISITOS DE SEGURIDAD E HIGIENE SUFICIENTES "
        f"PARA UTILIZARLO INMEDIATAMENTE COMO {destino}, ASÍ MISMO MANIFIESTA QUE SE ENCUENTRA CONFORME "
        f"DE LAS MENSUALIDADES QUE HABRÁN DE CUBRIRSE POR CONCEPTO DE PAGO DE RENTA MENSUAL MISMOS "
        f"CONCEPTOS QUE PODRÁN SER SUJETOS DE VARIACIÓN ANUALMENTE, ASÍ MISMO EL PAGO DE MANTENIMIENTO "
        f"DEL FRACCIONAMIENTO, Y LOS SERVICIOS DE ENERGÍA ELÉCTRICA, TELÉFONO, PODRÁN SER VARIABLES "
        f"DE ACUERDO A SU CONSUMO.")

    sub(f"3. TENER Y SEÑALAR EN ESTE ACTO SU DOMICILIO PARA OIR Y RECIBIR TODO TIPO DE NOTIFICACIONES "
        f"PERSONALES DE CARACTER LEGAL, ADMINISTRATIVO, FISCALES, ETC., PARA LOS EFECTOS DERIVADOS DEL "
        f"PRESENTE INSTRUMENTO EL UBICADO EN {dir_inm}, DE ESTA CIUDAD DE {mpio_inm}; ASÍ COMO EL "
        f"DOMICILIO SEÑALADO POR EL OBLIGADO SOLIDARIO.")

    par('3.- DECLARA EL OBLIGADO SOLIDARIO BAJO PROTESTA DE DECIR VERDAD:', bold=True, size=10, space_before=6, space_after=4)

    sub(f"3.1.- SER {mex_os}, MAYOR DE EDAD, ASÍ COMO TENER EL CARÁCTER, LA VOLUNTAD Y CAPACIDAD LEGAL "
        f"PARA OBLIGARSE EN LOS TÉRMINOS DEL PRESENTE INSTRUMENTO, CON SOLVENCIA MORAL Y ECONOMICA PARA "
        f"DAR CUMPLIMIENTO AL MISMO EN SU CARÁCTER DE OBLIGADO SOLIDARIO, ADEMÁS DE SER EN ESTE MOMENTO "
        f"PROPIETARIO DEL INMUEBLE UBICADO EN {dir_inm_os}, DE ESTA CIUDAD DE {mpio_inm_os}, MISMA "
        f"PROPIEDAD QUE GARANTIZA SU SOLVENCIA ECONÓMICA PARA OBLIGARSE EN LOS TÉRMINOS DE ESTE CONTRATO; "
        f"QUIEN SE IDENTIFICA A LA FIRMA DEL PRESENTE CONTRATO MEDIANTE CREDENCIAL PARA VOTAR, QUE EN "
        f"ORIGINAL EXHIBE, Y LA CUAL SE ANEXA AL PRESENTE INSTRUMENTO EN COPIA SIMPLE.")

    sub(f"3.2.- TENER Y SEÑALAR SU DOMICILIO PARA OIR Y RECIBIR TODO TIPO DE NOTIFICACIONES PERSONALES "
        f"DE CARACTER LEGAL, ADMINISTRATIVO, FISCALES, ETC., PARA LOS EFECTOS DERIVADOS DEL PRESENTE "
        f"INSTRUMENTO, EL UBICADO EN {dir_os}, DE ESTA CIUDAD DE {mpio_os}; ASÍ MISMO EL SEÑALADO POR "
        f"LA PARTE ARRENDATARIA.")

    par("Obligándose las partes a informar por escrito con anticipación cualquier cambio de domicilio y "
        "en caso de no hacerlo acuerdan que surtirá efecto legal cualquier comunicación, notificación, "
        "diligencia etc. que se les haga en los domicilios señalados.", size=10)

    par("LAS PARTES DECLARAN QUE ES SU VOLUNTAD OBLIGARSE RECIPROCAMENTE EN ESTE ACTO AL TENOR DE LAS SIGUIENTES:",
        bold=True, align=C, size=11, space_before=10, space_after=6)

    par('C L A U S U L A S :', bold=True, align=C, size=11, space_before=0, space_after=6)

    # ── CLÁUSULAS ──

    titulo('"OBJETO"')
    clausula("PRIMERA",
        "La parte arrendadora, en este acto, entrega en arrendamiento a la parte arrendataria "
        "y esta recibe de conformidad, a su entera satisfacción y bajo ese título el inmueble "
        "descrito en la declaración 1.1. En buen estado físico de conservación para servir al "
        "uso convenido. Incluyendo sus 2 respectivos cajones de estacionamiento y bodega. No "
        "reservándose la parte arrendataria derecho alguno que hacer valer ni en lo presente "
        "ni en lo futuro por este caso a la parte arrendadora.")

    titulo('"TERMINO CONTRATO"')
    clausula("SEGUNDA",
        f"El término del contrato de arrendamiento es por {plazo}, obligatorio para ambas partes. "
        f"Debiendo acordar por escrito un nuevo término en caso de querer continuar con el "
        f"arrendamiento. Iniciando el término antes citado el {fecha_ini} y finalizando el "
        f"{fecha_fin_s}. Con derecho de prórroga siempre y cuando la parte arrendataria se "
        f"encuentre al corriente del pago de las rentas y servicios. Debiendo la parte arrendataria "
        f"dar aviso en un plazo no mayor a 30 días antes de la fecha de vencimiento del presente "
        f"instrumento, su deseo de continuar con el arrendamiento, para que pueda ser valorado si "
        f"existen las condiciones para que sea renovado.")

    sub("No se entenderá renovado en términos de ley este contrato, ni en cuanto a la forma de "
        "pago, por el hecho de que la parte arrendadora reciba las prestaciones de renta u otras "
        "adicionales distintas a las estipuladas contractualmente, o admita abonos a cuenta de la misma.")

    clausula("TERCERA",
        f"Acuerdan las partes que en caso de que la parte arrendataria no cumpla con el término de "
        f"{plazo}, o incurra en una de las causas de rescisión del presente instrumento, deberá "
        f"pagarle a la parte arrendadora, una penalización equivalente a un mes de renta. "
        f"Independientemente de los meses de renta vencidos a la fecha que esto suceda.")

    titulo('"PRECIO RENTA"')
    clausula("CUARTA",
        f"La parte arrendataria se obliga a pagarle puntualmente, sin requerimiento previo alguno, "
        f"a la parte arrendadora el importe de la renta del inmueble objeto del presente instrumento, "
        f"por la cantidad de: {renta_f} ({renta_letra}) ya con mantenimiento incluido, por mes a "
        f"transcurrir, en esta ciudad de {mpio_inm}, mediante {forma_pago}. Pagaderos a más tardar "
        f"los días {dia_pago} de cada mes. Obligándose la parte arrendataria a mantener al corriente "
        f"los pagos de agua, luz, e internet, así como cualquier otro adeudo que se derive de la "
        f"ocupación del inmueble objeto de este contrato.")

    sub(f"Acuerdan ambas partes que el importe del pago de la renta del inmueble aumentará cada año "
        f"sin previo aviso a la parte arrendataria, de acuerdo al Índice Nacional de Precios al "
        f"Consumidor, es decir que ese incremento entrará en vigor el día {fecha_inpc} y así sucesivamente.")

    titulo('"FORMA DE PAGO"')
    clausula("QUINTA",
        f"La parte arrendataria se obliga con la parte arrendadora a cumplir puntualmente con el "
        f"pago de las rentas mensuales, así como cualquier otra prestación que se derive del "
        f"presente instrumento, en esta ciudad de {mpio_inm}. Los pagos de la renta deberán "
        f"realizarse en moneda nacional.")

    sub("En caso de incumplimiento en el pago puntual de las rentas o de alguna otra prestación "
        "inherente a este instrumento, ocasionará en perjuicio de la parte arrendataria la obligación "
        "de pagar a la parte arrendadora el 10% diez por ciento MENSUAL de interés moratorio respecto "
        "del importe de la renta que se encuentre vigente en ese momento, desde la constitución en mora "
        "y hasta la total liquidación de todas y cada una de las obligaciones contraídas.")

    sub("La parte arrendataria no podrá retener la renta en ningún caso ni bajo ningún título judicial "
        "o extrajudicial, ni por falta de composturas ni reparaciones que la parte arrendadora hiciere "
        "sino que la pagará íntegramente y en la fecha estipulada cumpliendo además las obligaciones "
        "que previenen el Código Civil del Estado de Michoacán, en cuanto a arrendatario le competen.")

    titulo('"DESTINO, OBJETO"')
    clausula("SEXTA",
        f"La parte arrendataria deberá destinar únicamente el inmueble arrendado exclusivamente para "
        f"fines y objeto de {destino} en el caso de variar el fin y objeto será motivo de rescisión "
        f"SIN PREVIO AVISO del presente contrato de arrendamiento. Además que desde la suscripción de "
        f"este contrato la parte arrendataria deslinda a la parte arrendadora de cualquier "
        f"responsabilidad de cualquier índole, que pudiera derivarse del uso indebido del inmueble, "
        f"entre ellas la realización de actividades ilícitas. Y, para este caso la parte arrendataria "
        f"se obliga a no realizar actividades ilícitas dentro del inmueble arrendado, que pongan en "
        f"peligro la propiedad, ni tampoco a desarrollar actividades catalogadas como ilícitas en la "
        f"Ley Nacional de Extinción de Dominio, por lo que se hará responsable de pagar el valor "
        f"comercial del inmueble, en caso de que el inmueble se vea afectado por dicha ley.")

    titulo('"RESTRICCIONES"')
    clausula("SÉPTIMA",
        "La parte arrendataria no podrá realizar en el inmueble arrendado modificaciones ni obras sin "
        "el consentimiento por escrito del arrendador. En caso de ser autorizadas, dichas obras serán "
        "hechas a costa de la parte arrendataria y sin que tenga derecho a compensación o remuneración "
        "alguna, quedando dicha obra en beneficio del inmueble al terminar el arrendamiento o de ser "
        "voluntad del arrendador deberán de ser retiradas y el inmueble deberá de ser reparado para "
        "dejarlo en su estado original. Si la parte arrendataria instala en algunas puertas chapas "
        "adicionales o sustituya las que existan al recibir el inmueble, al desocupar el mismo no "
        "podrá retirarlas, quedando estas a beneficio del inmueble arrendado.")

    titulo('"RESCISION"')
    clausula("OCTAVA",
        "El incumplimiento de cualquiera de las obligaciones contraídas por la parte arrendataria "
        "en este contrato, la falta de pago de una o más rentas vencidas, será causa de rescisión "
        "del mismo, bastando tan sólo que la parte arrendadora notifique por escrito a la parte "
        "arrendataria con una semana de anticipación su deseo de dar por rescindido el contrato, "
        "precisando la razón o motivo de esta causal, o bien mencionar el incumplimiento de alguna "
        "de las cláusulas establecidas en el presente instrumento contractual. Independientemente "
        "de que la parte arrendadora pueda ejercitar las acciones legales necesarias para que la "
        "parte arrendataria restablezca el inmueble al estado en que lo recibió y reclamar legalmente "
        "el pago de los daños y perjuicios ocasionados por su incumplimiento.")

    sub(f"Si la parte arrendataria desea rescindirlo después de {plazo} se avisará a la parte "
        f"arrendadora con 30 días de anticipación y sin penalización alguna.")

    clausula("NOVENA",
        "La parte arrendataria NO podrá subarrendar en parte el inmueble arrendado y no podrá ceder "
        "ni traspasar en forma alguna los derechos y obligaciones adquiridos en este contrato, sin "
        "previo consentimiento dado por escrito por la parte arrendadora. En caso de que la parte "
        "arrendataria haga caso omiso a esta restricción contractual, será motivo de rescisión.")

    sub("Así mismo, es causa de rescisión del contrato de arrendamiento el concurso mercantil al cual "
        "sea sometido la parte arrendataria. También lo será el hecho de que la parte arrendataria "
        "explote o use de manera distinta a la prevista en este contrato. Además de la falta de pago "
        "de uno o más rentas acordadas. Así como que la parte arrendataria de manera parcial o total, "
        "onerosa o gratuita, ceda los derechos consagrados en este instrumento.")

    clausula("DÉCIMA",
        "Expresamente se estipula que la parte arrendataria no podrá almacenar sustancias peligrosas, "
        "corrosivas, deletéreas o inflamables en el inmueble arrendado que puedan producir incendio u "
        "explosión etc. En caso de producirse siniestro en el inmueble arrendado por contravenir lo "
        "dispuesto en esta cláusula, la parte arrendataria deberá cubrir a la parte arrendadora todos "
        "los daños y perjuicios que le ocasione por su incumplimiento, asumiendo también el pago total "
        "de los daños y perjuicios que se ocasionen a terceros, como lo previene el Código Civil.")

    titulo('"RESPONSABILIDADES"')
    clausula("DECIMOPRIMERA",
        "La parte arrendadora no será responsable de la seguridad de los bienes muebles que introduzca "
        "la parte arrendataria al inmueble arrendado, por lo tanto no se podrá culpar a la parte "
        "arrendadora en ningún caso por los robos o daños surgidos en bienes propiedad de la parte "
        "arrendataria. La parte arrendataria queda obligado a salvaguardar sus pertenencias dentro "
        "del inmueble arrendado, quedando facultado desde este momento a cambiar la combinación de "
        "las chapas y a colocar chapas adicionales al inmueble para su propia protección. Así mismo, "
        "la parte arrendataria queda obligada a informar a la parte arrendadora de toda usurpación o "
        "novedad dañosa que otro haya hecho al inmueble, bajo pena de pagar los daños y perjuicios "
        "que su omisión cause.")

    sub("La parte arrendataria asume de manera enunciativa más no limitativa, toda la responsabilidad "
        "civil, laboral, penal, fiscal o de cualquier otra naturaleza, eximiendo a la parte arrendadora "
        "de todo género de responsabilidad derivada de sus actividades y de la ocupación del inmueble.")

    sub("Nada en este contrato será considerado o interpretado para constituir a las partes como socios, "
        "agentes, empleados uno del otro y ninguna de las disposiciones de este contrato será "
        "interpretado para arrendador o a su representante legal para que en caso de que fuera "
        "obligatorio presente el aviso a que se refiere el artículo 23 y 24 de la Ley de referencia. "
        "Y, que de igual manera, en términos de la fracción III del artículo 3 de la Ley Federal para "
        "la Prevención e Identificación de Operaciones con Recursos de Procedencia Ilícita, respecto "
        "a concepto de beneficiario controlador y los de presunción de existencia, manifiesta que no "
        "tiene conocimiento ni existen beneficiarios distintos a los comparecientes, ya que en el "
        "presente instrumento actúa en su nombre y por cuenta propia, ello por ser quien se beneficia "
        "de los actos que en el mismo se contienen y quien ejerce los derechos de uso, goce y disfrute, "
        "aprovechamiento o disposición del objeto del presente convenio, siendo él mismo quien ejerce "
        "los actos establecidos en el inciso b) del precepto legal en cita.")

    titulo("''DEVOLUCIÓN DEL INMUEBLE''")
    clausula("DECIMOCUARTA",
        "Independientemente de la causa de rescisión del contrato o por su terminación, la parte "
        "arrendataria queda obligada a hacer la devolución del inmueble, de manera personal, "
        "entregándolo en buen estado de conservación y funcionamiento en que le fue entregado, ya "
        "que en caso contrario la parte arrendataria deberá cubrir a la parte arrendadora todas las "
        "reparaciones, composturas, reposiciones o adecuaciones que sean necesarias para devolver el "
        "inmueble al estado en el que estaba, después de los desperfectos y mal uso que haya hecho "
        "la parte arrendataria al inmueble; además de todos los daños y perjuicios que le ocasione ésta.")

    clausula("DECIMOQUINTA",
        "Así mismo, se obliga la parte arrendataria a entregar el inmueble al corriente de los pagos "
        "en los servicios de agua, luz, gas, internet o cualquier otro servicio que derive de la "
        "ocupación del inmueble.")

    sub(f"Para el caso en que al término de la vigencia forzosa del presente instrumento, no se elabore "
        f"un nuevo contrato de arrendamiento, la parte arrendataria deberá entregar el inmueble "
        f"personalmente y totalmente desocupado a la parte arrendadora, a más tardar el {fecha_fin_s}. "
        f"Por lo que si no lo hace, se obliga a pagar la cantidad de {pena_f} ({pena_letra}) diarios, "
        f"hasta que la desocupe y entregue, como pena convencional por la no devolución del mismo.")

    titulo('"DEPOSITO"')
    clausula("DECIMOSEXTA",
        f"La parte arrendataria entregará a la parte arrendadora, la cantidad de {dep_f} ({dep_letra}), "
        f"por concepto de DEPÓSITO EN GARANTÍA, sin que el mismo genere intereses y sin que pueda ser "
        f"aplicado a ninguna mensualidad por concepto de renta para garantizar las obligaciones a su "
        f"cargo emanadas del presente contrato, sirviendo el presente instrumento como el comprobante "
        f"de pago más amplio que en derecho corresponda. La cantidad depositada será devuelta a los 30 "
        f"(treinta) días después de finalizar el contrato de arrendamiento siempre y cuando se haya "
        f"cumplido la totalidad de las obligaciones de renta y de entrega del inmueble. En caso "
        f"contrario, la cantidad en depósito se aplicará al pago de los adeudos que correspondan "
        f"hasta donde fueran suficientes.")

    sub("LA PARTE ARRENDATARIA, autoriza a LA PARTE ARRENDADORA, para que en caso de desocupación "
        "anticipada de EL INMUEBLE arrendado, pueda este último disponer del depósito de garantía "
        "para hacer los arreglos necesarios y de pintura de EL INMUEBLE.")

    clausula("DECIMOSÉPTIMA",
        "Cualquier pago que efectúe LA PARTE ARRENDATARIA a favor de LA PARTE ARRENDADORA se "
        "aplicará, primeramente, a cubrir los gastos que erogue LA PARTE ARRENDADORA y que "
        "correspondan a LA PARTE ARRENDATARIA en los términos del presente contrato, después serán "
        "imputados al pago de los intereses moratorios, y por último, al pago de las rentas "
        "generadas y no cubiertas.")

    titulo("''DERECHO DEL TANTO Y TRANSMISIÓN DE LA PROPIEDAD''")
    clausula("DECIMOCTAVA",
        "La parte arrendataria renuncia expresamente al derecho de preferencia o derecho del tanto, "
        "es decir, para la compra del inmueble.")

    clausula("DECIMONOVENA",
        "Si durante la vigencia del contrato de arrendamiento se verificare la transmisión de la "
        "propiedad inmueble arrendado, en virtud de que la parte arrendataria renunció a su derecho "
        "de preferencia o del tanto, en los términos establecidos en este contrato. El arrendamiento "
        "subsistirá en los mismo términos que establece el presente contrato, respecto al pago de "
        "las rentas la parte arrendataria tendrá la obligación de pagar al nuevo propietario la renta "
        "estipulada en el contrato, desde la fecha en que se le notifique judicial o extrajudicialmente "
        "ante fedatario público y/o ante dos testigos, haberse otorgado el correspondiente título "
        "de propiedad.")

    titulo('"OBLIGADO SOLIDARIO"')
    clausula("VIGÉSIMA",
        f"El obligado solidario se constituye como responsable de todas y cada una de las obligaciones "
        f"contraídas por la parte arrendataria, haciendo todas las renuncias que la parte arrendataria "
        f"tiene hechas, y los beneficios que de orden y exclusión consignadas en el Código Civil del "
        f"estado de {estado_inm}, no cesando la responsabilidad de este sino hasta cuando la parte "
        f"arrendadora se dé por recibido de la localidad de todo cuanto se le deba, por virtud de este "
        f"contrato aun cuando el arrendamiento haya concedido prórrogas o esperas subsistiendo la "
        f"obligación del obligado solidario a pesar de que no se le notifique. Así como se obliga a "
        f"hacer la entrega si la parte arrendataria no lo hiciere del inmueble así como lo que éste "
        f"haya recibido en el inventario y de reponer lo que le faltare, pagando el costo de los "
        f"desperfectos que por mal uso fueren causados por la parte arrendataria.")

    sub("Si la parte arrendataria no cumpliese con entregar el inmueble en las condiciones en que le "
        "fue entregado, el obligado solidario queda obligado a realizar el mantenimiento al inmueble "
        "para que éste quede en el mismo estado de conservación y funcionamiento en que fue entregado "
        "a la parte arrendataria. Quien en señal de aceptación del cargo firma en compañía de las partes.")

    titulo('"CONFIDENCIALIDAD"')
    clausula("VIGESIMOPRIMERA",
        "Las partes se obligan a mantener de forma confidencial toda la información y documentación "
        "relativa al presente instrumento y a la operación que prometen llevar a cabo, a no divulgar "
        "a terceros sin el consentimiento previo y por escrito de cualquiera de ellas. La obligación "
        "de confidencialidad aquí establecida no aplicará respecto de aquella información que por su "
        "naturaleza se encuentre o hubiere estado en dominio público por algún motivo que no constituya "
        "un acto u omisión de cualquiera de las partes.")

    titulo('"COMPETENCIA LEGAL CONTRACTUAL"')
    clausula("VIGESIMOSEGUNDA",
        "Si cualquier parte de este contrato se considera inválida o no exigible por un tribunal "
        "competente, entonces en la medida en que sea razonable y posible, las demás partes de este "
        "contrato se considerarán válidas y exigibles, y se dará efecto a la intención manifestada "
        "en la parte inválida o no exigible. La falta de cualquiera de las partes de exigir contra "
        "la otra los términos y condiciones de este contrato no se considerarán como una renuncia al "
        "derecho de dicha parte de reclamar a la otra tal término o estipulación o cualquier otro.")

    clausula("VIGESIMOTERCERA",
        f"Para todas las cuestiones relativas al alcance de la interpretación y cumplimiento de las "
        f"obligaciones y derechos que se consignan en este contrato, las partes contratantes se someten "
        f"expresamente a las leyes y a los tribunales competentes en la ciudad de {mpio_inm}, "
        f"renunciando al fuero que por sus domicilios actuales o futuros o que por cualquier otra "
        f"razón pudiera corresponderles. Conviniendo que serán a cargo de la parte arrendataria, "
        f"todos los gastos y costas judiciales y extrajudiciales a que dieran lugar por incumplimiento "
        f"del contrato en caso de controversia judicial.")

    titulo("''TÉRMINOS''")
    for termino in [
        "Encabezados. Todos los encabezados de las cláusulas del presente contrato son para fines de "
        "conveniencia y no modifican, definen o limitan, de modo alguno, los términos o disposiciones "
        "que en ellas se contienen.",
        "Pronombres y términos. En el presente contrato el singular incluirá el plural y el plural "
        "incluirá el singular, y el uso de cualquier género será aplicable a todos los géneros.",
        "Entendimiento único. El presente contrato contiene el acuerdo total y completo de las partes, "
        "quienes acuerdan sujetarse a los términos y condiciones establecidos en el mismo, y sustituye "
        "y deja sin efecto cualquier convenio o negociación previa oral o escrita entre las partes.",
        "Obligatoriedad. Este contrato, y los respectivos derechos y obligaciones de las partes serán "
        "obligatorios y tendrán efecto para el beneficio de las partes y de sus respectivos sucesores, "
        "representantes y cesionarios permitidos.",
        "Modificaciones. Este contrato no podrá ser modificado, cambiado o terminado salvo que exista "
        "un convenio por escrito y firmado por las partes.",
        "Acuerdos adicionales. Cada uno de los contratantes acuerda firmar y entregar a la contraparte, "
        "cualquier documento adicional que se requiera y/o llevar a cabo cualquier acción que pueda ser "
        "necesaria para consumar en una forma más efectiva los propósitos y objetivos del presente contrato.",
        "Gastos. Cada una de las partes asumirá todos los gastos relacionados con la preparación y "
        "celebración del presente contrato.",
        "Impuestos. Las partes por este medio convienen en que cada parte será responsable de pagar sus "
        "respectivos impuestos que se causen en relación con este contrato.",
    ]:
        sub(termino)

    # ── CLÁUSULAS ESPECIALES NUMERADAS ──
    ORDINALES = ['VIGESIMOCUARTA','VIGESIMOQUINTA','VIGESIMOSEXTA','VIGESIMOSÉPTIMA',
                 'VIGESIMOCTAVA','VIGESIMONOVENA','TRIGÉSIMA','TRIGESIMOPRIMERA',
                 'TRIGESIMOSEGUNDA','TRIGESIMOTERCERA']
    clausulas_esp = datos.get('clausulas_especiales', [])
    import re as _re
    for idx_cl, cl_text in enumerate(clausulas_esp):
        if not cl_text.strip(): continue
        ordinal = ORDINALES[idx_cl] if idx_cl < len(ORDINALES) else f'CLÁUSULA {idx_cl+24}'
        lines = [l.strip() for l in cl_text.strip().split('\n') if l.strip()]
        # Extract subject from AI header line if present
        subject = ''
        body_lines = lines
        m = _re.match(r'^[A-ZÁÉÍÓÚÑ\s]+\.-\s*([A-ZÁÉÍÓÚÑ\s]+)\.-\s*(.*)', lines[0])
        if m:
            subject = m.group(1).strip()
            rest = m.group(2).strip()
            body_lines = ([rest] if rest else []) + lines[1:]
        if subject:
            titulo(f'"{subject}"')
        clausula(ordinal, ' '.join(body_lines))

    # ── CIERRE ──
    par('', space_before=10, space_after=4)
    par(f"Manifestando ambas partes bajo protesta de decir verdad, que en este contrato no existe "
        f"dolo, lesión, mala fe, o algún vicio del consentimiento que pueda afectarle de nulidad, "
        f"lo leen y habiendo quedado plenamente enteradas del contenido y los alcances legales de "
        f"todas y cada una de las cláusulas de este contrato, lo firman de absoluta conformidad por "
        f"duplicado, al margen de cada página anterior y al calce de ésta, en la ciudad de "
        f"{mpio_inm}, a {fecha_firma}.", size=10)

    # ── FIRMAS EN RECUADROS ──
    par('', space_before=20, space_after=4)

    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def firma_box(nombre, titulo_firma):
        """Add a signature box as a table cell"""
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        cell.width = Inches(2.8)
        # Title
        tp = cell.paragraphs[0]
        tp.alignment = C
        r = tp.add_run(titulo_firma)
        r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(9)
        # Signature line
        p2 = cell.add_paragraph()
        p2.alignment = C
        p2.paragraph_format.space_before = Pt(30)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run('_' * 30)
        r2.font.name = 'Arial'; r2.font.size = Pt(10)
        # Name
        p3 = cell.add_paragraph()
        p3.alignment = C
        r3 = p3.add_run(f'C. {nombre}')
        r3.bold = True; r3.font.name = 'Arial'; r3.font.size = Pt(9)
        return table

    # Row 1: Arrendador centered
    firma_box(nombre_arr, 'LA PARTE ARRENDADORA')
    par('', space_before=8, space_after=4)

    # Row 2: Arrendatario + OS side by side
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = 'Table Grid'
    for col_idx, (nom, tit) in enumerate([
        (nombre_arrt, 'LA PARTE ARRENDATARIA'),
        (nombre_os,   'EL OBLIGADO SOLIDARIO'),
    ]):
        cell = tbl2.cell(0, col_idx)
        tp = cell.paragraphs[0]
        tp.alignment = C
        r = tp.add_run(tit)
        r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(9)
        p2 = cell.add_paragraph()
        p2.alignment = C
        p2.paragraph_format.space_before = Pt(30)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run('_' * 28)
        r2.font.name = 'Arial'; r2.font.size = Pt(10)
        p3 = cell.add_paragraph()
        p3.alignment = C
        r3 = p3.add_run(f'C. {nom}')
        r3.bold = True; r3.font.name = 'Arial'; r3.font.size = Pt(9)

    doc.save(output_path)
    print(f"✓ Contrato de arrendamiento generado: {output_path}")



def generar_promesa(d, output_path):
    doc = setup_doc()

    precio_num, precio_letra  = fmt_monto(d['precio_total'])
    arras_num, arras_letra    = fmt_monto(d['monto_arras'])
    saldo_num, saldo_letra    = fmt_monto(d['monto_saldo'])
    pena_num, pena_letra      = fmt_monto(d.get('pena_convencional', d['monto_arras']))

    fecha         = d['fecha_contrato']
    nombre_vend   = d['nombre_vendedor'].upper()
    nombre_comp   = d['nombre_comprador'].upper()
    dir_inmueble  = d['direccion_inmueble'].upper()
    col_inmueble  = d['colonia_inmueble'].upper()
    cp_inmueble   = d['cp_inmueble']
    escritura_num = d.get('escritura_numero', '___')
    notario_nombre= d.get('notario_nombre', '___')
    notario_num   = d.get('notario_numero', '___')
    tomo          = d.get('tomo_registro', '___')
    registro      = d.get('registro', '___')
    dom_vend      = d['domicilio_vendedor'].upper()
    dom_comp      = d['domicilio_comprador'].upper()
    fecha_limite  = d['fecha_limite_escritura']
    forma_pago    = d.get('forma_pago_saldo', 'efectivo').lower()

    # ── ENCABEZADO ──
    heading(doc, "CONTRATO PRIVADO DE PROMESA DE COMPRAVENTA DE BIEN INMUEBLE")
    doc.add_paragraph()

    p(doc,
      f"CONTRATO PRIVADO DE PROMESA DE COMPRAVENTA QUE CELEBRAN POR UNA PARTE "
      f"{nombre_vend}, PROPIETARIO DEL INMUEBLE UBICADO EN {dir_inmueble}, "
      f"COLONIA {col_inmueble}, CÓDIGO POSTAL {cp_inmueble}, CORRESPONDIENTE AL "
      f"MUNICIPIO DE MORELIA, MICHOACÁN, A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ "
      f"\"EL PROMITENTE VENDEDOR\", Y POR LA OTRA PARTE {nombre_comp}, A QUIEN EN "
      f"LO SUCESIVO SE LE DENOMINARÁ \"EL PROMITENTE COMPRADOR\", SUJETÁNDOSE LAS "
      f"PARTES A LAS SIGUIENTES DECLARACIONES Y CLÁUSULAS:",
      bold=True)

    # ── DECLARACIONES ──
    p(doc, "- - - - - - - - - - - - - D E C L A R A C I O N E S - - - - - - - - - - - - -",
      align=WD_ALIGN_PARAGRAPH.CENTER)

    p(doc,
      f"I.- Declara EL PROMITENTE VENDEDOR, bajo protesta de decir verdad, ser mexicano(a), "
      f"mayor de edad, que es su voluntad celebrar este contrato promisorio y en su oportunidad "
      f"el contrato definitivo respectivo, que tiene las facultades necesarias y que no tiene "
      f"ningún impedimento legal para vender y quien se identifica con su credencial para votar "
      f"emitida por el Instituto Nacional Electoral, que en original exhibe y que se anexa al "
      f"presente instrumento en copia simple.\n\n"
      f"Así mismo, declara bajo protesta de decir verdad, ser el legítimo propietario del "
      f"INMUEBLE UBICADO EN {dir_inmueble}, COLONIA {col_inmueble}, CÓDIGO POSTAL {cp_inmueble}, "
      f"CORRESPONDIENTE AL MUNICIPIO DE MORELIA, MICHOACÁN, lo que demuestra con la escritura "
      f"pública número {escritura_num} pasada ante la fe del {notario_nombre}, notario público "
      f"número {notario_num} en el estado de Michoacán, y debidamente inscrita en el Registro "
      f"Público de la Propiedad bajo el tomo {tomo} y registro {registro} del libro de propiedad; "
      f"que este se encuentra libre de todo gravamen y que no existe impedimento legal alguno para "
      f"vender dicho inmueble.\n\n"
      f"Así mismo EL PROMITENTE VENDEDOR señala como domicilio para recibir cualquier tipo de "
      f"notificación el ubicado en {dom_vend}, CORRESPONDIENTE AL MUNICIPIO DE MORELIA, MICHOACÁN.")

    p(doc,
      f"II.- Declara EL PROMITENTE COMPRADOR, bajo protesta de decir verdad, ser mexicano(a), "
      f"mayor de edad, que es su voluntad celebrar este contrato promisorio y en su oportunidad "
      f"el contrato definitivo respectivo, que tiene las facultades necesarias para comprar, que "
      f"conoce el estado físico y jurídico del inmueble objeto de este contrato y los acepta, y "
      f"quien se identifica con credencial para votar emitida por el Instituto Nacional Electoral, "
      f"que en original exhibe y que se anexa en copia simple al presente instrumento.\n\n"
      f"Además, bajo protesta de decir verdad, manifiesta que los recursos con los que pretende "
      f"adquirir el inmueble objeto de este contrato, son de procedencia lícita.\n\n"
      f"Así mismo señala como domicilio para recibir y oír notificaciones el ubicado en "
      f"{dom_comp}, CORRESPONDIENTE AL MUNICIPIO DE MORELIA, MICHOACÁN.")

    p(doc,
      "III.- Declaran LAS PARTES, bajo protesta de decir verdad, que se reconocen la identidad "
      "de acuerdo a las identificaciones que se describen anteriormente y que se exhiben el uno "
      "al otro en original, que es su voluntad sujetarse en los términos del presente instrumento "
      "y que en este contrato no existe dolo, mala fe, vicios en el consentimiento, ni ningún otro "
      "que lo invalide. Además declaran que no obtienen enriquecimiento ilegítimo.")

    p(doc, "- - - - - - - - - - - - - C L Á U S U L A S - - - - - - - - - - - - -",
      align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── CLÁUSULAS ──
    clausula(doc, "PRIMERA.-", "OBJETO",
        f"El C. {nombre_vend} promete VENDER, y el C. {nombre_comp} promete COMPRAR para sí, "
        f"el inmueble descrito en la declaración I, en el estado físico en que se encuentra, "
        f"que EL PROMITENTE VENDEDOR entregará libre de gravamen, al corriente en sus pagos "
        f"de servicios e impuestos. Así mismo ambos se obligan a celebrar contrato definitivo "
        f"de compraventa ante la fe de un notario público.")

    clausula(doc, "SEGUNDA.-", "PRECIO Y FORMA DE PAGO",
        f"El contrato definitivo de compraventa tendrá un precio pactado de {precio_num} "
        f"({precio_letra}), mismo que será cubierto de la siguiente forma:\n\n"
        f"A) A la firma del presente contrato la cantidad de {arras_num} ({arras_letra}) en "
        f"efectivo, cantidad que será recibida como depósito a título de arras para que en su "
        f"caso dicha cantidad sea aplicada como parte del pago del precio.\n\n"
        f"B) La cantidad de {saldo_num} ({saldo_letra}) mediante {forma_pago} a más tardar "
        f"el {fecha_limite}, previo a la firma de la escritura que certifique el contrato de "
        f"compraventa o simultáneamente a esta.")

    clausula(doc, "TERCERA.-", "ESCRITURA",
        f"Ambas partes aceptan, entienden y se obligan a que la firma de la escritura pública "
        f"que certifique la compraventa sobre el inmueble objeto del presente contrato se celebre "
        f"a más tardar el {fecha_limite}. EL PROMITENTE VENDEDOR se reserva el dominio y "
        f"propiedad del inmueble materia del presente contrato hasta que hayan recibido el importe "
        f"total del precio.")

    clausula(doc, "CUARTA.-", "RESCISIÓN",
        "Serán causas de rescisión del presente instrumento, si alguna de las declaraciones hechas "
        "por las partes resultan falsas; que alguna de las partes no entregue en su totalidad la "
        "documentación requerida para formalizar la compraventa; si la documentación entregada al "
        "notario que formalizará la compraventa es contraria a derecho o falsa; si no se formalizara "
        "el contrato de compraventa a más tardar a la fecha pactada por las partes.")

    clausula(doc, "QUINTA.-", "PENA CONVENCIONAL",
        f"En caso de rescisión del presente contrato por causas imputables a EL PROMITENTE COMPRADOR, "
        f"pagará a EL PROMITENTE VENDEDOR por concepto de pena convencional, la cantidad de {pena_num} "
        f"({pena_letra}), a más tardar 3 días naturales posteriores a la notificación de su "
        f"incumplimiento; mismos que podrán ser pagados con el depósito a título de arras.\n\n"
        f"En el caso de rescisión por causas imputables a EL PROMITENTE VENDEDOR, deberá pagar a "
        f"EL PROMITENTE COMPRADOR la cantidad de {pena_num} ({pena_letra}), a más tardar 3 días "
        f"posteriores a la notificación de su incumplimiento, además de devolver íntegramente todas "
        f"las cantidades que le hayan sido entregadas.")

    clausula(doc, "SEXTA.-", "GASTOS E IMPUESTOS",
        "Acuerdan los contratantes que los gastos, impuestos, derechos y honorarios que se originen "
        "con motivo de la escritura definitiva de compraventa correrán por parte de EL PROMITENTE "
        "COMPRADOR, a excepción del impuesto sobre la renta que en caso de generarse, lo cubrirá "
        "EL PROMITENTE VENDEDOR.")

    clausula(doc, "SÉPTIMA.-", "CONFIDENCIALIDAD",
        "Las partes se obligan a mantener de forma confidencial toda la información y documentación "
        "relativa al presente instrumento y a la operación que prometen llevar a cabo, a no divulgar "
        "a terceros sin el consentimiento previo y por escrito de cualquiera de ellas.")

    clausula(doc, "OCTAVA.-", "VALIDEZ",
        "Si cualquier parte de este contrato se considera inválida o no exigible por un tribunal "
        "competente, las demás partes de este contrato se considerarán válidas y exigibles. "
        "La falta de cualquiera de las partes de exigir los términos y condiciones de este contrato "
        "no se considerarán como una renuncia al derecho de dicha parte de reclamarlos.")

    clausula(doc, "NOVENA.-", "JURISDICCIÓN",
        "Para la interpretación y cumplimiento de cualquier controversia que se pudiera suscitar "
        "con motivo de cumplimiento de las obligaciones que las partes contraen en este contrato, "
        "ambos se someten expresamente a la jurisdicción y tribunales competentes de la ciudad de "
        "Morelia, Michoacán, renunciando a cualquier fuero presente o futuro que les pudiera "
        "corresponder por razón de domicilio.")

    # ── CIERRE ──
    doc.add_paragraph()
    p(doc,
      f"Manifestando ambas partes bajo protesta de decir verdad, que en este contrato no existe "
      f"dolo, mala fe, o algún vicio del consentimiento que pueda afectarle de nulidad, lo leen y "
      f"habiendo quedado enteradas del contenido y los alcances legales de todas y cada una de las "
      f"cláusulas de este contrato, lo firman por duplicado, al margen de cada página anterior y "
      f"al calce de esta, en la ciudad de Morelia, Michoacán, a {fecha}.",
      bold=True)

    # ── FIRMAS ──
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    for cell in table.rows[0].cells:
        cell.width = Cm(9)

    cells = table.rows[0].cells

    def sig_cell2(cell, label, nombre):
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run('\n\n\n_________________________\n').font.size = Pt(10)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(nombre.upper())
        r3.font.size = Pt(9)

    sig_cell2(cells[0], "EL PROMITENTE VENDEDOR", nombre_vend)
    sig_cell2(cells[1], "EL PROMITENTE COMPRADOR", nombre_comp)

    # Remove table borders
    for cell in table.rows[0].cells:
        for side in ['top','left','bottom','right']:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
            tcPr.append(tcBorders)


    # ── CLÁUSULAS ESPECIALES REDACTADAS POR IA ──
    clausulas_esp = datos.get('clausulas_especiales', [])
    if clausulas_esp:
        p(doc, '', space_before=12)
        p(doc, 'CLÁUSULAS ESPECIALES', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        p(doc, '', space_before=4)
        for cl_text in clausulas_esp:
            if cl_text.strip():
                # Each clause block — already formatted by AI
                for line in cl_text.strip().split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    is_header = (line.isupper() and len(line) < 60) or line.endswith('.-')
                    p(doc, line, bold=is_header, space_before=(8 if is_header else 0))
        p(doc, '', space_before=6)

    doc.save(output_path)
    print(f"✓ Promesa de compraventa generada: {output_path}")


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python3 generar_contrato.py <arrendamiento|promesa> <datos.json> <output.docx>")
        sys.exit(1)

    tipo, datos_path, output = sys.argv[1], sys.argv[2], sys.argv[3]

    with open(datos_path) as f:
        datos = json.load(f)

    if tipo == 'arrendamiento':
        generar_arrendamiento(datos, output)
    elif tipo == 'promesa':
        generar_promesa(datos, output)
    else:
        print(f"Tipo desconocido: {tipo}")
        sys.exit(1)

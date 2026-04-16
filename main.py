from fastapi import FastAPI, HTTPException, Query, Request, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import httpx
import os
import time
import re
import asyncio
import base64
import uuid as _uuid
import io
import json
import concurrent.futures
from typing import Optional, List
from datetime import datetime
from pathlib import Path

# Pillow
try:
    from PIL import Image, ImageEnhance
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# OpenCV
try:
    import cv2
    import numpy as np
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False

_thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

CONFIG_FILE = Path(__file__).parent / "config.json"

def load_config() -> dict:
    try:
        if CONFIG_FILE.exists():
            return json.loads(CONFIG_FILE.read_text())
    except Exception:
        pass
    return {}

def save_config(data: dict):
    try:
        CONFIG_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))
    except Exception:
        pass

_config = load_config()

EB_API_KEY       = os.environ.get("EB_API_KEY", "") or _config.get("eb_api_key", "")
GROQ_API_KEY     = os.environ.get("GROQ_API_KEY", "")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
GEMINI_API_KEY    = os.environ.get("GEMINI_API_KEY", "")
EB_BASE          = "https://api.easybroker.com/v1"
GROQ_BASE        = "https://api.groq.com/openai/v1"
ANTHROPIC_BASE   = "https://api.anthropic.com/v1"
GEMINI_BASE      = "https://generativelanguage.googleapis.com/v1beta"
APIFY_API_KEY = os.environ.get("APIFY_API_KEY", "")
GOOGLE_PLACES_KEY = os.environ.get("GOOGLE_PLACES_KEY", "")
SUPABASE_URL      = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY      = os.environ.get("SUPABASE_ANON_KEY", "")

# In-memory PDF store: token → (bytes, filename). Max 50 entradas.
_pdf_store: dict = {}

# ── CACHE EN MEMORIA (TTL 6h) ──
_cache: dict = {}
CACHE_TTL = 21600  # 6 hours default
_cache_ttl: dict = {}  # per-key TTL overrides

def cache_get(key):
    if key in _cache:
        data, ts = _cache[key]
        ttl = _cache_ttl.get(key, CACHE_TTL)
        if time.time() - ts < ttl:
            return data
        del _cache[key]
        _cache_ttl.pop(key, None)
    return None

def cache_set(key, data, ttl=None):
    _cache[key] = (data, time.time())
    if ttl is not None:
        _cache_ttl[key] = ttl

def eb_headers(key: str = None):
    k = key or EB_API_KEY
    return {"X-Authorization": k, "accept": "application/json"}

# ────────────────────────────────────────────
# EASYBROKER — BASE ENDPOINTS
# ────────────────────────────────────────────
@app.get("/")
def root():
    return {"status": "Brokr API activa", "version": "4.0"}

# ────────────────────────────────────────────
# CONFIG — EB API KEY PERSISTENCE
# ────────────────────────────────────────────
class EbKeyRequest(BaseModel):
    key: str

@app.post("/config/eb-key")
async def set_eb_key(req: EbKeyRequest):
    global EB_API_KEY, _config
    EB_API_KEY = req.key.strip()
    _config["eb_api_key"] = EB_API_KEY
    save_config(_config)
    return {"ok": True, "saved": True}

@app.get("/config/eb-key")
async def get_eb_key():
    if EB_API_KEY and len(EB_API_KEY) > 4:
        masked = "*" * (len(EB_API_KEY) - 4) + EB_API_KEY[-4:]
    else:
        masked = ""
    return {"configured": bool(EB_API_KEY), "masked": masked}

# ────────────────────────────────────────────
# GROQ CHAT PROXY
# ────────────────────────────────────────────
class ChatRequest(BaseModel):
    messages: list
    model: str = "llama-3.3-70b-versatile"
    max_tokens: int = 1024
    temperature: float = 0.7

@app.post("/chat")
async def chat_proxy(req: ChatRequest):
    if not GROQ_API_KEY:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY no configurada en el servidor")
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.post(
            f"{GROQ_BASE}/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model":       req.model,
                "messages":    req.messages,
                "max_tokens":  req.max_tokens,
                "temperature": req.temperature,
            }
        )
        if r.status_code != 200:
            raise HTTPException(status_code=r.status_code,
                detail=f"Error Groq: {r.text}")
        return r.json()


# ────────────────────────────────────────────
# CLAUDE CHAT PROXY — SHAARK IA SUPERINTELIGENTE
# ────────────────────────────────────────────
SHAARK_SYSTEM_PROMPT = """Eres Shaark, el asistente de inteligencia artificial de BROKR®, la plataforma inmobiliaria más avanzada de México, especializada en Morelia y Michoacán.

Eres un experto inmobiliario que conoce a fondo:
- LISR (Ley del Impuesto Sobre la Renta) — artículos de enajenación de inmuebles
- ISR por enajenación: exención de 700,000 UDIS, deducciones permitidas, INPC
- Código Civil Federal y de Michoacán — contratos de compraventa y arrendamiento
- SAT: obligaciones fiscales del vendedor y comprador
- Mercado inmobiliario de Morelia: colonias, plusvalía, precios por zona
- Avalúos y valuación de inmuebles (método de mercado, hedónico, físico)

PERSONALIDAD:
- Hablas en español mexicano, natural y cercano
- Eres directo, preciso y profesional — nunca redundante
- Cuando el usuario habla por voz, respondes con oraciones cortas y claras
- Nunca inventes cifras ni datos legales

REGLA DE ORO:
Cuando el usuario pide realizar una tarea, recopila los datos OBLIGATORIOS de UNO EN UNO, de forma conversacional. NUNCA ejecutes la acción con datos incompletos. Cuando tengas todo, di un resumen breve y ejecuta la acción. Los datos opcionales que el usuario no conozca se omiten (usa 0 o "").

══════════════════════════════════════════════════
ACCIÓN 1: CALCULAR ISR POR ENAJENACIÓN
══════════════════════════════════════════════════
Datos OBLIGATORIOS (pregunta uno por uno):
1. Tipo de inmueble: casa habitación, terreno, o comercial
2. Precio de venta (MXN)
3. Mes y año de la venta
4. Precio de compra original (MXN)
5. Mes y año de la compra
6. Si es casa: ¿usó la exención en los últimos 3 años? (sí / no / no sabe)
7. ¿Mejoras o ampliaciones? (monto o "no")
8. ¿Escrituración al comprar? (monto o "no sé")
9. ¿Comisión del agente en esta venta? (monto o "no aplica")

La pregunta 6 SOLO aplica a casa/departamento. Para terrenos y comerciales usa "no" automáticamente.

Cuando tengas todo:
[ACCION]{"tipo":"llenar_isr","precio_venta":NUMERO,"precio_compra":NUMERO,"anio_venta":NUMERO,"mes_venta":NUMERO,"anio_compra":NUMERO,"mes_compra":NUMERO,"inmueble":"casa","exencion":"no","mejoras":NUMERO,"escrituracion":NUMERO,"comision":NUMERO}[/ACCION]

Valores "inmueble": "casa" | "terreno" | "comercial"
Valores "exencion": "no" | "si" | "nose"
mes_venta y mes_compra son números 1-12. Datos opcionales desconocidos = 0.

══════════════════════════════════════════════════
ACCIÓN 2: OPINIÓN DE VALOR CON BÚSQUEDA WEB
══════════════════════════════════════════════════
Cuando el usuario pide valuar, tasar, dar un precio o dar opinión de valor de un inmueble.

Datos OBLIGATORIOS (pregunta uno por uno si faltan):
1. Colonia o fraccionamiento
2. Tipo de inmueble: casa, departamento, terreno, local, oficina, bodega
3. Operación: venta o renta
4. Superficie: m² de construcción (casas/deptos/locales) o m² de terreno (terrenos)

Datos OPCIONALES que si el usuario menciona debes capturar: recámaras, baños, estacionamientos, condición del terreno (plano/pendiente), ciudad (default Morelia).

Cuando tengas los datos OBLIGATORIOS, emite la acción opinion_valor_web:
[ACCION]{"tipo":"opinion_valor_web","colonia":"Vistas Altozano","tipo_inmueble":"terreno","operacion":"venta","m2_terreno":183,"m2_construccion":0,"recamaras":0,"banos":0,"ciudad":"Morelia","condicion_terreno":"plano"}[/ACCION]

Valores "tipo_inmueble": "casa" | "departamento" | "terreno" | "local" | "oficina" | "bodega"
Valores "operacion": "venta" | "renta"
Valores "condicion_terreno": "plano" | "pendiente" | "irregular" | "" (solo para terrenos)
Para casas/deptos: usa m2_construccion. Para terrenos: usa m2_terreno. Ciudad default "Morelia".
Omite campos opcionales que no tengas (usa 0 o "").

══════════════════════════════════════════════════
ACCIÓN 3: GENERAR CONTRATO DE ARRENDAMIENTO
══════════════════════════════════════════════════
Cuando el usuario pide contrato de renta/arrendamiento.
Datos OBLIGATORIOS:
1. Calle del inmueble arrendado
2. Número exterior
3. Colonia del inmueble
4. C.P. (código postal)
5. Municipio y estado (ej: "Morelia, Michoacán")
6. Nombre completo del arrendador (dueño) — EN MAYÚSCULAS
7. Nombre completo del arrendatario (inquilino) — EN MAYÚSCULAS
8. Renta mensual (MXN)
9. Depósito en garantía (si no sabe, usa el mismo valor que la renta)
10. Fecha de inicio (día/mes/año)

Cuando tengas todo:
[ACCION]{"tipo":"llenar_contrato","subtipo":"arrendamiento","calle_inmueble":"AV. CAMELINAS","num_ext":"123","num_int":"","colonia":"CHAPULTEPEC","cp":"58260","municipio_estado":"MORELIA, MICHOACÁN","arrendador":"SALVADOR BOLAÑOS NAVARRO","arrendatario":"GABRIELA NAVARRO PÉREZ","renta":8500,"deposito":8500,"dia_pago":5,"fecha_inicio":"2026-05-01"}[/ACCION]

dia_pago: día límite del mes para pagar (default 5). fecha_inicio en formato YYYY-MM-DD.

══════════════════════════════════════════════════
ACCIÓN 4: GENERAR PROMESA DE COMPRAVENTA
══════════════════════════════════════════════════
Cuando el usuario pide contrato de compraventa o promesa de venta.
Datos OBLIGATORIOS:
1. Dirección del inmueble (calle y número)
2. Colonia
3. C.P.
4. Nombre del vendedor (promitente vendedor)
5. Nombre del comprador (promitente comprador)
6. Precio total de venta
7. Monto de arras/enganche
8. Fecha límite para escriturar

Cuando tengas todo:
[ACCION]{"tipo":"llenar_contrato","subtipo":"promesa","dir":"Cipres 167","colonia":"Melchor Ocampo","cp":"58160","vendedor":"JUAN PÉREZ GARCÍA","comprador":"MARÍA LÓPEZ HERNÁNDEZ","precio":2500000,"arras":250000,"fecha_limite":"2026-06-30"}[/ACCION]

fecha_limite en formato YYYY-MM-DD.

══════════════════════════════════════════════════
ACCIÓN 5: FICHA TÉCNICA DESDE EASYBROKER
══════════════════════════════════════════════════
Cuando el usuario quiere hacer una ficha de una propiedad de EasyBroker y da el ID (formato EB-XXXX).
[ACCION]{"tipo":"crear_ficha","id_easybroker":"EB-KH4322"}[/ACCION]

Si el usuario no da el ID, navega al módulo y pídele el ID:
[ACCION]{"tipo":"navegar","modulo":"ficha"}[/ACCION]

══════════════════════════════════════════════════
ACCIÓN 6: FICHA TÉCNICA MANUAL
══════════════════════════════════════════════════
Cuando el usuario quiere hacer una ficha técnica sin ID de EasyBroker y da los datos del inmueble.
Datos mínimos: tipo, operación, precio, colonia.
[ACCION]{"tipo":"crear_ficha_manual","tipo_inmueble":"casa","operacion":"venta","precio":3500000,"colonia":"Chapultepec","ciudad":"Morelia","calle":"Av. Madero 123","recamaras":3,"banos":2,"m2_construccion":180,"m2_terreno":220,"estacionamientos":2,"descripcion":""}[/ACCION]

Valores "operacion": "venta" | "renta". Omite campos que no tengas.

══════════════════════════════════════════════════
ACCIÓN 7: BUSCAR PROPIEDAD EN MIS INMUEBLES
══════════════════════════════════════════════════
Cuando el usuario pide ver, buscar o encontrar una propiedad en su cartera.
[ACCION]{"tipo":"buscar_propiedad","query":"Chapultepec"}[/ACCION]

══════════════════════════════════════════════════
NAVEGACIÓN DIRECTA
══════════════════════════════════════════════════
Para ir a un módulo sin datos adicionales:
[ACCION]{"tipo":"navegar","modulo":"isr"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"ficha-manual"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"ficha"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"contratos"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"avm"}[/ACCION]
[ACCION]{"tipo":"navegar","modulo":"props"}[/ACCION]

══════════════════════════════════════════════════
EJEMPLOS DE CONVERSACIÓN CORRECTA
══════════════════════════════════════════════════

EJEMPLO ISR:
Usuario: "calcula el ISR de una casa que vendí en 3.2 millones en marzo 2026, la compré en un millón en enero 2015, comisión 96 mil"
Shaark: "Perfecto. Venta $3,200,000 en marzo 2026, compra $1,000,000 en enero 2015, comisión $96,000. ¿Usaste la exención de casa habitación en los últimos 3 años?"
Usuario: "no"
Shaark: "Listo. Calculando ISR ahora."
[ACCION]{"tipo":"llenar_isr","precio_venta":3200000,"precio_compra":1000000,"anio_venta":2026,"mes_venta":3,"anio_compra":2015,"mes_compra":1,"inmueble":"casa","exencion":"no","mejoras":0,"escrituracion":0,"comision":96000}[/ACCION]

EJEMPLO CONTRATO DE RENTA:
Usuario: "dame un contrato de renta para Av. Camelinas 123, Chapultepec, Morelia, el dueño es Salvador Bolaños y el inquilino es Gabriela Navarro, renta de 8,500 al mes"
Shaark: "¿Cuál es el código postal de Chapultepec?"
Usuario: "58260"
Shaark: "¿Desde qué fecha empieza el contrato?"
Usuario: "1 de mayo de 2026"
Shaark: "Perfecto. Generando contrato de arrendamiento."
[ACCION]{"tipo":"llenar_contrato","subtipo":"arrendamiento","calle_inmueble":"AV. CAMELINAS","num_ext":"123","num_int":"","colonia":"CHAPULTEPEC","cp":"58260","municipio_estado":"MORELIA, MICHOACÁN","arrendador":"SALVADOR BOLAÑOS NAVARRO","arrendatario":"GABRIELA NAVARRO PÉREZ","renta":8500,"deposito":8500,"dia_pago":5,"fecha_inicio":"2026-05-01"}[/ACCION]

EJEMPLO AVM:
Usuario: "valúa un terreno plano de 183m² en Vistas Altozano"
Shaark: "¿Es para venta o renta?"
Usuario: "venta"
Shaark: "Perfecto. Buscando comparables en el mercado ahora mismo."
[ACCION]{"tipo":"opinion_valor_web","colonia":"Vistas Altozano","tipo_inmueble":"terreno","operacion":"venta","m2_terreno":183,"m2_construccion":0,"recamaras":0,"banos":0,"ciudad":"Morelia","condicion_terreno":"plano"}[/ACCION]

EJEMPLO AVM CASA:
Usuario: "dame el valor de una casa de 180m² construcción, 3 recámaras, 2 baños en Chapultepec"
Shaark: "¿Operación venta o renta?"
Usuario: "venta"
Shaark: "Listo. Analizando el mercado de Chapultepec."
[ACCION]{"tipo":"opinion_valor_web","colonia":"Chapultepec","tipo_inmueble":"casa","operacion":"venta","m2_construccion":180,"m2_terreno":0,"recamaras":3,"banos":2,"ciudad":"Morelia","condicion_terreno":""}[/ACCION]

EJEMPLO FICHA EB:
Usuario: "haz la ficha de la propiedad EB-KH4322"
Shaark: "Generando ficha técnica de EB-KH4322."
[ACCION]{"tipo":"crear_ficha","id_easybroker":"EB-KH4322"}[/ACCION]

Responde siempre en español. Nunca uses markdown en respuestas conversacionales (sin **, sin #, sin listas con guiones)."""

class ClaudeChatRequest(BaseModel):
    messages: list
    max_tokens: int = 1200
    temperature: float = 0.7
    context: str = ""  # Módulo/pantalla activa — se inyecta al system prompt

@app.post("/chat-claude")
async def chat_claude_proxy(req: ClaudeChatRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY no configurada en el servidor")

    # Construir system prompt con contexto dinámico del módulo activo
    system_content = SHAARK_SYSTEM_PROMPT
    if req.context:
        system_content += f"\n\n═══════════════════════════════════════\nCONTEXTO ACTUAL DEL USUARIO\n═══════════════════════════════════════\nEl usuario está en: {req.context}\nAdapta tu respuesta y acciones a este módulo cuando sea relevante."

    user_messages = [m for m in req.messages if m.get("role") != "system"]

    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(
            f"{ANTHROPIC_BASE}/messages",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": req.max_tokens,
                "system": system_content,
                "messages": user_messages,
            }
        )
        if r.status_code != 200:
            raise HTTPException(status_code=r.status_code,
                detail=f"Error Claude: {r.text}")

        data = r.json()
        reply_text = data.get("content", [{}])[0].get("text", "Sin respuesta.")
        return {
            "choices": [
                {"message": {"role": "assistant", "content": reply_text}}
            ]
        }


@app.post("/isr-pdf")
async def generar_isr_pdf(p: dict):
    """Recibe HTML del cálculo ISR y lo convierte a PDF con Playwright."""
    from playwright.async_api import async_playwright  # noqa: re-import ok here (lazy)
    html = p.get("html", "")
    if not html:
        raise HTTPException(status_code=400, detail="HTML vacío")
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        page = await browser.new_page()
        await page.set_content(html, wait_until="domcontentloaded")
        await page.wait_for_timeout(300)
        pdf_bytes = await page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "20mm", "right": "20mm", "bottom": "20mm", "left": "20mm"}
        )
        await browser.close()
    token = str(_uuid.uuid4()).replace("-","")[:16]
    filename = p.get("filename", "ISR_Brokr.pdf")
    _pdf_store[token] = (pdf_bytes, filename)
    if len(_pdf_store) > 50:
        oldest = list(_pdf_store.keys())[0]
        del _pdf_store[oldest]
    from fastapi.responses import JSONResponse
    return JSONResponse({"token": token, "filename": filename})


@app.get("/propiedad/{property_id}")
async def get_propiedad(property_id: str, request: Request):
    user_key = request.headers.get("X-EB-Key", "").strip()
    if not user_key:
        raise HTTPException(status_code=400, detail="Configura tu API key de EasyBroker en Perfil → API EasyBroker para usar este módulo.")
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(f"{EB_BASE}/properties/{property_id}",
                             headers=eb_headers(user_key))
        if r.status_code == 401:
            raise HTTPException(status_code=401, detail="API key de EasyBroker inválida. Verifica tu configuración en Perfil → API EasyBroker.")
        if r.status_code == 404:
            raise HTTPException(status_code=404, detail="Propiedad no encontrada")
        if r.status_code != 200:
            raise HTTPException(status_code=r.status_code, detail="Error EasyBroker")
        return r.json()

@app.get("/propiedades")
async def get_propiedades(page: int = 1, limit: int = 20):
    if not EB_API_KEY:
        raise HTTPException(status_code=500, detail="EB_API_KEY no configurada")
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(f"{EB_BASE}/properties", headers=eb_headers(),
                             params={"page": page, "limit": limit})
        if r.status_code != 200:
            raise HTTPException(status_code=r.status_code, detail="Error EasyBroker")
        return r.json()

# ────────────────────────────────────────────
# COLONIAS AUTOCOMPLETE
# ────────────────────────────────────────────
async def fetch_all_properties() -> list:
    """Fetch all properties from EB and cache them."""
    cached = cache_get("all_properties")
    if cached is not None:
        return cached

    all_props = []
    page = 1
    async with httpx.AsyncClient(timeout=30) as client:
        while True:
            r = await client.get(f"{EB_BASE}/properties", headers=eb_headers(),
                                 params={"limit": 50, "page": page})
            if r.status_code != 200:
                break
            data = r.json()
            props = data.get("content", [])
            if not props:
                break
            all_props.extend(props)
            # Stop if we have enough or no more pages
            total = data.get("pagination", {}).get("total", 0)
            if len(all_props) >= min(total, 3000):  # cap at 3000 for speed
                break
            if not data.get("pagination", {}).get("next_page"):
                break
            page += 1
            if page > 60:  # safety cap
                break

    cache_set("all_properties", all_props)
    return all_props

def extract_colonia(location_str: str) -> str:
    """Extract colonia from 'Colonia, Ciudad, Estado' string."""
    if not location_str:
        return ""
    parts = [p.strip() for p in location_str.split(",")]
    return parts[0] if parts else location_str.strip()

def normalize(s: str) -> str:
    for a, b in [('á','a'),('é','e'),('í','i'),('ó','o'),('ú','u'),('ü','u'),('ñ','n')]:
        s = s.lower().replace(a, b)
    return s

@app.get("/colonias")
async def get_colonias(q: str = Query("", min_length=2), ciudad: str = "Morelia"):
    """Return unique colonias matching search query — fast direct EB search."""
    if not EB_API_KEY:
        raise HTTPException(status_code=500, detail="EB_API_KEY no configurada")

    cache_key = f"colonias_{normalize(ciudad)}"
    colonias_map = cache_get(cache_key)

    if colonias_map is None:
        # Build index: paginate EB and collect all colonias
        colonias_map = {}
        page = 1
        async with httpx.AsyncClient(timeout=30) as client:
            while page <= 80:  # up to 4000 properties
                r = await client.get(
                    f"{EB_BASE}/properties",
                    headers=eb_headers(),
                    params={"limit": 50, "page": page}
                )
                if r.status_code != 200:
                    break
                data = r.json()
                props = data.get("content", [])
                if not props:
                    break
                for p in props:
                    loc = p.get("location", "")
                    if not loc or normalize(ciudad) not in normalize(loc):
                        continue
                    # Status field empty in this EB plan — no filter
                    # Date: January 2025 onwards
                    # No date filter — all properties included
                    col = extract_colonia(loc)
                    if col and len(col) > 2:
                        colonias_map[col] = colonias_map.get(col, 0) + 1
                if not data.get("pagination",{}).get("next_page"):
                    break
                page += 1
        cache_set(cache_key, colonias_map)

    q_norm = normalize(q)
    matches = [
        {"colonia": col, "count": cnt}
        for col, cnt in colonias_map.items()
        if q_norm in normalize(col)
    ]
    matches.sort(key=lambda x: -x["count"])
    return {"colonias": matches[:12], "total_colonias": len(colonias_map)}

# ────────────────────────────────────────────
# AVM — HELPERS
# ────────────────────────────────────────────
class AVMRequest(BaseModel):
    colonia: str
    ciudad: str
    tipo: str
    operacion: str
    m2_construccion: Optional[float] = None
    m2_terreno:      Optional[float] = None
    recamaras:       Optional[int]   = None
    banos:           Optional[float] = None
    estado:          Optional[str]   = "bueno"
    anio_construccion: Optional[int] = None

def parse_price(val) -> Optional[float]:
    if not val:
        return None
    try:
        v = float(str(val).replace(",", ""))
        if 50_000 <= v <= 999_000_000:
            return v
    except:
        pass
    return None

TIPO_MAP = {
    "casa":          ["Casa"],
    "departamento":  ["Departamento"],
    "terreno":       ["Terreno"],
    "local":         ["Local comercial"],
    "comercial":     ["Local comercial","Oficina","Bodega"],
    "oficina":       ["Oficina"],
    "bodega":        ["Bodega"],
}
OP_MAP = {
    "venta": "sale",
    "renta": "rental",
}


TIPO_SIMILAR = {
    "casa":         ["Casa","Departamento"],
    "departamento": ["Departamento","Casa"],
    "terreno":      ["Terreno"],
    "local":        ["Local comercial","Oficina","Bodega"],
    "comercial":    ["Local comercial","Oficina","Bodega"],
    "oficina":      ["Oficina","Local comercial"],
    "bodega":       ["Bodega","Local comercial"],
}

async def get_comparables_eb(colonia: str, ciudad: str,
                              tipo: str, operacion: str) -> list:
    cache_key = f"comp_{colonia}_{ciudad}_{tipo}_{operacion}".lower().replace(" ","_")
    cached = cache_get(cache_key)
    if cached is not None:
        return cached

    # Map tipo to EB property_type values
    tipo_labels = TIPO_MAP.get(tipo.lower(), [tipo.capitalize()])
    op_type     = OP_MAP.get(operacion.lower(), "sale")

    comparables = []
    page = 1

    def norm(s):
        for a,b in [("á","a"),("é","e"),("í","i"),("ó","o"),("ú","u"),("ñ","n")]:
            s = s.lower().replace(a,b)
        return re.sub(r"[^a-z0-9 ]", "", s).strip()

    async with httpx.AsyncClient(timeout=60) as client:
        while len(comparables) < 50 and page <= 160:
            r = await client.get(
                f"{EB_BASE}/properties",
                headers=eb_headers(),
                params={"limit": 50, "page": page}
            )
            if r.status_code != 200:
                break
            data = r.json()
            props = data.get("content", [])
            if not props:
                break

            for p in props:
                # ── 1. COLONIA FILTER (strict) ──
                loc = p.get("location", "")
                if not loc:
                    continue
                if colonia and norm(colonia) not in norm(loc):
                    continue
                if norm(ciudad) not in norm(loc):
                    continue

                # ── 2. TIPO FILTER ──
                prop_type = p.get("property_type", "")
                tipo_match = any(norm(t) in norm(prop_type) for t in tipo_labels)
                if not tipo_match:
                    continue

                # ── 3. OPERATION FILTER (strict) ──
                ops = p.get("operations", [])
                matching_op = None
                for op in ops:
                    if op.get("type") == op_type:
                        matching_op = op
                        break
                if not matching_op:
                    continue  # wrong operation type — skip

                # ── 4. DATE: use created_at for appreciation calculation ──
                created_at   = p.get("created_at", "")
                published_at = p.get("published_at", "") or p.get("updated_at", "")
                # created_at = when property was first entered in EB (true age)
                pub_year = 2026  # default
                if created_at:
                    try:
                        pub_year = int(created_at[:4])
                    except:
                        pass

                # ── 5. PRICE ──
                price = parse_price(matching_op.get("amount"))
                if not price:
                    continue

                col_prop = extract_colonia(loc)
                comparables.append({
                    "precio":          price,
                    "titulo":          p.get("title", "")[:80],
                    "m2_construccion": p.get("construction_size"),
                    "m2_terreno":      p.get("lot_size"),
                    "recamaras":       p.get("bedrooms"),
                    "banos":           p.get("bathrooms"),
                    "colonia":         col_prop,
                    "fuente":          "EasyBroker",
                    "public_id":       p.get("public_id", ""),
                    "publicado":       created_at[:10] if created_at else (published_at[:10] if published_at else ""),
                    "pub_year":        pub_year,
                    "tipo_exacto":     norm(tipo_labels[0]) in norm(prop_type),
                })

            if not data.get("pagination", {}).get("next_page"):
                break
            page += 1

    # Remove outliers
    if len(comparables) >= 3:
        prices = sorted(c["precio"] for c in comparables)
        median = prices[len(prices)//2]
        comparables = [c for c in comparables
                       if median * 0.25 <= c["precio"] <= median * 4.0]

    cache_set(cache_key, comparables[:30])
    return comparables[:30]


# ────────────────────────────────────────────
# HEDONIC MODEL
# ────────────────────────────────────────────
APRECIACION_ANUAL = 0.04  # 4% annual real estate appreciation in Morelia
ANIO_ACTUAL = 2026

def ajuste_hedonico(comp: dict, sujeto: dict) -> dict:
    precio_base = comp["precio"]
    ajustes = []
    factor  = 1.0

    # ── 0. PRICE UPDATE BY APPRECIATION (4% annual) ──
    pub_year = comp.get("pub_year", ANIO_ACTUAL)
    anos_transcurridos = max(0, ANIO_ACTUAL - pub_year)
    if anos_transcurridos > 0:
        factor_apreciacion = (1 + APRECIACION_ANUAL) ** anos_transcurridos
        factor *= factor_apreciacion
        ajustes.append(f"actualización {anos_transcurridos} año{'s' if anos_transcurridos>1 else ''} "
                       f"(+{round((factor_apreciacion-1)*100,1)}% a 4%/año)")

    # m² construction (sqrt scaling)
    m2s = sujeto.get("m2_construccion")
    m2c = comp.get("m2_construccion")
    if m2s and m2c and m2c > 0 and abs(m2s - m2c) > 5:
        ratio = (m2s / m2c) ** 0.5
        factor *= ratio
        diff = m2s - m2c
        ajustes.append(f"m² ({'+' if diff>0 else ''}{diff:.0f}): "
                       f"{'+' if ratio>1 else ''}{(ratio-1)*100:.1f}%")

    # Bedrooms (4% per room)
    rs = sujeto.get("recamaras")
    rc = comp.get("recamaras")
    if rs and rc and rs != rc:
        diff = rs - rc
        factor *= (1 + diff * 0.04)
        ajustes.append(f"recámaras ({'+' if diff>0 else ''}{diff}): "
                       f"{'+' if diff>0 else ''}{diff*4}%")

    # Conservation state
    estado_adj = {"malo":-0.15,"regular":-0.07,"bueno":0.0,"excelente":0.08}
    adj_e = estado_adj.get(sujeto.get("estado","bueno"), 0.0)
    if adj_e != 0:
        factor *= (1 + adj_e)
        ajustes.append(f"estado ({sujeto.get('estado')}): "
                       f"{'+' if adj_e>0 else ''}{adj_e*100:.0f}%")

    # Age (1.5% per decade over 10 years)
    anio = sujeto.get("anio_construccion")
    if anio:
        anos = datetime.now().year - anio
        age_adj = max(-0.20, min(0.15, -0.015 * ((anos - 10) / 10)))
        if abs(age_adj) > 0.01:
            factor *= (1 + age_adj)
            ajustes.append(f"antigüedad ({anos} años): "
                           f"{'+' if age_adj>0 else ''}{age_adj*100:.1f}%")

    # EB properties are already real transaction prices
    # No offer-to-close discount needed (unlike portal listings)
    if not ajustes:
        ajustes.append("sin ajustes — comparable directo")

    return {
        **comp,
        "precio_ajustado": round(precio_base * factor, -3),
        "factor_total":    round(factor, 4),
        "ajustes":         ajustes,
    }

# ────────────────────────────────────────────
# AVM ENDPOINT
# ────────────────────────────────────────────


@app.post("/avm")
async def calcular_avm(req: AVMRequest):
    if not EB_API_KEY:
        raise HTTPException(status_code=500, detail="EB_API_KEY no configurada")

    comparables_raw = await get_comparables_eb(
        req.colonia, req.ciudad, req.tipo, req.operacion
    )

    nivel = 1
    nivel_msg = ""

    # If < 3 exact tipo matches, try similar tipos in same colonia
    exact_matches = [c for c in comparables_raw if c.get("tipo_exacto", True)]
    if len(exact_matches) < 3 and req.tipo.lower() in TIPO_SIMILAR:
        similar_tipos = TIPO_SIMILAR[req.tipo.lower()]
        for tipo_alt in similar_tipos[1:]:  # skip first (same as original)
            alt_comps = await get_comparables_eb(
                req.colonia, req.ciudad, tipo_alt.lower(), req.operacion
            )
            for c in alt_comps:
                if c not in comparables_raw:
                    comparables_raw.append(c)
        if len(comparables_raw) >= 3:
            nivel_msg = (f"{len(exact_matches)} comparables exactos en {req.colonia}. "
                         f"Se complementó con tipos similares en la misma colonia.")

    if len(comparables_raw) < 3:
        nivel = 2
        comparables_raw = await get_comparables_eb(
            "", req.ciudad, req.tipo, req.operacion
        )
        nivel_msg = (f"Pocos comparables en {req.colonia} con datos ene 2025–mar 2026. "
                     f"Se amplió a {req.ciudad} — filtrado por precio/m².")

    if len(comparables_raw) < 2:
        raise HTTPException(
            status_code=422,
            detail=(f"No se encontraron comparables de {req.tipo} en {req.operacion} "
                    f"en {req.ciudad}. Verifica el tipo de operación e inmueble.")
        )

    sujeto = {
        "m2_construccion":   req.m2_construccion,
        "m2_terreno":        req.m2_terreno,
        "recamaras":         req.recamaras,
        "banos":             req.banos,
        "estado":            req.estado,
        "anio_construccion": req.anio_construccion,
    }

    # Apply hedonic adjustments
    ajustados = []
    for comp in comparables_raw:
        try:
            ajustados.append(ajuste_hedonico(comp, sujeto))
        except:
            continue

    if not ajustados:
        raise HTTPException(status_code=422, detail="Error procesando comparables")

    # Filter by price/m² if we have m2 data (nivel 2 only)
    if nivel == 2 and req.m2_construccion and req.m2_construccion > 0:
        pm2s = [(c, c["precio_ajustado"] / req.m2_construccion)
                for c in ajustados]
        if len(pm2s) >= 5:
            vals = sorted(p for _, p in pm2s)
            median_pm2 = vals[len(vals)//2]
            ajustados = [c for c, pm2 in pm2s
                         if median_pm2 * 0.65 <= pm2 <= median_pm2 * 1.35]

    # Calculate value range
    precios = sorted(c["precio_ajustado"] for c in ajustados)
    n       = len(precios)
    trim    = max(1, n // 10)
    p_trim  = precios[trim: n-trim] if n > 4 else precios

    valor_minimo   = round(min(p_trim), -3)
    valor_probable = round(sum(p_trim) / len(p_trim), -3)
    valor_maximo   = round(max(p_trim), -3)

    # Price per m²
    pm2_list = []
    for c in ajustados:
        m2 = c.get("m2_construccion") or req.m2_construccion
        if m2 and m2 > 0:
            pm2_list.append(c["precio_ajustado"] / m2)
    pm2_prom = round(sum(pm2_list) / len(pm2_list)) if pm2_list else None

    nivel_labels = {
        1: f"Alta confianza — {len(ajustados)} comparables en {req.colonia}",
        2: f"Confianza media — {len(ajustados)} comparables en {req.ciudad} (filtrado por precio/m²)",
    }

    return {
        "colonia":            req.colonia,
        "ciudad":             req.ciudad,
        "tipo":               req.tipo,
        "operacion":          req.operacion,
        "nivel":              nivel,
        "nivel_mensaje":      nivel_labels.get(nivel, nivel_msg),
        "fuentes":            ["EasyBroker"],
        "num_comparables":    len(ajustados),
        "valor_minimo":       valor_minimo,
        "valor_probable":     valor_probable,
        "valor_maximo":       valor_maximo,
        "precio_m2_promedio": pm2_prom,
        "comparables":        ajustados[:10],
        "nota": ("Valores calculados con base en propiedades publicadas en la bolsa "
                 "EasyBroker — comparables actualizados al 2026 con apreciación del 4% anual, más ajustes hedónicos por m², recámaras, "
                 "estado y antigüedad. El valor definitivo requiere inspección física "
                 "y avalúo formal."),
        "timestamp": time.strftime("%Y-%m-%d %H:%M"),
    }


# ────────────────────────────────────────────
# AVM — CLAUDE AI OPINION DE VALOR
# ────────────────────────────────────────────

class AvmClaudeRequest(BaseModel):
    # Ubicación
    estado: str
    ciudad: str
    colonia: str = ""
    direccion: str = ""
    tipo_zona: str = ""      # residencial, comercial, industrial, mixta, turistica
    nse: str = ""            # A, B, C+, C, D+, D, E
    # Inmueble
    tipo: str                # casa, departamento, terreno, local, oficina, bodega, edificio
    operacion: str = "venta" # venta | renta
    m2_construccion: float = 0
    m2_terreno: float = 0
    recamaras: int = 0
    banos_completos: float = 0
    medios_banos: int = 0
    estacionamientos: int = 0
    nivel_piso: int = 0
    # Estado y acabados
    antiguedad: int = 0
    conservacion: str = "bueno"  # excelente, bueno, regular, malo
    acabados: str = "medio"      # lujo, residencial_plus, residencial, medio, economico
    remodelado: bool = False
    descripcion_remodelacion: str = ""
    # Amenidades
    amenidades: list = []        # alberca, jardin, bodega, cuarto_servicio, elevador, seguridad, gimnasio, salon
    # Contexto
    precio_lista: float = 0
    motivo_valuacion: str = ""
    comentarios: str = ""

@app.post("/api/avm-claude")
async def avm_claude(req: AvmClaudeRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY no configurada en el servidor")

    # Construir descripción detallada de la propiedad
    tipo_labels = {
        "casa": "Casa habitación", "departamento": "Departamento/Condominio",
        "terreno": "Terreno", "local": "Local comercial",
        "oficina": "Oficina", "bodega": "Bodega/Nave industrial", "edificio": "Edificio"
    }
    conservacion_labels = {
        "excelente": "Excelente / Como nuevo", "bueno": "Bueno",
        "regular": "Regular / Necesita detalles", "malo": "Malo / Requiere remodelación"
    }
    acabados_labels = {
        "lujo": "Lujo / Residencial Plus", "residencial_plus": "Residencial Plus",
        "residencial": "Residencial", "medio": "Estándar / Medio", "economico": "Económico / Interés social"
    }

    partes = []
    partes.append(f"TIPO DE INMUEBLE: {tipo_labels.get(req.tipo, req.tipo)}")
    partes.append(f"OPERACIÓN: {req.operacion.upper()}")
    partes.append(f"\nUBICACIÓN:")
    partes.append(f"  - Estado: {req.estado}")
    partes.append(f"  - Ciudad/Municipio: {req.ciudad}")
    if req.colonia: partes.append(f"  - Colonia/Fraccionamiento: {req.colonia}")
    if req.direccion: partes.append(f"  - Dirección: {req.direccion}")
    if req.tipo_zona: partes.append(f"  - Tipo de zona: {req.tipo_zona}")
    if req.nse: partes.append(f"  - Nivel socioeconómico de la zona: {req.nse}")

    partes.append(f"\nDIMENSIONES:")
    if req.m2_construccion > 0: partes.append(f"  - Superficie construida: {req.m2_construccion} m²")
    if req.m2_terreno > 0: partes.append(f"  - Superficie de terreno: {req.m2_terreno} m²")
    if req.recamaras > 0: partes.append(f"  - Recámaras: {req.recamaras}")
    if req.banos_completos > 0: partes.append(f"  - Baños completos: {req.banos_completos}")
    if req.medios_banos > 0: partes.append(f"  - Medios baños: {req.medios_banos}")
    if req.estacionamientos > 0: partes.append(f"  - Estacionamientos: {req.estacionamientos}")
    if req.nivel_piso > 0: partes.append(f"  - Piso/Nivel: {req.nivel_piso}")

    partes.append(f"\nESTADO DEL INMUEBLE:")
    partes.append(f"  - Antigüedad aproximada: {req.antiguedad} años")
    partes.append(f"  - Estado de conservación: {conservacion_labels.get(req.conservacion, req.conservacion)}")
    partes.append(f"  - Calidad de acabados: {acabados_labels.get(req.acabados, req.acabados)}")
    if req.remodelado:
        partes.append(f"  - Remodelado recientemente: SÍ")
        if req.descripcion_remodelacion:
            partes.append(f"  - Descripción remodelación: {req.descripcion_remodelacion}")

    if req.amenidades:
        amenidad_labels = {
            "alberca": "Alberca/Pool", "jardin": "Jardín", "bodega": "Bodega",
            "cuarto_servicio": "Cuarto de servicio", "elevador": "Elevador",
            "seguridad": "Seguridad/Vigilancia 24h", "gimnasio": "Gimnasio",
            "salon": "Salón de eventos", "roof_garden": "Roof garden",
            "terraza": "Terraza", "vista": "Vista panorámica", "acceso_playa": "Acceso a playa",
        }
        am_list = [amenidad_labels.get(a, a) for a in req.amenidades]
        partes.append(f"\nAMENIDADES: {', '.join(am_list)}")

    if req.precio_lista > 0:
        partes.append(f"\nPRECIO DE LISTA ACTUAL: ${req.precio_lista:,.0f} MXN")
    if req.motivo_valuacion:
        partes.append(f"MOTIVO DE LA VALUACIÓN: {req.motivo_valuacion}")
    if req.comentarios:
        partes.append(f"COMENTARIOS ADICIONALES: {req.comentarios}")

    descripcion = "\n".join(partes)

    system_prompt = """Eres el mejor perito valuador de bienes raíces de México, certificado por la Sociedad Hipotecaria Federal y el INDAABIN, con 30 años de experiencia valuando propiedades en todo el territorio nacional. Tu análisis es utilizado por bancos, notarías y juzgados para transacciones de millones de pesos. La vida financiera del usuario que solicita esta opinión de valor depende de la precisión de tu análisis.

Tu misión: proporcionar la opinión de valor más precisa, fundamentada y útil posible basándote en:
1. Tu conocimiento profundo del mercado inmobiliario mexicano por región, ciudad y colonia
2. Tendencias y precios actuales del mercado (hasta tu fecha de corte de conocimiento)
3. Factores macroeconómicos: inflación, tasas de interés, INPP, INPC
4. El Método Comparativo de Mercado (enfoque principal)
5. El Enfoque Físico o de Costos (edificaciones)
6. El Enfoque de Capitalización de Rentas (cuando aplique)
7. Ajustes hedónicos por ubicación, características, estado y acabados

IMPORTANTE: Responde ÚNICAMENTE con un objeto JSON válido (sin texto antes ni después, sin markdown, sin ```json), con exactamente esta estructura:
{
  "valor_estimado": <número en pesos MXN sin comas ni signos>,
  "valor_minimo": <número>,
  "valor_maximo": <número>,
  "valor_por_m2_construccion": <número o 0 si no aplica>,
  "valor_por_m2_terreno": <número o 0 si no aplica>,
  "nivel_confianza": "<alta|media|baja>",
  "razon_confianza": "<por qué ese nivel>",
  "resumen_ejecutivo": "<2-3 oraciones concretas sobre el valor>",
  "analisis_ubicacion": "<análisis del valor de la zona y su impacto>",
  "analisis_propiedad": "<análisis de las características físicas y su impacto>",
  "factores_positivos": ["<factor 1>", "<factor 2>", ...],
  "factores_negativos": ["<factor 1>", "<factor 2>", ...],
  "recomendaciones": ["<recomendación 1>", "<recomendación 2>", ...],
  "mercado_actual": "<descripción del mercado actual en esa zona>",
  "metodologia": "<metodología aplicada y justificación>",
  "advertencias": "<advertencias o limitaciones de esta opinión>"
}"""

    user_msg = f"""Por favor valúa la siguiente propiedad y proporciona tu opinión de valor profesional:

{descripcion}

Recuerda: responde ÚNICAMENTE con el JSON, sin ningún texto adicional."""

    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(
            f"{ANTHROPIC_BASE}/messages",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": 2000,
                "temperature": 0.3,
                "messages": [{"role": "user", "content": user_msg}],
                "system": system_prompt,
            },
        )

    if r.status_code != 200:
        raise HTTPException(status_code=502, detail=f"Error de Claude: {r.text[:300]}")

    raw = r.json().get("content", [{}])[0].get("text", "")
    # Limpiar posibles markdown wrappers
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1]
        if raw.endswith("```"):
            raw = raw[:-3]
    raw = raw.strip()

    try:
        resultado = _json.loads(raw)
    except Exception:
        raise HTTPException(status_code=502, detail=f"Claude no devolvió JSON válido: {raw[:500]}")

    resultado["timestamp"] = time.strftime("%Y-%m-%d %H:%M")
    resultado["propiedad_descripcion"] = f"{tipo_labels.get(req.tipo, req.tipo)} en {req.colonia or req.ciudad}, {req.estado}"
    return resultado


# ────────────────────────────────────────────
# AVM — OPINIÓN DE VALOR CON WEB SEARCH
# ────────────────────────────────────────────

class AvmWebSearchRequest(BaseModel):
    colonia: str
    tipo_inmueble: str          # casa | departamento | terreno | local | oficina | bodega
    operacion: str = "venta"    # venta | renta
    m2_construccion: float = 0
    m2_terreno: float = 0
    recamaras: int = 0
    banos: float = 0
    estacionamientos: int = 0
    condicion_terreno: str = "" # plano | pendiente | irregular
    ciudad: str = "Morelia"
    estado: str = "Michoacán"
    comentarios: str = ""

@app.post("/api/avm-websearch")
async def avm_websearch(req: AvmWebSearchRequest):
    """Genera opinión de valor buscando comparables reales en internet con web_search tool de Claude."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY no configurada")

    tipo_labels = {
        "casa": "Casa habitación", "departamento": "Departamento/Condominio",
        "terreno": "Terreno", "local": "Local comercial",
        "oficina": "Oficina", "bodega": "Bodega/Nave industrial",
    }
    tipo_label = tipo_labels.get(req.tipo_inmueble, req.tipo_inmueble)
    es_terreno = req.tipo_inmueble == "terreno"

    # Construir descripción del sujeto
    partes = [f"INMUEBLE A VALUAR: {tipo_label} en {req.operacion.upper()}"]
    partes.append(f"Ubicación: {req.colonia}, {req.ciudad}, {req.estado}")
    if req.m2_terreno > 0:
        cond = f" ({req.condicion_terreno})" if req.condicion_terreno else ""
        partes.append(f"Superficie de terreno: {req.m2_terreno} m²{cond}")
    if req.m2_construccion > 0:
        partes.append(f"Superficie construida: {req.m2_construccion} m²")
    if req.recamaras > 0: partes.append(f"Recámaras: {req.recamaras}")
    if req.banos > 0: partes.append(f"Baños: {req.banos}")
    if req.estacionamientos > 0: partes.append(f"Estacionamientos: {req.estacionamientos}")
    if req.comentarios: partes.append(f"Comentarios: {req.comentarios}")
    descripcion_sujeto = "\n".join(partes)

    system_prompt = """Eres el mejor perito valuador de bienes raíces de México, con 30 años de experiencia y certificación de la Sociedad Hipotecaria Federal. Tu análisis es utilizado por bancos, notarías y juzgados. La precisión de tu opinión tiene consecuencias financieras reales para el usuario.

PROCESO OBLIGATORIO — SIGUE ESTOS PASOS EN ORDEN, SIN SALTARTE NINGUNO:

PASO 1 — BÚSQUEDA MÚLTIPLE DE COMPARABLES
Ejecuta mínimo 4 búsquedas web diferentes con variaciones de query:
- Query 1: "[tipo] en venta [colonia] [ciudad] precio"
- Query 2: "terrenos [colonia] [ciudad] lamudi 2025" (o el tipo que aplique)
- Query 3: "[colonia] [ciudad] vivanuncios precio metro cuadrado"
- Query 4: "[fraccionamiento o zona] [ciudad] trovit inmuebles24"
Busca también en portales adyacentes si la zona tiene submercados (ej: "Rio Altozano", "Campo Golf Altozano" si el sujeto está en "Vistas Altozano").

PASO 2 — RECOPILACIÓN DE COMPARABLES REALES
De los resultados, extrae TODOS los comparables que encuentres con:
- Precio de oferta publicado (precio real, no estimado)
- Superficie en m² (construcción y/o terreno según aplique)
- Fraccionamiento o colonia exacta
- Portal donde aparece
Recopila mínimo 5 comparables. Si un comparable no tiene precio explícito, descártalo.

PASO 3 — FILTRADO Y SELECCIÓN
Selecciona los 4-6 comparables más representativos siguiendo estas reglas:
- PRIORIDAD 1: Comparables en el MISMO fraccionamiento o colonia exacta del sujeto
- PRIORIDAD 2: Comparables en fraccionamientos inmediatamente adyacentes de nivel similar
- EXCLUIR: Lotes en Campo de Golf si el sujeto es residencial sin golf (son submercado diferente, 30-50% más caros)
- EXCLUIR: Outliers con precio/m² más del 40% por encima o debajo del promedio sin justificación
- NOTA: Lotes pequeños (<150m²) tienden a tener precio/m² más alto — aplica ajuste descendente si el sujeto es más grande

PASO 4 — CÁLCULO DEL PRECIO UNITARIO
Para cada comparable seleccionado:
a) Calcula precio/m² = precio_oferta ÷ superficie_relevante
b) Para terrenos usa m² de terreno; para construcciones usa m² de construcción
c) Calcula el PROMEDIO del precio/m² de los comparables seleccionados
d) EXCLUYE del promedio los lotes <150m² si el sujeto es ≥150m² (distorsión de precio unitario)

PASO 5 — APLICACIÓN DE FACTORES DE AJUSTE (en este orden)
Aplica cada factor y explica el impacto:
1. FACTOR NEGOCIACIÓN: -5% siempre (los precios de oferta en México cierran 5-8% abajo)
2. FACTOR TOPOGRAFÍA: terreno plano = 0% ajuste; pendiente leve = -5%; pendiente pronunciada = -10 a -15%; irregular = -8%
3. FACTOR TAMAÑO: si el sujeto es significativamente más grande que los comparables, precio/m² tiende a bajar (economías de escala inversas). Ajusta -3% por cada 20% adicional de superficie vs. promedio de comparables.
4. FACTOR UBICACIÓN INTERNA: esquina = +8%; frente a área verde = +5%; cul-de-sac privado = +3%; sin dato = 0%
5. FACTOR SUBMERCADO: si los comparables son de zona más premium que el sujeto, aplica descuento -5 a -15%

PASO 6 — CÁLCULO DEL VALOR
a) Precio/m² base = promedio de comparables filtrados
b) Precio/m² ajustado = precio/m² base × (1 + suma de factores de ajuste)
c) Valor estimado = precio/m² ajustado × superficie del sujeto
d) Redondea al millar más cercano
e) Valor mínimo = valor estimado × 0.92 (precio mínimo negociable)
f) Valor máximo = valor estimado × 1.08 (techo de mercado)

PASO 7 — NIVEL DE CONFIANZA
- ALTA: 5+ comparables directos en el mismo fraccionamiento, mercado activo
- MEDIA: 3-4 comparables, algunos de zonas adyacentes
- BAJA: menos de 3 comparables o todos de zonas diferentes

FORMATO DE RESPUESTA — responde ÚNICAMENTE con un JSON válido (sin texto antes ni después, sin markdown, sin ```json), con esta estructura exacta:
{
  "valor_estimado": <número MXN entero sin comas>,
  "valor_minimo": <número entero>,
  "valor_maximo": <número entero>,
  "valor_por_m2": <número entero — precio/m² ajustado final>,
  "precio_m2_base": <número entero — promedio de comparables ANTES de ajustes>,
  "nivel_confianza": "<alta|media|baja>",
  "razon_confianza": "<explica cuántos comparables directos encontraste y de qué fuentes>",
  "resumen_ejecutivo": "<3 oraciones: (1) valor con rango, (2) precio/m² de mercado y cuántos comparables, (3) factor más importante que afecta el valor>",
  "comparables": [
    {
      "descripcion": "<fraccionamiento o colonia exacta + características clave>",
      "superficie_m2": <número>,
      "precio": <número entero>,
      "precio_m2": <número entero>,
      "fuente": "<portal>",
      "incluido_en_promedio": <true|false>
    }
  ],
  "factores_ajuste": [
    {
      "factor": "<nombre del factor>",
      "descripcion": "<qué aplica exactamente al sujeto y por qué>",
      "porcentaje": <número — ej: -5 para -5%, 0 para neutro>,
      "impacto": "<positivo|negativo|neutro>"
    }
  ],
  "precio_m2_ajustado_calculo": "<muestra el cálculo: ej: $10,379 × (1 - 0.05 - 0.03) = $9,550>",
  "analisis_zona": "<análisis del mercado, plusvalía, demanda y tendencia de la zona>",
  "recomendaciones": ["<rec 1>", "<rec 2>", "<rec 3>"],
  "advertencias": "<limitaciones de esta opinión de valor>",
  "fecha": "<fecha de hoy en formato DD/MM/YYYY>"
}"""

    # Construir queries de búsqueda específicas según el tipo y zona
    tipo_busqueda = {
        "terreno": "terreno", "casa": "casa", "departamento": "departamento",
        "local": "local comercial", "oficina": "oficina", "bodega": "bodega"
    }.get(req.tipo_inmueble, req.tipo_inmueble)

    user_msg = f"""Genera una opinión de valor profesional siguiendo el proceso de 7 pasos de tu metodología.

INMUEBLE SUJETO:
{descripcion_sujeto}

BÚSQUEDAS SUGERIDAS (ejecuta todas o variantes):
1. "{tipo_busqueda} en venta {req.colonia} {req.ciudad} precio"
2. "{tipo_busqueda} {req.colonia} {req.ciudad} lamudi"
3. "{tipo_busqueda} {req.colonia} {req.ciudad} vivanuncios"
4. "{tipo_busqueda} {req.colonia} {req.ciudad} inmuebles24"
5. Si la colonia es parte de un ecosistema más grande (ej: Vistas Altozano → busca también "Rio Altozano", "Altozano Morelia"), busca los submercados adyacentes para tener más comparables.

IMPORTANTE: 
- Calcula el precio/m² de CADA comparable encontrado y muéstralo explícitamente.
- Excluye del promedio los outliers y explica por qué.
- Muestra el cálculo del valor final paso a paso en "precio_m2_ajustado_calculo".
- Responde ÚNICAMENTE con el JSON, sin texto antes ni después."""

    # Llamada a Claude con web_search tool
    async with httpx.AsyncClient(timeout=120) as client:
        r = await client.post(
            f"{ANTHROPIC_BASE}/messages",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": 6000,
                "temperature": 0.1,
                "system": system_prompt,
                "tools": [{"type": "web_search_20250305", "name": "web_search", "max_uses": 6}],
                "messages": [{"role": "user", "content": user_msg}],
            },
        )

    if r.status_code != 200:
        raise HTTPException(status_code=502, detail=f"Error de Claude: {r.text[:400]}")

    # Extraer el texto final de la respuesta (puede venir después de tool_use blocks)
    content_blocks = r.json().get("content", [])
    raw = ""
    for block in content_blocks:
        if block.get("type") == "text":
            raw = block.get("text", "")

    if not raw:
        raise HTTPException(status_code=502, detail="Claude no devolvió respuesta de texto")

    # Limpiar posibles markdown wrappers
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1]
        if raw.endswith("```"):
            raw = raw[:-3]
    raw = raw.strip()

    try:
        resultado = json.loads(raw)
    except Exception:
        # Intentar extraer JSON si viene con texto extra
        import re as _re
        match = _re.search(r'\{.*\}', raw, _re.DOTALL)
        if match:
            try:
                resultado = json.loads(match.group())
            except Exception:
                raise HTTPException(status_code=502, detail=f"Claude no devolvió JSON válido: {raw[:500]}")
        else:
            raise HTTPException(status_code=502, detail=f"Claude no devolvió JSON válido: {raw[:500]}")

    # Enriquecer con metadata de la solicitud
    resultado["tipo_inmueble"] = tipo_label
    resultado["operacion"] = req.operacion
    resultado["colonia"] = req.colonia
    resultado["ciudad"] = req.ciudad
    resultado["m2_construccion"] = req.m2_construccion
    resultado["m2_terreno"] = req.m2_terreno
    resultado["recamaras"] = req.recamaras
    resultado["banos"] = req.banos
    resultado["condicion_terreno"] = req.condicion_terreno
    resultado["timestamp"] = time.strftime("%Y-%m-%d %H:%M")

    return resultado


# ────────────────────────────────────────────
# AVM — PDF DE OPINIÓN DE VALOR
# ────────────────────────────────────────────

@app.post("/avm-pdf")
async def generar_avm_pdf(p: dict):
    """Recibe el resultado del AVM websearch y genera un PDF profesional con Playwright."""
    from playwright.async_api import async_playwright

    resultado = p.get("resultado", {})
    agente = p.get("agente", "Powered by Brokr")

    if not resultado:
        raise HTTPException(status_code=400, detail="Resultado vacío")

    def fmt_mx(n):
        try:
            return "${:,.0f}".format(float(n))
        except Exception:
            return str(n)

    # Comparables HTML
    comps_html = ""
    for c in resultado.get("comparables", []):
        comps_html += f"""
        <tr>
          <td>{c.get('descripcion','—')}</td>
          <td class="num">{c.get('superficie_m2','—')} m²</td>
          <td class="num">{fmt_mx(c.get('precio',0))}</td>
          <td class="num">{fmt_mx(c.get('precio_m2',0))}/m²</td>
          <td class="src">{c.get('fuente','—')}</td>
        </tr>"""

    # Factores HTML
    factores_html = ""
    for f in resultado.get("factores_ajuste", []):
        imp = f.get("impacto", "neutro")
        color = "#1D9E75" if imp == "positivo" else "#E24B4A" if imp == "negativo" else "#888"
        dot = f'<span style="display:inline-block;width:8px;height:8px;border-radius:50%;background:{color};margin-right:6px;"></span>'
        factores_html += f"""
        <tr>
          <td>{dot}{f.get('factor','—')}</td>
          <td>{f.get('descripcion','—')}</td>
        </tr>"""

    # Recomendaciones HTML
    recs_html = "".join(f"<li>{r}</li>" for r in resultado.get("recomendaciones", []))

    # Superficie display
    m2c = resultado.get("m2_construccion", 0)
    m2t = resultado.get("m2_terreno", 0)
    sup_parts = []
    if m2t: sup_parts.append(f"{m2t} m² terreno")
    if m2c: sup_parts.append(f"{m2c} m² construcción")
    superficie_str = " · ".join(sup_parts) if sup_parts else "—"

    confianza = resultado.get("nivel_confianza", "media")
    conf_color = "#1D9E75" if confianza == "alta" else "#EF9F27" if confianza == "media" else "#E24B4A"
    conf_bg    = "#E1F5EE" if confianza == "alta" else "#FAEEDA" if confianza == "media" else "#FCEBEB"

    fecha_hoy = resultado.get("fecha", time.strftime("%d/%m/%Y"))

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8"/>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; color: #1a2035; background: white; font-size: 13px; }}
  .page {{ padding: 32px 36px; max-width: 760px; margin: 0 auto; }}

  /* HEADER */
  .hdr {{ display: flex; justify-content: space-between; align-items: flex-start; border-bottom: 2px solid #0f1829; padding-bottom: 16px; margin-bottom: 24px; }}
  .hdr-logo {{ font-size: 26px; font-weight: 900; color: #0f1829; letter-spacing: 2px; }}
  .hdr-logo span {{ color: #2a9db5; }}
  .hdr-right {{ text-align: right; }}
  .hdr-right .doc-title {{ font-size: 15px; font-weight: 700; color: #0f1829; }}
  .hdr-right .doc-sub {{ font-size: 11px; color: #9aa0ad; margin-top: 2px; }}

  /* HERO VALOR */
  .hero {{ background: #0f1829; border-radius: 14px; padding: 24px 28px; margin-bottom: 20px; display: flex; justify-content: space-between; align-items: center; }}
  .hero-left .hero-lbl {{ font-size: 10px; color: rgba(255,255,255,0.5); text-transform: uppercase; letter-spacing: 2px; margin-bottom: 6px; }}
  .hero-left .hero-val {{ font-size: 38px; font-weight: 900; color: #2a9db5; line-height: 1; }}
  .hero-left .hero-range {{ font-size: 12px; color: rgba(255,255,255,0.4); margin-top: 6px; }}
  .hero-right {{ text-align: right; }}
  .hero-right .conf-badge {{ display: inline-block; padding: 5px 14px; border-radius: 20px; font-size: 11px; font-weight: 700; text-transform: uppercase; background: {conf_bg}; color: {conf_color}; }}
  .hero-right .vpm {{ font-size: 18px; font-weight: 700; color: white; margin-top: 8px; }}
  .hero-right .vpm-lbl {{ font-size: 10px; color: rgba(255,255,255,0.4); }}

  /* INFO PROPIEDAD */
  .prop-grid {{ display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 10px; margin-bottom: 20px; }}
  .prop-cell {{ background: #f5f6f8; border-radius: 10px; padding: 10px 12px; }}
  .prop-cell .lbl {{ font-size: 10px; color: #9aa0ad; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 3px; }}
  .prop-cell .val {{ font-size: 13px; font-weight: 600; color: #1a2035; }}

  /* SECCIONES */
  .section {{ margin-bottom: 20px; }}
  .section-title {{ font-size: 10px; font-weight: 700; color: #2a9db5; text-transform: uppercase; letter-spacing: 1px; border-bottom: 1px solid #e8eaee; padding-bottom: 6px; margin-bottom: 12px; }}
  .resumen {{ background: #f5f6f8; border-left: 3px solid #2a9db5; border-radius: 0 8px 8px 0; padding: 12px 16px; font-size: 13px; line-height: 1.7; color: #1a2035; }}

  /* TABLA COMPARABLES */
  table {{ width: 100%; border-collapse: collapse; font-size: 12px; }}
  th {{ background: #f5f6f8; font-weight: 600; color: #5a6070; text-align: left; padding: 7px 10px; border-bottom: 1px solid #e8eaee; font-size: 11px; }}
  td {{ padding: 8px 10px; border-bottom: 1px solid #f0f2f7; color: #1a2035; vertical-align: top; }}
  td.num {{ text-align: right; font-variant-numeric: tabular-nums; font-weight: 600; }}
  td.src {{ color: #9aa0ad; font-size: 11px; }}
  tr:last-child td {{ border-bottom: none; }}

  /* FACTORES */
  .factores-table td {{ padding: 6px 10px; }}

  /* RECOMENDACIONES */
  .recs ul {{ padding-left: 18px; }}
  .recs li {{ margin-bottom: 5px; line-height: 1.5; color: #1a2035; }}

  /* ZONA */
  .zona-txt {{ font-size: 12px; line-height: 1.7; color: #5a6070; }}

  /* ADVERTENCIA */
  .advertencia {{ background: #FAEEDA; border-left: 3px solid #EF9F27; border-radius: 0 8px 8px 0; padding: 10px 14px; font-size: 11px; color: #633806; line-height: 1.6; margin-bottom: 20px; }}

  /* FOOTER */
  .footer {{ border-top: 1px solid #e8eaee; padding-top: 12px; margin-top: 8px; display: flex; justify-content: space-between; font-size: 10px; color: #9aa0ad; }}
</style>
</head>
<body>
<div class="page">

  <div class="hdr">
    <div>
    </div>
    <div class="hdr-right">
      <div class="doc-title">Opinión de Valor</div>
      <div class="doc-sub">{fecha_hoy}</div>
      <div class="doc-sub" style="margin-top:3px;">{agente}</div>
    </div>
  </div>

  <div class="hero">
    <div class="hero-left">
      <div class="hero-lbl">Valor comercial estimado</div>
      <div class="hero-val">{fmt_mx(resultado.get('valor_estimado',0))}</div>
      <div class="hero-range">Rango: {fmt_mx(resultado.get('valor_minimo',0))} — {fmt_mx(resultado.get('valor_maximo',0))}</div>
    </div>
    <div class="hero-right">
      <div class="conf-badge">Confianza {confianza}</div>
      <div class="vpm">{fmt_mx(resultado.get('valor_por_m2',0))}/m²</div>
      <div class="vpm-lbl">Valor unitario</div>
    </div>
  </div>

  <div class="prop-grid">
    <div class="prop-cell"><div class="lbl">Inmueble</div><div class="val">{resultado.get('tipo_inmueble','—')}</div></div>
    <div class="prop-cell"><div class="lbl">Operación</div><div class="val">{resultado.get('operacion','venta').upper()}</div></div>
    <div class="prop-cell"><div class="lbl">Superficie</div><div class="val">{superficie_str}</div></div>
    <div class="prop-cell"><div class="lbl">Colonia</div><div class="val">{resultado.get('colonia','—')}</div></div>
    <div class="prop-cell"><div class="lbl">Ciudad</div><div class="val">{resultado.get('ciudad','Morelia')}, {resultado.get('estado','Michoacán') if resultado.get('estado') else 'Michoacán'}</div></div>
    <div class="prop-cell"><div class="lbl">Fecha de análisis</div><div class="val">{fecha_hoy}</div></div>
  </div>

  <div class="section">
    <div class="section-title">Resumen ejecutivo</div>
    <div class="resumen">{resultado.get('resumen_ejecutivo','—')}</div>
  </div>

  <div class="section">
    <div class="section-title">Comparables de mercado utilizados</div>
    <table>
      <thead><tr><th>Comparable</th><th>Superficie</th><th style="text-align:right">Precio</th><th style="text-align:right">$/m²</th><th>Fuente</th></tr></thead>
      <tbody>{comps_html}</tbody>
    </table>
  </div>

  <div class="section">
    <div class="section-title">Factores de homologación aplicados</div>
    <table class="factores-table">
      <thead><tr><th>Factor</th><th>Descripción</th></tr></thead>
      <tbody>{factores_html}</tbody>
    </table>
  </div>

  <div class="section">
    <div class="section-title">Análisis de zona y plusvalía</div>
    <div class="zona-txt">{resultado.get('analisis_zona','—')}</div>
  </div>

  <div class="section recs">
    <div class="section-title">Recomendaciones</div>
    <ul>{recs_html}</ul>
  </div>

  <div class="advertencia">
    <strong>Nota importante:</strong> {resultado.get('advertencias','Esta opinión de valor tiene fines informativos y se basa en oferta activa de mercado. No sustituye un avalúo pericial certificado para efectos notariales o fiscales.')}
  </div>

  <div class="footer">
    <span>Brokr AI</span>
    <span>{fecha_hoy}</span>
  </div>

</div>
</body>
</html>"""

    async with async_playwright() as pw:
        browser = await pw.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        page = await browser.new_page()
        await page.set_content(html, wait_until="domcontentloaded")
        await page.wait_for_timeout(400)
        pdf_bytes = await page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "10mm", "right": "10mm", "bottom": "10mm", "left": "10mm"}
        )
        await browser.close()

    token = str(_uuid.uuid4()).replace("-", "")[:16]
    colonia_slug = resultado.get("colonia", "propiedad").replace(" ", "_")[:20]
    filename = f"Opinion_Valor_{colonia_slug}_{time.strftime('%Y%m%d')}.pdf"
    _pdf_store[token] = (pdf_bytes, filename)
    if len(_pdf_store) > 50:
        oldest = list(_pdf_store.keys())[0]
        del _pdf_store[oldest]

    from fastapi.responses import JSONResponse
    return JSONResponse({"token": token, "filename": filename})


@app.get("/avm-pdf/{token}")
async def descargar_avm_pdf(token: str):
    from fastapi.responses import StreamingResponse
    import io as _io
    if token not in _pdf_store:
        raise HTTPException(status_code=404, detail="PDF no encontrado o expirado")
    pdf_bytes, filename = _pdf_store[token]
    return StreamingResponse(
        _io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "application/pdf",
        }
    )


# ────────────────────────────────────────────
# CONTRATOS
# ────────────────────────────────────────────
from fastapi.responses import FileResponse
import tempfile, os, subprocess, json as _json

class ContratoRequest(BaseModel):
    tipo: str   # arrendamiento | promesa
    datos: dict
    clausulas_especiales: list = []  # plain-language clauses to be drafted by AI


@app.get("/img")
async def proxy_image(url: str):
    """Proxy image from EasyBroker to avoid CORS issues in PDF printing."""
    import base64
    from fastapi.responses import Response
    try:
        headers = {
            "User-Agent": "Mozilla/5.0",
            "Referer": "https://www.easybroker.com/",
        }
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
            r = await client.get(url, headers=headers)
            if r.status_code == 200:
                content_type = r.headers.get("content-type", "image/jpeg")
                return Response(content=r.content, media_type=content_type,
                    headers={"Access-Control-Allow-Origin": "*",
                             "Cache-Control": "public, max-age=3600"})
    except Exception as e:
        pass
    raise HTTPException(status_code=404, detail="Image not available")

@app.post("/contrato")
async def generar_contrato(req: ContratoRequest):
    """Generate a DOCX contract from form data, with AI-drafted special clauses."""

    # ── STEP 1: Draft special clauses with AI (abogado mexicano) ──
    clausulas_redactadas = []
    if req.clausulas_especiales:
        tipo_label = "arrendamiento" if req.tipo == "arrendamiento" else "promesa de compraventa"
        lista_clausulas = "\n".join(
            f"{i+1}. {c}" for i, c in enumerate(req.clausulas_especiales)
        )
        prompt_clausulas = (
            "Eres un abogado especialista en derecho inmobiliario mexicano con 20 años de experiencia "
            "redactando contratos conforme al Código Civil Federal y los códigos civiles estatales.\n\n"
            f"El usuario quiere incluir las siguientes cláusulas especiales en un contrato de {tipo_label}. "
            "Para cada una, redacta una cláusula jurídicamente correcta, con lenguaje formal, precisa y "
            "ejecutable ante tribunales mexicanos. Usa numeración romana (PRIMERA ESPECIAL, SEGUNDA ESPECIAL, etc.).\n\n"
            "No incluyas explicaciones ni comentarios — solo la cláusula redactada lista para insertarse en el contrato.\n\n"
            "Cláusulas a redactar:\n"
            + lista_clausulas
        )

        try:
            import httpx
            headers = {
                "Authorization": f"Bearer {os.environ.get('GROQ_API_KEY','')}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "llama-3.3-70b-versatile",
                "messages": [{"role": "user", "content": prompt_clausulas}],
                "max_tokens": 2000,
                "temperature": 0.3
            }
            async with httpx.AsyncClient(timeout=30) as client:
                r = await client.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers=headers, json=payload
                )
            if r.status_code == 200:
                ai_text = r.json()["choices"][0]["message"]["content"].strip()
                clausulas_redactadas = [ai_text]
        except Exception as e:
            print(f"AI clause drafting error: {e}")
            # Fallback: use plain text
            clausulas_redactadas = req.clausulas_especiales

    # ── STEP 2: Write datos + clausulas to temp JSON ──
    datos_completos = dict(req.datos)
    datos_completos["clausulas_especiales"] = clausulas_redactadas

    # Write datos to temp JSON
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
        _json.dump(datos_completos, f, ensure_ascii=False)
        json_path = f.name

    output_path = json_path.replace('.json', '.docx')

    try:
        script = os.path.join(os.path.dirname(__file__), 'generar_contrato.py')
        result = subprocess.run(
            ['python3', script, req.tipo, json_path, output_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            raise HTTPException(status_code=500,
                detail=f"Error generando contrato: {result.stderr}")

        nombres = {
            'arrendamiento': 'Contrato_Arrendamiento.docx',
            'promesa': 'Promesa_Compraventa.docx',
        }
        filename = nombres.get(req.tipo, 'Contrato.docx')

        return FileResponse(
            output_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename,
            background=None
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        try: os.unlink(json_path)
        except: pass


# ── CONTRATOS PERSONALIZADOS (MACHOTES) ─────────────────────────

from fastapi import Form as FastAPIForm

@app.post("/contrato/analizar")
async def analizar_machote(
    file: UploadFile = File(...),
    tipo: str = FastAPIForm(default=""),
):
    """
    Analiza un DOCX subido por el usuario y detecta los campos variables.
    Soporta: {{campo}}, {campo}, [CAMPO], <<campo>>, y blancos (___).
    Si no detecta patrones, usa IA para identificar los campos variables.
    """
    import io, re
    from docx import Document as DocxDocument

    content = await file.read()
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el archivo DOCX: {e}")

    # Extraer todo el texto (párrafos + celdas de tabla)
    partes = []
    for p in doc.paragraphs:
        if p.text.strip():
            partes.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        partes.append(p.text)
    full_text = "\n".join(partes)

    # Patrones de detección de variables (orden de prioridad)
    patrones_regex = [
        (r'\{\{([^}]{1,60})\}\}',   '{{{}}}'),   # {{campo}}
        (r'\{([^}]{1,60})\}',        '{}'),        # {campo}
        (r'<<([^>]{1,60})>>',        '<<{}>>'),    # <<campo>>
        (r'\[([A-ZÁÉÍÓÚÜÑ][^[\]]{0,58})\]', '[{}]'),  # [CAMPO] mayúsculas
        (r'\[([a-záéíóúüñ][^[\]]{0,58})\]', '[{}]'),  # [campo] minúsculas
    ]

    campos = []
    patron_usado = None

    for regex, fmt in patrones_regex:
        matches = re.findall(regex, full_text, re.IGNORECASE)
        if matches:
            seen = set()
            for m in matches:
                nombre_original = m.strip()
                slug = re.sub(r'[^a-z0-9_]', '_', nombre_original.lower().strip())
                slug = re.sub(r'_+', '_', slug).strip('_') or 'campo'
                if slug not in seen:
                    seen.add(slug)
                    campos.append({
                        "id": slug,
                        "label": nombre_original.replace('_', ' ').strip(),
                        "tipo_input": "text",
                        "patron_texto": nombre_original,
                        "patron_fmt": fmt,
                    })
            patron_usado = fmt
            break

    # Detección de blancos (líneas de subrayado: ___ 3+ guiones bajos consecutivos)
    if not campos:
        blancos = re.findall(r'_{3,}', full_text)
        if blancos:
            for i, _ in enumerate(set(map(len, blancos)), start=1):
                campos.append({
                    "id": f"campo_{i}",
                    "label": f"Campo {i}",
                    "tipo_input": "text",
                    "patron_texto": None,
                    "patron_fmt": "blank",
                })
            patron_usado = "blank"

    # Si no se detectaron patrones, usar IA
    if not campos and os.environ.get('GROQ_API_KEY'):
        tipo_label = tipo if tipo else "contrato"
        prompt_ia = (
            "Eres un asistente que analiza contratos legales mexicanos.\n\n"
            f"Analiza el siguiente texto de un {tipo_label} e identifica TODOS los campos "
            "variables (nombres de partes, fechas, montos, direcciones, plazos, etc.).\n\n"
            "Devuelve ÚNICAMENTE un JSON válido con esta estructura (sin explicaciones extra):\n"
            '{"campos": [{"id": "nombre_snake_case", "label": "Nombre legible", "tipo_input": "text|number|date|currency"}]}\n\n'
            f"Texto del contrato (primeros 3000 caracteres):\n{full_text[:3000]}"
        )
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                r = await client.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers={"Authorization": f"Bearer {os.environ.get('GROQ_API_KEY','')}",
                             "Content-Type": "application/json"},
                    json={"model": "llama-3.3-70b-versatile",
                          "messages": [{"role": "user", "content": prompt_ia}],
                          "max_tokens": 1000, "temperature": 0.1}
                )
            if r.status_code == 200:
                txt = r.json()["choices"][0]["message"]["content"].strip()
                # Extraer JSON aunque venga con texto extra
                json_match = re.search(r'\{.*\}', txt, re.DOTALL)
                if json_match:
                    ia_data = _json.loads(json_match.group())
                    for c in ia_data.get("campos", []):
                        c.setdefault("patron_texto", None)
                        c.setdefault("patron_fmt", "ia")
                        campos.append(c)
                    patron_usado = "ia"
        except Exception as e:
            print(f"Error IA analizar_machote: {e}")

    # Inferir tipo_input por el nombre del campo
    TIPO_HINTS = {
        "fecha": "date", "date": "date", "dia": "date",
        "monto": "currency", "precio": "currency", "renta": "currency",
        "pago": "currency", "importe": "currency", "valor": "currency",
        "cantidad": "number", "plazo": "number", "dias": "number",
        "meses": "number", "años": "number", "superficie": "number",
        "metros": "number", "m2": "number",
    }
    for c in campos:
        if c.get("tipo_input") in (None, "text"):
            label_lower = c.get("label", "").lower()
            for hint, tipo_inp in TIPO_HINTS.items():
                if hint in label_lower:
                    c["tipo_input"] = tipo_inp
                    break

    return {
        "campos": campos,
        "patron_usado": patron_usado,
        "detectado_automaticamente": bool(campos),
        "texto_preview": full_text[:600],
    }


@app.post("/contrato/generar-machote")
async def generar_desde_machote(
    file: UploadFile = File(...),
    datos: str = FastAPIForm(...),
    tipo: str = FastAPIForm(default="contrato_personalizado"),
):
    """
    Rellena un DOCX machote con los datos proporcionados.
    Reemplaza {{campo}}, {campo}, <<campo>>, [CAMPO] con los valores del formulario.
    """
    import io, re
    from docx import Document as DocxDocument
    from copy import deepcopy

    content = await file.read()
    try:
        valores = _json.loads(datos)
    except Exception:
        raise HTTPException(status_code=400, detail="El campo 'datos' debe ser JSON válido.")

    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el archivo DOCX: {e}")

    def reemplazar_texto(texto: str, vals: dict) -> str:
        for campo_id, valor in vals.items():
            valor_str = str(valor) if valor is not None else ""
            # Probar todos los patrones posibles para ese campo
            # Buscamos tanto por id (slug) como por el label original
            patrones_campo = [
                "{{" + campo_id + "}}",
                "{" + campo_id + "}",
                "<<" + campo_id + ">>",
                "[" + campo_id + "]",
                "[" + campo_id.upper() + "]",
                "[" + campo_id.replace('_', ' ').title() + "]",
                "{{" + campo_id.replace('_', ' ') + "}}",
                "<<" + campo_id.replace('_', ' ') + ">>",
            ]
            # También reemplazar por el label original si se pasó
            label_original = vals.get(f"__label_{campo_id}")
            if label_original:
                patrones_campo += [
                    "{{" + label_original + "}}",
                    "{" + label_original + "}",
                    "<<" + label_original + ">>",
                    "[" + label_original + "]",
                    "[" + label_original.upper() + "]",
                ]
            for patron in patrones_campo:
                if patron in texto:
                    texto = texto.replace(patron, valor_str)
        return texto

    def reemplazar_run(run, vals):
        if run.text:
            run.text = reemplazar_texto(run.text, vals)

    # Reemplazar en párrafos
    for p in doc.paragraphs:
        for run in p.runs:
            reemplazar_run(run, valores)
        # Manejar caso donde el patrón está partido entre runs
        texto_completo = p.text
        texto_reemplazado = reemplazar_texto(texto_completo, valores)
        if texto_reemplazado != texto_completo and p.runs:
            p.runs[0].text = texto_reemplazado
            for run in p.runs[1:]:
                run.text = ""

    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        reemplazar_run(run, valores)
                    texto_completo = p.text
                    texto_reemplazado = reemplazar_texto(texto_completo, valores)
                    if texto_reemplazado != texto_completo and p.runs:
                        p.runs[0].text = texto_reemplazado
                        for run in p.runs[1:]:
                            run.text = ""

    # Guardar DOCX en archivo temporal
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name
    doc.save(output_path)

    tipo_limpio = re.sub(r'[^a-zA-Z0-9_]', '_', tipo)
    filename = f"Contrato_{tipo_limpio}.docx"

    return FileResponse(
        output_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=filename,
    )


# ── PDF GENERATION ──────────────────────────────────────────────
from playwright.async_api import async_playwright
import base64, asyncio
from pydantic import BaseModel
from typing import List, Optional

class FotoItem(BaseModel):
    url: Optional[str] = None
    original: Optional[str] = None

class PropData(BaseModel):
    id: Optional[str] = None
    public_id: Optional[str] = None
    title: Optional[str] = None
    property_type: Optional[str] = None
    description: Optional[str] = None
    operations: Optional[list] = None
    location: Optional[dict] = None
    address: Optional[str] = None
    bedrooms: Optional[float] = None
    bathrooms: Optional[float] = None
    half_bathrooms: Optional[float] = None
    construction_size: Optional[float] = None
    lot_size: Optional[float] = None
    parking_spaces: Optional[float] = None
    floors: Optional[float] = None
    age: Optional[float] = None
    amenities: Optional[list] = None
    property_images: Optional[list] = None
    status: Optional[str] = None


def build_ficha_html(p: dict, images_b64: dict) -> str:
    import re as _re
    LOGO = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAA7oAAAEGCAYAAABLk+lKAAC3DUlEQVR4nOz993NkZ5rni30S3hXKe8MqFllF722zu2nasP2YnZmdnTV3r26EpJAUCpkI/ag/QVIoVu5KV1drZndmd6YtyTYku+ma3nsWy7O8RQEoeCD1w/c8+x5kJYA8iXMyz0k8n4gMAFVAmmPe9/k+tlQul3GcFqMNaAc6ou8B5oCZ6OEXveM4juM4juO0MB3NfgOOs0zagW6gM3r0AoPA6uhrD1ACJoERYAi4AlwFJqLHbKPftOM4juM4juM42eFC1ykaJSReu5Go3QzsANYBA8Ca6Pt1SOz2oqjuJBK4l4Dz0eMU8CXwFTDWwM/gOI7jOI7jOE6GuNB1ioIJ2/XALmA7sAXYDVwHbAD6CSLYorxtSBxb6vIEMA4MI4H7LvAy8H70b57W7DiO4ziO4zgFx4Wuk2c6CKnIu4A9wA3AfiRw10f/1wt0oTTmttijVPF8ZSR4Z4Hp6Dl3I9HcB7yGUpsdx3Ecx3EcxykwLnSdPNKDhOdmJGqvB24B9gFbUXpyH0HcVgrahShFv98e/W0PEsl96F4YB15HUV/HcRzHcRzHcQqKC10nL5SQ6FwN3Iiit/uB+1DENZ6aXC1aWw/t0XNuj17nBHA0ejiO4ziO4ziOU1Bc6DrNpg2J19VI2N4GPISit5tR9LYXXatpiNtqr9+L6n33I4F9EqU2O47jOI7jOI5TQFzoOs3CuievQwLzZuBh4HZgG+qg3NOg99Ievd4OVAP8LurQ7DiO4ziO4zhOAXGh6zSaEuqGvBYJ3LuAR6Lvt6PmUl1NeF8dKHq8MXp/juM4juM4juMUFBe6TiPpQEJ2N3Av8A3gThTBbZbANdqQwO1EEV7HcRzHcRzHcQqKC12nEVijqW0EgfsgShVeQ+NSlGuhRDa1wI7jOI7jOI7jNAgXuk7WdKA05X1I4H4X1cFuQOI3b5Sjh+M4juM4juM4BcWFrpMVJaAb1d0+DHwbjfDZCawim6hpGZiLHm0kG0NUBmajx1wG781xHMdxHMdxnAbhQtfJgjbCPNzvEKK465H4TRsTqVOoW/IY0Ic6OndRm9idAyaBq7jQdRzHcRzHcZxC40LXSZt2NJP2YeB70dedaHxPmlFcSzGeAyaA82j+7TFgGInsu5DYrVXojgKX8Rm6juM4juM4jlNoXOg6aVFCUdTdwJNI5N4MbCLdbsombqdR5PY8ErdvAa9F3/cAfw3sRc2u2mp43lkkdIdwoes4juM4juM4hcaFrpMGJTR/9h7gh8CjwHWkW4trEVyLup4GjgIvAu8Ah1HaMsBNKEW61rRlE8/j0fPPpvSeHcdxHMdxHMdpAi50neXSgcYEfR/4AXAHsJn0anGtA/JV4BIStG8DvwM+BS6g2tr4+xlAc3l7qC2aWyZEiIdxoes4juM4juM4hcaFrrMcOlF68F8jkXsDakJVi7hcirjAvQIcB94Efgq8D4xQfQxQJxpdtAZd37VGdF3oOo7jOI7jOE6L4ELXqZcBVIP7VyiauwfNxU0jVblMEJ3HgXeBv0MCd5TFuyJ3ogZUg6gxVi3MoY7N49HDcRzHcRzHcZwC40LXqYd1qJvyn6J63J0oVXm5Itc6KF9BNbgfA/8u+noRmKnhOTpRVHkVtV/fVp87HL2+4ziO4ziO4zgFxoWuk5StaC7uj4AH0CihTpYncssoojoEnAI+QxHcd4FzSOBWS1OuRgeK5vZTewr1bPTap5lf7+s4juM4juM4TgFxoeskYRdKVf4ecDuK7LazPJE7iyKp51BzqZ+jWtzjKMpaq8A1ulF9bh+1py7PoEZXp3Ch6ziO4ziO4ziFx4WuUwvtaFzQ36DxQTeh1OA26he5FsW9CBwEXgKeQWJ3OSN+elA0t9ZU6jkkboeQ2PUZuo7jOI7jOI5TcFzoOkvRjhpN/QskcvcjIVmifpE7g6K4Z4EPgX9AUdwzSPwuhz7UFCtJx+UpVBd8mdrqgB3HcRzHcRzHyTEudJ3F6OBakdtH/eODTFSeBz4HXgaeRzW5QyzeTbkWOoG1hPrcWiO6U6jL89gyX99xHMdxHMdxnBzgQtdZiE7geuCfI5G7j+WJ3FkkJE8DrwG/QNHcE0hoJq3FrUYXsB6lVScZLTQZvTfvuOw4juM4juM4LYALXacaXcBeQk3ujSxP5E6j1OAjwO+Bp4BPUPpyvbW41egkNKKq9b1alHkMT1t2HMdxHMdxnJbAha5TSQcStn8DfJ/lidwyipZeAD4CnkWpyl+i6GkaUdw47ag+t4fa369FdCdIV3Q7juM4juM4jtMkXOg6cTpQuvLfAD9AIreX+kTuHBKPZ4DXURT3DUKqcha0oW7LSeb6zqExRqN4RNdxHMdxHMdxWgIXuo7RhhpP/XOCyO2jvs7Ks8BV4CvgBeBXwPtolFCWYrIdRXO7qF2cz6K06gu40HUcx3Ecx3GclsCFrgPzRe5PWJ7InQFGgAMoVfnXqB53hOV3VV6KDvS+u6lN6JbR+x3FRws5juM4juM4TsvQDKHbQZhxWkbiZzb6mnbNprM0Nif3b1i+yJ0GLiFh+zvgt0jwjtOYc9uBRG4Xtc/QnUUp1lfxGl3HcRzHcRzHaQmyFrrtaKZpb/QYQF1xVyExYnWcYyh9dCT6eRpF16ai712AZEMJ2IXm5P6EMEKoHpFr83HfAX6DUpYPo0ZPjaId1ee2k6xGdyp6ZB1xdhzHcRzHcRynAWQpdNegSOEtwPbo50FgNRK8nQSRMY5Erond0ejrZeBs9BhGIngSiV9n+QwA3wV+xPJE7iSh6dTTwCvASbJrOrUQJnSTjBaaRdfTNJ5R4DiO4ziO4zgtQRZCtw3YDDwGPATcFv3cT0grjUfc4mJjihDhHQOGkMg9FX09hwTUSSR8r0a/7ySnH7gHNZ66nuWJ3JPAy4TOyqdpTr1rJ8kiuvFrz7MGHMdxHMdxHKdFyELobkURwh8DtwLrkMC1ulx7xLFI2lzsYQJkAkV4h1DX3pOom+9J4BhwJPq/USSOnaXpBG4G/gKJ3UGWJ3L/APwSeAulLzerqZONFkoyDsmuNRe6juM4juM4jtMipC10B4EngD8H7kbpyiZwF8P+vz16VD7nOmAb89OcLyHBewiJrSPApyjdeYzG1oYWiU5gL/CnwLdQtL3ymNfCNEpXfgH4ORK5F2ieYCwRMgbaqD2iG2+G5jiO4ziO4zhOC5C20L0OiadbgbUpPX+JkJLai4TvelT3eyNwPxK9x4APgOOoCdIBFOn1KG+gA9iNHBE/jL7vrON5ZlAa+R/RjNy3aa7IBYnbpDN0LXV5Bo/oOo7jOI7jOE7LkKbQLaFaz90okltPlLDW14kL39UoKrkb1QOfBw4C76MxNx8jITxBazQbKhEilvZ1oXRwe8yi87ETpZX/CXADOn5JMZH7GiFdudkiF3Qsugipy0m6Ls/QGteG4ziO4ziO4zikK3TbUXrxGiQ26qn5TEqJkO7cjcYWbUWpuXcCn6NOwO+jCG9RBG8HQczb3GHrKNyFIpfdsf9vZ764i6fkTkaPPuBrKGV5H2pGlZQZJGpfA34KvIqahDWrJjeOCd1aUuUNT112HMdxHMdxnBYkTaHbRuh42wxKBFHYj0TvNmA/arj0NvAe8CVqatXo0TfVMPEaF7J9qCZ5A0r/HkCR157o0R/9Ti9B2HUQhG48kjuDhP1E9He3o3TveppPlVFt9HsokvsyqtHNg8iF+U6PJJ/Nmp/l3fnhOI7jOI7jOE6NpCl0Z1Hn47zMuDVx2I+ivLehKO8bwDuopneIxr3fDkJqrQnatUjUrkGCth8J3M3Aluj7gejv7G+riduFalLjEct2gkCuhzmUFv4KiuTmSeQa1VK4l8KOkeM4juM4juM4LULaQvcUmm87iwREI9KXF6OEhF0vEowbgZuQ6H0beBc1rxohfbHThtKLu5HAXI8E7CYkbDdFj41I0K6Kfq+bkJpsotZSk+M1uY1mjtD1uhuJ9AnkKLBmTs2OitZzbOK1zo7jOI7jOI7jtABpd10+iWo4p5BYyxMmIgeBHcAdSPS+hBpWnUe1rMsRax0EcTsI7IoeW1GX6F3Ra69G0VubLxx/NFPMLkYbEuaPI5F7DDiNzvdlwizjSZoXIa3nmLWTrK7XcRzHcRzHcZyck7bQvYzEzyiKoDarXnch2gmR00EkQK9H9aZvoLFEV6g9Jdc6QHehqPEGgrDdCdyMuhtbra3V2SbtDJwH2lEk+mFU93wenWt7nEDR8a/QdTCJHB6NTG+22uRaMYdC0c6F4ziO4ziO4ziLkLbQHQI+Q919N5BcQFQKlawim+0oqtqLROgNKLr7IurQfAql6FaLTFrTrXiXZ3vcgMTtTlR7O4BEdVdGn6PRdKHPNYgi1PuRU2MEiduvgC+AI0j8nkKCeBwJ3yxFrzXfStpYyoWu4ziO4ziO47QYaQvdUTR65gkUKe1K8LdlVO85jmo/if7e0nvbYo+06CI0fNqExu68jBoufYK6M88SOjqbMN6OUpCvQwJ5N6H21iLGnSm+zzxhY3y6UPr1OkJ35xuBe5HoPQMcQl2uLdJ7GhhDkd6005tno/cwRXCY1CJe076mHMdxHMdxHMdpMmkLXdBc1ROoKdVAgteYRQLpc+AgEixrUAOnzYS61j7mN2dKgx7CDODdSKT/GngTRST7ov/fg6KYd6BZvRujv7HmUXlL1W4EceE7iI7JFBK9d6Eo/1l0Xj9A5/YommmcZmrzLIosT1C7iI6nLrvYdRzHcRzHcZwWIQuhewWJmbMoWlrrXNM5JHRfBX6GoqnrkPC8AaUDb0E1sOuRqDIhnUbqaVv0fLsJY362obrdzcDdKC15S/T6q5C49ZTX+VhTLZsHvBVFvvejlPYjqNv1xyjSexJFeadZXpS3HD3PBBK9tdJGmL/rOI7jOI7jOE4LkIXQvYpSf+9F4rTWSKelB8+gKOpR1Nn3MyQ8VyHBeQtKMd6JooZbkfBMS/B2Rq/zCBJoQyiavDl6Dz149K9W2ggNuFaj9O5dwO0otflzNObpABLAFuWtV/BaAyyr013qWrBormUIOI7jOI7jOI7TAmQhdEGRug9Qiu8aahOgHdHv7kDprydRlO9q9DiHxNBHSHBuQWL6oeh1dqL6WRMtyxG87dFzrULRQRtB49RPOyH13CK916Muzl+hSL6J3nPUF+GdI/k83zbCrGLHcRzHcRzHcVqArMTbJRTVvR9FRWvpOlxCzZ6uQ6nKB1AadJw5VIc5ghobHYhe50HgMeA21ChqDemM8LE03LyxnFm/zU61LhFqei1SvhM19fo2Erwvoe7N5widlGuhjIRukhpdc2LYLN3lHFvHcRzHcRzHcXJAliLuJKHp0Cpq68DcTRA9r3Kt0I1TRl2eP0cRwdeR4H0CpTfbiB+rEW62wEtCeYmvcxWPavNjS1w7J7by+1Lsd6nyfda0IedGD4qgb0SOjkeAd4DnUer6GRThXYoyycYLWbq8NRNzoes4juM4juM4LUCWQvcyGtHzNdTUqZPa0pfXooZQO1DUdmqJv4kL3iPACyid+Ueo6+8OJLTzOis1LmLtMYOE3Qz6/BOxh9WhTse+znKtGIYgbDsIs397UPpwPxKZHbFHJ+GaqOYcyOr4WWR1LSHKuxelNb8B/Aqlwo+yeKOpWUJEtxbBakLXasDbSH/skeM4juM4juM4DSZLoTuBOuseRhHaHmqrg+xHzaZuRdG8SwlecxI1sDqFIoI/Ab6FBO9G5ou4ZlEpSCejxzgScpdRA6xxJGSHUQfrUyiV9zI6tiZ0TeyaQJtj/sgcS801oTuIHA+7oq8WUe1BYs+6WfdHv98V/a2lgsfJ4ji2Re9hFYrIb0I12L9FM44/Rd2VqwnZOZKlOpfQ5xpEAruT9MYdOY7jOI7jOI7TJLKuP72A0k6HCWnES9GFBNi+6G+SCF1jGo04+jfAW8B/h8YD7USippHR3UpBZqJ2BB2bSyhF+yzwJao7vhj9ziQStWPRzzMkb7ZUjXZ0nLujxwASuv3IIbALHf+NBMG5AR07E8bdsefL4liWCKJ7I4rMPwr8ParLPsW1onQGHS/rvFzL9dZO6AjdhY6z4ziO4ziO4zgFJmuhewWlFJ9CTaJqqdM14XEjElsnqV98TKHU17PAD1D97p2o4681H8qScvQexpkvbs+iVNx3ou+voGiuCdokc2DrYTb2nojeQ5wSEpgmNDejqPhN0febUdfrNUgcd5HdeJ42QkrzpujxPeAXyIkRF7wmdCeoParbjsT8emq7Ph3HcRzHcRzHyTlZC92rwPuodvYWFA2sRRD1oBrNu4D3WF6UbQpFSv9fqMHVv0bNjnYhgZP2WBmrsR1FkdlTaC7wEeDD6Ov56P+Go9+t1kyqmZSRKB+Nfv4SiUoTnNuA+1B6+S7kxFhHEL1pOxAsFXt99Bq70fXxKqrf/RA5C6YJToMZ5kedF6IN1SwP4kLXcRzHcRzHcRqFZdm2EZrKptYvpxGjcy6iplJXUASwFjHRiUTNbiSqzrG8KGcZieV3UI3rpygqeDdKi+1cxnMbs0gcDqPPexR1gn4PRUyHkPAfZ+FOyXkjXk88jiKlZ5DwfQelM+9EovdOJD63ITHcRzrH1TDxbGnTq1Bk/mbgaeBF5EAYQse51pukLXq+vui5HcdxHMdxHMdJDyubtEkrW5DWW4vs8DakpUZR6etJpDlGqW3ySlUaIXSt+/LDSJjUInRLSMjcgsToISRglsts9Fz/EY0k+mvUFXpLje9rIWaQoP8Uidu3otcxgTtNMYTtUthnmEOfaxg1/3ofpRTfANwD3I4E8Fbk3EhT9Fo36L7oNVah87cPeA5FZm2kVK3P1xU9T29K79FxHMdxHMdxVjqdKAC2GzWXvQ2NEl1P6BFkdrtlxY4h/XgcNTb+ADUoHkZBt5pphNAdR2/uMBKufSydvmziYwuKEg6SjtAFibTLKAI4Fj0ep/Ya4mrMomjiH1CzpDPR86bROCrPWHrBZXTxHQXeRDW8e5CT4i50QW9FF3ov6aQ220iidegmGUA3T2/02rWkLdvzdEXvzep0lxpp5TiO4ziO4zhOdXpRludtwAOEINgGZLPbNBd7mDawgNo00hGPoEzZQ0hnvYbKQmuK8jZC6IJSj4+jDsNrqb0pVX/0+6tSfj9llEr9OhKjJeAx6he7Vue5Onq+q2TfUCpv2AzbC0j4HkSR7R2oidVDhIt8PTpey73+LLrbj8TtVkIqcq3PbSOGNqDz34MLXcdxHMdxHMdJSjsqC30ABRLvQv18LDDVTojc2ojUMiGAZSNRe5BWWIsCZncA96Po7j+iCSxXWCKg2CihO4TSeo8j4VNr+nIPOlibgC9Id8apNVx6Bx3QMkHs1hoNNNpRiu6NKGp9ltDIaaVhF+8MOg4XUU3vm6ie9n6U3rwbRX4HWH4TKBOr8U7aSaLGdv62omtueJnvx3Ecx3Ecx3FWEv3A9Ujgfhtpoo2E0sAJFPw8h6K0F5AWm0a2eB+yxzdGj3Uo2NkXPcdqFDDbhYJoz6BI7+RCb6hRQncMRfhOoFTmfpKnL79DNgJkGEV2Z1Co/HEkxpOIXYsq7kdi+XN04FPrGlZQyszvhHwKdUj+A2pe9TV0zLaxfMFr0d16aEPnb90y34PjOI7jOI7jrDTWAPeiZr9fR4J3EGmBy8xvZvsJajZlIteyYDuQoF2DAo/7UET4ZoJWWIeixZuQ2P23wNssMKGnUUK3jNKWzyHBs5baxgx1oDTXXUiIZCF0y9Hzvk1obf048hgkqfPsRKL8TpRPbim8jrBOamPoYv8UXez3oRviBnQRryLdbs210IYEbi8udB3HcRzHcRynVtaipsM/QRrINNQo6t/zFtJZX6Ls3osoCmsC19KPLWjVjlKUX0Ta6h4UHLsHlSquRiLYRpv+e5TKfI1ObJTQJXrxIyiddQu1iRmrfV1Dth1xTey+FXvdbyHhVavwsU7ANwDfQCd0mJVXq7sUc4RRRUMo0v8eqt/9BvLabKGxgteyByw1wnEcx3Ecx3GcxVmDxO2fRV+3I6F6Humq56Kvx7g2gluJCd5ZVL87ikTxMeCPKJL7I8LEnE3AE0hUl4GXqSgdbaTQvYrSeU+iUHMt3XetsdAA6TekqsQaVL2LBE8/8E10EGs5TqXo9zYgsXYfGmF0JYs32wLYbN5TKNr/GfLe3M18wTtAsnFB9WDX2TpUE+Cdlx3HcRzHcRxnYQaR3vkRys7cHv37URSN/Q0KZp1DEdx6JtHMEEaanoue+yvgx6jfz/rotS2I9gbKHgUaK3SnUcrq2ejNrGZp8RIf/bIRRfjqHhpcA5ZH/hYSrBsI3ZRrEVolJJh2oBD+H3GhuxQmeC3CewD4COX5P4oi5JsJdd1ZCd5uJKz3oJvSha7jOI7jOI7jXEsvajb1AyQ0tyEb/TDwa+ApVKZ4ldBF2UYHWZdle5gAniNEcyujvnNIJ7yNykOvAn+JdMI6FJy0cacfEzWoaqTQBYWTLxNmzNbSkKoD5X5vRqI3S6ELOpDn0JymXYROvD01/r3Ndt2PWmGfJOZZcBbEBO84upA/RxfqwyhFwS7kPtIXu/Fo/G4kqt1B4TiO4ziO4zjzaUPR228jgbkL6Z+jwG+BX6Cg1QRhlu5mZG9PI63VSRgHas1rJ5D9fQoFRoe5tt/RBKr1/U9IF/4VGj+0GfVYOoaE8AlgttFCdzx68WEUiq6lBrMdpS2vr/H302AGHagX0clbg6J97TX8rUV1dyGR9houdJMyFj2GUA3vBygH34rQB9HFnabgbUdp0hujr47jOI7jOI7jzGcV8BAqNdyNbPIzwAsokvsxqsftQXb7D1DwrwOJ3DlCI1jTdnNI7F5FWvEkEsuvIE02EXv9WVQO+4+oxPRHKFi1B/guEsKXgeFGC91J9OavUHuTpjYUHl9FYyPQYyjk/gqhw1etAqgdCfN96AScyuINrgBG0cV6Bl3kH6AmYXcQ6ndr6d5dC23ohlyLxO4hvJGY4ziO4ziO4xgWzX0Y6Zx+JGrfRtHcD1FA09KTe1BU92aUlWmCdoaQUdlJCCba/48BDwK3ogjx68yflzuDxhT9FInth6Lnvw0J8APAgUYL3Rmk1C11uRasTreX2iKqaVFGnb7eQyOD9qCTVesx60FibD9qtLTgMGNnUeaQY+R9JHaPIrH7EHAjiranEd21m3ETumHeR9eq4ziO4ziO4zjKqrwXCcr1yE4/grJg30fZmPHuyZaq3IXs7Glk159AerCXEEzsJqQ0b4z+fSsqXRxHmixewmrNp55FmbR7ovd0P/AqcKLRQncOFRhPRd/Xgs1T6iS96F2tTKPOXu+jgcXrqC2qW0LvdT26EJ7Dhe5ymUL5+i9HXw+h2oBb0U1g3ZnrxWYhbwL2Rs/nQtdxHMdxHMdxxGY05mc3Eq5DSCe9jezzuL4royBnPLg5gSKxT6OeSDb1ZC0KXm1AEeD9SEdtB74Tvc4ZpMvi3ZsvAX9Ao422oQjzbtQo681mCd1pkrWYjnfmajRDKEf8I+QpqDWy3IbE0m7k/biYzdtbcVxFuf9nkDfoG6hZ1T50k3RQ/3Vijc+uQ4L3PLU7ZBzHcRzHcRynVelB9vZ+ZC+XkD3+Icq6rAzqlZHItU7LoCjuZ8DvUGmnlQ72Id20GmmnH6AMznXIJv8Wit7+tOJ1yqifz7soKNkbvbebgC3NELqTJIvoxmmG0J0m1Ifeh7wLvTX8naXCrkfeiaPUNz/KuZZp4DQqej+Bzs930Qzejcg7VE/0vw3Vgu9F3qpTuIPCcRzHcRzHcdYhMbkD2dpTyAY/iNKRq+kcaz5l/zeNanqvoGAihEzY9ujr56hPz06UJt0Zff9dZPufqXitYRRVPosizr0oOLmr0anA8RB2EtFXjj2awTDwBWqMNJrgfVg4fi+1jydyamcURXd/Dvwtqg84Fv17PY2krPHZHtSi/AGUAuE4juM4juM4K5l1KKK7AYnSMRTIO0X1Ek3TbpXBTStLrfw9q+m9CLyJUpzHo98ZQOnIW7k28DkTvY/z0fft0Xvc1uiIrnXXaiNZdNYGCDdL6E6iyOEBVOC8jqXTl+2zbkTNrF4gnCwnPWbQDfYc8hBdQqnMe1H6Q9K63Q4Uhb8D3WhXUPG7nzvHcZyVhUUYSlxbQhV3vseNuXjkwnEcp1Ww3kMbUFCoDdndx5G9vFSAKb4uVitHLUfPaYL3CsretO7MbSi9ebDK34Km+pxDUeZeopGhjRa6negN9lF7amlc4Tdr3EsZzWM6gsLiO9BnWIoSElt7Ub74mazeoMNlNArKhks/jjw/VrdbKyV0g+xErdNHkMfqE+Z3enMcx3GKTXy0RQcStR3RowftBX0oO8v+3ww0c8DPIUNsGhlY48g5PkkYoWFfZ3AR7DhOMelEQtfGvc4hgXsK9c9ZaG1bqM9SNR1oz2ER37j9Hu/zVI1RFJS8irRmN7Cm0UK3D3XE2oTaTNdCmbB5zGT0vmphHJ3M06hjWK1Ctxsd8MHs3poTMYqir1eR4J1A0fQNJBO78VrdR1ENwSV0A3lzqmR0onvdhoKboQhLN5grL/J1of+r9vdGqeLrQv9W+bPNglsqmhOP6NhjlmDgrtS5zB3o/Fc2iqt2Hdj3Vqez2PVR7f8WuqbKVb6vdv5q/f/Kn+1827n2dSKftDN/fMUqZLitQ07pfoK4XYUiAn2E0YLx69Lub7vHp5C4HUN70Sjai0bQHnI5eoyivcnE8EpdF2qhnXDM2yt+Xuw+j9+L5mRw5tNW8aiWtVBi/noXHxWTBh3oPqzMvKu239ayF1TuLwv9HH9+E1tLlUdW20PiP9s+EN8P4vt/K1yDXSh41IuO3TSKop5H6181Ko9z5f+1oWNVec47kO2+G63XpgUvI3u82rmaRDptGNXpdgD9jRS6JbShXI/my3bX+HezaLO4THMjatPRe7hAslFB7eji8FrPxjCOurldjb6fRU2qNqCbpVbakeGzD40xOgf8Gt3QTm10ont9L1p01iKj0QyV+OYap5qgiG8ilSmC8f+r9jwwf7GtttkttkGWK/622uvGN7Xp6GFG7zBamK9E/xbf+FrdCOtA2RG70dgAuwcrjaz4wyJslSPl4uekLfZvcQ9wNSEN1xop1a6taud1IQdG3JAx428cnesr6LxPIUEzXfFwGkcbQdT2oVKibWg/sE6eW6N/X4uEbS9BDFcKq0rj277G14HK6K7d/xejx3lkR5wBTkY/XyVEgD3iK9rROdkQfV2Nzk8P4dxUYudgCt2Dl1Bw4iQre1ygra12L/QTHDm9BEdk3Jlg17ld01NobTuFjudyr9NeZBvsjt5HfF+O2wgw39m8EJX7eeUjbnPE9xv7rCbkF3OgVOsXVLn/22OS4Oi6QtgXJgiOb1srinLPd6Drpgsdpxn0GcdY2GFXLVBg12JP9HyVv9OB1uevoxGiPdHvjCD7/hzVnS0zSKeNR//fBnQ3UuiuQjNlb0aCt5bXLqM3bovVQh6DRmCCO2mjI7tpa41gO8tnChWl/45gkN6LjJmkkd216Lr9U3STPY8MF2dxSiiL4T7UEt6aF/QRBMxS0beFomjVhEgtInchj/Vikd5qG1C1167mwbWN7jJamM+hyI5FfK4QjN8RwozxVhK+/aivwaNobJeVrdi6GDdo4lGbyjRRqnytZDEjKM5C3vhqD/v/ShE8y/zzPYXO9VD0GCWUUQxH/2Yi5wohotdK5zovlNB+24OMpR1IzG5HRvVu5HhbQxBOlenLSfuIVBIXvXGHlgmwK8imOYxKok4g8XASXSvjrGynSAkJoeuQo3ofOodr0ZrSxcLrgwndEXQ83wF+j471Srjf4s7CLiRse9C1vh5d+1uir5aGas6DyuvfbPAJdM0eQcfyDMs/ltuA76OeKhaIqBShVq8ZF0IL7fUL7emLCd/418rnqXz+hewSWHj/HyVkc1wgOLzGCJmHF6L/HyM4yPIqfNvQNRV3QNheWMvfGj3ALnRv7yKUi5gN0Ifu/UeRo7wdHc+TwEvo3q6GRX1No5WA9kYJ3W40z+hx4HZ0Y9Wyicyhk38WDQhuptAtowM9QTB6avkMdiMliSY6y2cGdWD+DeGiv49kaczmWdqI2qmbaHkdXQfOwpiT4CE0xHsX80UuLC4mK6mWNrRYKpGxWKS22v271D292PuoFuG1aJ9F98ajr9Y4zda2r5AIOo2MCIsINnPNS4MulMVzJ+pm3s/CzX3iRsdCkfasWMxRspChs1BE37z1FtEbJZzrU8hYPIMaeFht0wSewrpcLA1yPTKSdgI3onmPu1D0dhAZWfGIbRbEHTmV2Wu2JtyA9qRh5AQ7gqIVB6LvTxIcYCtBoMXpQULoTrSHmKM0HtGttlaYYJlD99Se6G8uIMFxnvyKiOVQIojZVciJY9Hw9dH36wgidx3XRnPjws+YQ9fqMFqvTiIbaLmpy23o/N6CGn+up7oDfDlrfrW9o/L54v++lE2/kL1RuS9Uy/axDK9xtLebs/sy2vuPob3/fPTzEKEWNe/7gq1ztWDHdwAFn1ahY2IOGbsW+9A1uzb69wm0Rr4MvM3CTsB4NgBE56MRQrcDeVCfQONaNlO76JtFC/0Z8rFAJfFeGLZwNLoe2tG5Ook6MtsNljSyW0LX62bgHmSwXkAzvlaa8ZGEElqsdiNP/BoWbkBXq8MoD9QrvCqjgfGaPnOgnEUG7gG02Z1FkZ4hwvzxomH3jxlfVtuTNxYygJaD7RfxVNZJFBk5i87zR8jIOYwMnavk26OfRzrRdbULCZvbkGPyerTWW8SqWi1gMzDDsAel425EUeebUXTrJNpfPkLXhUV7R1kZ10YH2m/vQ3bjHdHPA+gc1hJtL6Nrogcd60toTb1CstKzItCF1tbdKKC0kxCx3YhE7SDXpuSbuF2sMewcEiIXgUPAH9A1uVyhaymwawkp6c2+N5e6ppbjdI0L4Hja8hTzs39OIWfXl4QssPNIB03S/HvfHEhx2zeefVUNe8/x+7YTZdisiX6u7OESd1oNozXwdTRGdLG0+XZ0rVtq9RwwlbX4KiFPzSMomnsDutlqZRZ9yPPkY7zLQmmSixFPk3Aazyy6SX5LGFF1H9oAFrs545i3dAfyLn+JFv7KgdVOwI7ZapJ1WW9V4uuARZ76CevJdhR9uoew5n2Fun1/jhb3I4QNL+9eXsM8rJWNyFYClhFiHXxB53oTMkZvAR5Dxs3HwFtI/B5G53kCX18Wo4NQ57cPicS70LFdg+6vyjrvPGKpen3ofW9BUehvoGvjC9Rk8TO0BlwiH0ZvVthe+yASubtILoRs3elFe/0t6Br5nNYRuhb52okc+Hejz7mT0A/DBEQ9qfiWtnwZOV1+jeaajqXw3tsIKdWWht7KxDM8KgN96whR35tRqc8lJHIPIxvA9oXLNLfkZZbgcJsjOOys90o17Hd6Yr9jTmBLM7YAojUutZT1aWRv/xp4Bu2Tizn8uwlOMWtyNZG10LXw9JMoZXmQ2jedOXRCL6DFPi9Cdzb2fa0LR2UKntNYLLL7W4I37QFkcNZ6D7Sh6/k6ZJx+gYzR0ZTfa6tgkTwby+FUx9YG2wgGCULoZmTsnUM1528iQ+0IWhPHyH9WgZVtrDSRuxAmfgeixzqUwrcfCbXP0Ji0L1AE5QKhXMYRbehe2Y3uke8igWsRApvvWETakUDvR0ECWwe+hq6JN1Dq3hFCY8xWE7wDyHlxCxK8q6j/fNq1YrXavSiqW3QsInYzoTzoBnTNmKBYznprIuMicrL8HEVz0+pPYutgERxRWWPO0F5CF/jtaN2/A53bI8gR+glKIT+LzkWjG1lNIRFuzac60b21Edl61Rq+dSC7ZhVB6E4hR/5HKGOtK/qdrcjhvxpdF5PR7/xblPG0lA40x4+tGZPAcJYGaA/yMP0FMtaSjngxT5Kl8uWhJrLeyOxyG1s4y2cO3Sgmdkskvy5LyJC6E/gOWmw+Y2U3DVmMeHqUUxvm+bUIjwmhG5HT8CvgQ+CPyLt5Am06eRVCJnR9DaxOXPSuRZv0nch7/wZybnyMjIsidefMii50P+xDjWweRM7HdcxvktIK2LVhoncHSsv+Bkrjew3ZR+coZlnDQqxBjp+d6LMv55yazWY1fz2L/3oh6Eb7weOoWc+dKAsgrcwpi+QOob3m58huupTCcxvxWkrfFwIWIOhE176VNuxE3YdPAQeRs+sTVPYyRIiwZs0UsqMvR9/3RO9vFxKX1a4RazAVvzavor3NUpFBa90twD9Bgai1hLm97SytAS2DeDth/OsYcCEroduBagX+Ci3KW0jWjMnqAk4QPBh5MOQs7bCeCG2rp2YUgTJKCf09wZt4H+FGWooSMqa2IqFrTYOO4WK3EjMwKhc4JxnxaO9aQlrjA8jQfQ54n7Dx5E0IVTaHcBbGzvNqgpi7H3WPfwOJmivkYy9sNCb69gHfA76JIlgbKHYEtxZKhEYtq9EacBOK8r4KvIiMX4vwFJlOlNGyBzkv0rBRTTz0UftYy7zSjQJIf4nO/w1oX0hLMFrW4jDwKfAUShu9kMJzx/GSvtqIR3vXIhG3H+0LXyIb4H2U/XOe7O//OUIjxStoTR4kjJCsNh3H0p1tYk1n9L2VZJyL/e4XSPv1ojKuPiR+fwz8j8jOWYhOQnf2jui1LgNnshC6Jea3Dd9OstE61t34NDqB77JwK+lGE5+5VSvWfKboG1CrEBe7Vr9yD2GzWIp4k6U/QUL3N8gpk/c00kYTn1vnLJ825DW1xk470CbwDBJCZuzmqX7XjBk3aGqni9AVdQNaa0zwvo4MjVZMWa2GpZ7uQsfge8jQ38byUlqLiAm29UjwbkaRvfvQfvY6Mn6L3MF7LYrQb0cRrTT2DnO2WSOmotKDzvX/HGX32DFK6x6whkkjqETmaeCXyBZPm5WwdqWJ1TR3I2FpUd7bkMh9IXocQSIyS2foRaTL7oneRx9q/LcX2SCVUd2Z6D1djb0vG8dXaTNfQZ9jHVrfb0Zr/Z+izN5fs/DaNoBKWDYjoXsVlXwdy0LoDqB0okeRV66b2hcrS5m4hCK5L7J08XGjqLemwGqNhzJ4T059zKG04xfQzdSP0kIGqb051erob/4EiYtnkdfTF/BA5ZxDJz36kddzEG149yCHyxs0f+a4kw5dhJEqm1D05iZUK/cp2ieLKmhqoRN5529DGTRfQ6J/HV7z3YFEodX0W433L1GU5DjFc663I/F2BzrvNlc1DSxIUdTMuh50fv+36H7YHP1bmvdApcj9KSqVyQoTPSv5Pq6HNrT39xKa1+1CTq9fo7nR58lub7iKUtqPoD1pHVqX7wY+QGK18rXjs+fNoTLDtfbyHIrwPoNE9CChI/2/QLrwaJW/KyEH2V1oXSxF7+MTMhC6Hegm/DO0WFnnq1qxVtKfIqPtZdIrfl8uVucxQEhfroVZFIpfLOTuNJ4ZtIj/gWBMWlfwWsXuWlQfcyZ6rrfIRy15nvBNLDvaCN1lNxLSh36H0unzcC2642f59CDDfxU6v3ehlMLnUH1T0QTNUliq7o3Aj4BvR99vIH3jvsiY8z0e4b0BpTP+B2TkFamuuxMFR24gNKNJmyJeO73Aw8D/ATV1tfGIaYvcqyhq9gzwX9AektW1U88EE+daugn26zok9p5GNdXHyUbsltG+8xnaiwaRHXI3crSfZH7DN7tObSawjVla6Pq1SSm/QAL6SbT33YkCS/8913b+HkDlLDciB7E1u/oEGEpT6LYh1f2vqK/5VBndaIdQdOw5FCLPC9ZdbAPJUrFnCYOinXwxhRb255AneRXyHtWa3tSGxO796Lo9jbxNnsIsVmItYTOwOc89BE/v0+iabLbYdWMmHUpofepGa85m5Mn/FUpXTWPkRx7oQCLnNtTj4+vIeFvdzDeVcyyleRMy+Dai/ez/iQy9ixRjLbZ1bD3p111b7WkRjkOcXtR193+PDP2NpF8OZLb3lwSRe4Rss0V8X0iPErpOLJV9PdIp/4CaGmaR3TWESiXuQ/fsKpS+fBeK9o4w/16zgN8VJHitxGqh63gaCemfIufXvug1nkRi+i2Cnd2G7o3voWPQjta8N1H560iaQncrKhh+hGRiAXTBTyAV/0cUZTtNvtKyepBhsYXaGxpYKvYUrTO7rdUYR6keqwiesQ3UvpH0oJvrG6gN+gV0M/si7ptZo1mNUlv/GVp/f45qZpq59sTPfxGjKXmjCxkyXSjDaD3wj2hDL/qos27kwX8URXLvINl+u9Kx/hF7COnuTwE/Q9GNvDtgO9A+PEi6acsgo9siSkXAzuXXgP8doR4y7TU0HmD6DUEcNeI4eVQ3Xax+3xzefcDfowZPaTu8J5EQfQet2ZZx8whymAyh9GkITaE+id5TL0pPPs/iInwCeAmVZTyB1oZp4M9RtPoUuna2oxreW6PnnkaOmtdQtuVcWkK3Dyn7x9Aim1TkTiOB8B7wPPkc6t2LNo6NJGtRP4UMkLx9HicwgrpXbkALRTe6qWrdVAYITUG+QBtHq6UTLgcXOI1jFUph/gt0Df4jyjJoloFnRoxfA+nRhpwae9FaZcbDG+Sn1CcpvUjY/jVKQduNMqicZFjjpc3omK5HEfH/Hu1NeRa77WgvzaKLtgndPH9+o4TE/teA/w2Kkm3K4HXKKBPkCMpq+wVyjDbKdnGBmz7WsHI/cob2AP8J9TpKO6v0IiqT2knIMroVOSkvoNLTEWR7HEGjhH5J0E8XqT53N84wKsN4C90TY7HnBK0Xf456OJgj6Ay6nv/rfpiW0N2HNqg7qb3G0YgXwD+LVHgePdP9yLu8mtqPWxldXKdpndSyVmUIXXs3onqDvdSeot6BRPJ3UH3CFPI4FWFTzRJLsXKR01gsjejPkSf15wTvajPwayAb+lETkifQWtUOvEI+98+FaEdG0gOo7Mm6yfY28021ACaW9hEi//8nlHWUVydsOzKC047mWtryDPlPXS4h4/1h4H+NIrkbMngds02PIlHwUySGGhWQsfPrYjcbepEN8ENkn/57lFKc5vm10r/nUUbvg8h2/hoSsKMofXgiet1zzB8lVAtzSD9V6/zdhXTn36AAazsS2K+iTBaL+KYidNegAb/3ELohJmEMpU08g7wDaQ6lTosuJHJ3kGwDnkMC6ige0c07s6j5wm9QNGE9yVKFBpA4/iu0UT9FMdLFssRHCzWPVaipy/eR8+UlmieAfF5idlht1tcJZTKvUwzHaicS6t9FTpnbSN7bw1kYm8G9Gx3rTuD/jEaD5LErexs691l1RraOr3klHsn9XyGben1GrzWBnPEvIEfoBzS2j0x8vrrbCNnQg7I5niT0CvqMdB1dV1Dk1up070B283cImZEvkf6I2PXAf4eE7l70WUfQdfz36HP+1zUujQ3FhhdvI/kCNYFU9wvRI+2h1GnRhU7kNmpPW7Z5wOdQF2kXuvlnHN0ov0DpGN3U3gSlhG6+W9ENNoQ6350n35trVpQW+N5pDDYC604kgo4g72ujU5jj0Vy/DrKhB+1NX0dr2FVUBpRHMWN0I9vhT1ATkZtQZNdJn3bkpDdR8X9FtXV5a5DZTshMyGKtyHNmiTWbe4gwJzeLmlyQ3X2SEPl6l8Y7xmzckztAs6ULOROfRAL3/4PS09PMbDiDUpJXoeyR/SjV/ttoX7oFBTLT6hdyH/C/RH1xtiEH3gjKSPgPhJTp/8pyhe424F+ibstJU5ZnCaOEXiO77mBp0I2i1Wupvf7YGmyNoFz0ojRBWOmMoIjILnRN34RSBGuhhK6TW9GIrVNoM8mbQdEosjQsrOZqmuCpjzsUFmt0UZk2tZAYq3z/1X6ufF6LXtpXezSDdmQsfR15OM/QnDFneTYwWwHrurkDNXIaRplRh8hnqmYvYQb591C5SK0zzBuFpbvaYy72sH+z9SNutFtEspP8zPq1tWgbSokF+DfA2+Qr8m/jklbaWCFzSj6AolT3kZ3IHUci93UkTl6nOZk+WZ5rJ2Cj2q5DmTPjwP+AovlpchL1A2lH6+I+dE3fi9adB1Bp6ovRa0+SbG/qQDb5n6F07FtRNnEbCip9CPw7NG3iGhtnOUK3C3XYuhel9Sa9YK8icft7lMeddmg7TdrQxdJN7Z8zXp97ZYnfdfLDHBKoz6Ao/gCqdajVwdGOIrv7gW+hBeAgKy+FOctuijPIkP8KRcwnuHYguRkJcxX/Fhdd8fdn4tQW6vjPRjzVKi6MS7HfNSO3izDqJ/7oonFGVwk5afagyO57hLb+jSLLz2piZJyQjlVL7Vf82qxslrWQKK90hJgDoyP2aCbWpfU6lDY2iuqyTjbzTVWhD12Lf4re5w1ojW22ELHUvgl0PY0jETgW/dsUusamYv82h8673e/dhAYwa5B47yd0Qe2keYa9rU9bkdidBf4fhBq6PFBi/tqd5vPmtXzCRO6DwH+L0pbrsadrYQyJjDdRJ+5XaK5tmmWmWxntdSPoWNq9F98fKh8LOa6rPXfl71V+NedXJ83fH6yEYQ/q53ARNYZKc3zrLNJzf4/WzidRAHQDKp3YiCK730MZBO8gu/gyoRt63Eaz7I51aI94AGnNfYTM2lngbPR8/xGlSF+hynW1nIO/D9Xm7iX58OoJJABfRDfbhWpvLmckjUqY0D2BPA5OcbD25M+gm3QttXc9tEVlC7o530Vi7BL5v8aLwhQ6Pz9D48jMSVatBisuahbbtBbarBaqISpV/E47YUNrR4btWuQs2YWibZuin9dG/9+IsSnWofcedC220no0g6LUxwhz+yoFrBE/h/HfiTtB7Dwutc7b39tm3I/Stgai703wJJm3ngbWyOZ6FNk9iVITmxHFr0YvcDsSuU8i26Gf5ohcG/1ngvYioVnK0ehxATkMTNha196Z6PvK66YDnXvrdrwLnYttyODbFP17f/Q7zTB+42J3DH2Wd8mH2C1VPNLEzlHesJrcf4UCR1mJ3Kvomn4JRXLfROtCs2ySuFM6beaQ4PkM2QkTBMFp63s8Y8Pehx33yn0h/rD/m4t+v4sQmY7/XhuyA/sJDi9zenWidSLJdJrlYo7Qveg6+xhprzSzaGeQ2P3PyKHybeTA2YU++w3IDroXjaI9i9bbCwQHvDkOe5HI3UFYOwejf29H6/ZX6Hr+B+TEv8wCTvx6F9oB1HTnHmS0JU1ZvoJmKr2GLsS8piwb8Qs/yY1pUSYXOMVjHN08PyWMalhV49+W0E25FzUDOo08WHkwJhpJXFCkabjMofNzCnVrH6rhPRiLRfySvsdqaczxr2b8DqKF+jq0Zt6GFvCtBNGbJRbpuxmlKw7TuJTWrGYlWmnIl8D/O/pqrxH/bAulmld7X4sZ2JUC2gwdc2hsRU2hdqP1YiOhS79F8rPGDKxBlFFitdlv0/y1pxulKH8fpdA1S+SauB1GDshTyDh7H11Dl5DBNIrskvi+X3kNLJQNYOfBmrGsQdfBjcDdSPxeh66RVcj4baQIa0fX5jeQLTaB7LE89BHJqszBsi+anTkQZxUSuf+aEMlN+zooo2v5MPAH4FfIsVE18tVAss70OY8+77+Lfo7XA1eL5pYJgjVu78P8+zy+HrQxP/26cq/pRff+duTs2oLsgDVoj9hKEL9ZlzmYPbIOZdT8CDk+Dqf8OrPIwfpr1BPkczQubh8hQ3IbOg5ThAyaKcIxtQh8N+HY2Hm5ioTx0eg1fkdwZix4PdcjdNtRfvRDhAZUSU7QKKodeg7NRirC3L+49yfJ4tCFPLid5GMTcWrHUl9eQ4bJepKNHLL6yDuQB+s4WgDyWDNXNMyQLFOM2YjDKPL4GaH++ybkWb0LXVeryS61rg2JsZtQ+tIJGlebl1UKuxkjV9F+8mHF/1WjWlR3uZgBEfdCr0FGzH3IqbEVOTbMs58lpei9bEIZJSfR2pN2TVYSrLvyd5DQvR45XxopOizl+DyKBHxMSJ87hyK6Y8yvx10uE4R7vw2N9nkZnZtb0fWxHx2bDeiYNCrK24UM8MeRsDenUbNHD2UR5YunQ+YlfdlE7r9E0fXNZCdyzd7+JVonh2l+8CV+PtJcB2xfGEOZPofqeF9LHZuFMsTiTvT4NdeJ7u1+JPTWIOF3P7r3txO6FvdUed60sL1hG1p7vkWIpqbJHLruPkVR27fRZ70Pfe4t6FjYOLFuru2XEj+WNjv3Ijqfr6Hr+PPo35Zcs5IuqiVkvP8F8k4mPSnW7e0PqElPEVKW4drGN7UQr5naRLFmGzpiFkVjnyeMvthY49+aR+86lEb4CcGYWmmkHdHNyuufFSbIp5H38gLydn6AxO6fofqVjWQT2bFrcTcSGY3usplVVBe0HlsqabOYZn7jQUtpfgeJ3N0oenYH2jfXkf6s0Dh2vnciA+Pj6H1dzej1FqMNXdePoXS1G5Gx1yjBMY323rMoCvA2Sts8EP2bidus6tYtRRoktq8iof0pSh20VML7kRNqAzICsxa88bq9J5EDYARFuFuxcWY8xbzZ9KPz/U8Jkdy0z3cZnc9DqA/Or8iPyIVwPrJYAy0wZVHCZgYXbG8YZf41+AFhLM+9yCl5I1qzB5H4y+LYxG2BB5Hj/aMMXgf02c8iR9qXSPPdhiLKO5Gw7yZExeONO+38jSJH5CHgCyRuj6N1dJIar+WkN1cnOhn3oM0ryWZlefOfA2+gTSfvKcvGLLpYraallto6u6C2IAPnGK25gbQ6EygSZ/W6/ciBUQtWH7kP3dwHkZGzUqK6WW+oRRK7Rhmte5cIi/gV4J+j6G7SdbVWOqPn3oWE1kUacx1mJXDtuS0ClxdMNFkE0SL5n6A14MdoP9hJtqm7Vq+7D0WMPiB5dCMN1iKR/xPkzBmkMSLXUpTPo3X3VeYL3HGasx/Hr49RtB8cQAbnQyjqsRcZwFlHvduQsbkPpTKeRWMeh2ieGMrq2oj3U2gmvej+/wlycGwlG5Fr6crPI5H7AfkRuUZWDmtzLlnDuDxg+6A5ZifROTqK7v/X0Dr9CCox2orWyrQdMxZlXo+uw8dQhldWfRzK6LOa4D2IPutGtD9ZrbLVT9ueHj9Glwh2kvVKSESSG6wNiTZLPUo6Tsg6vr2Ebro8d1muZBod7Itog6ol/cyE7nbkMX0fXdRO8RhBN+fdhOZCtW6YndHfPISu+/OszKhu2mTVtKSRTKGMgWejn9uQE3F9Bq/VhjaWnWgdP0pjyimyjOZC/oRunDl0jqeQg/ckEjZ/howaG8GQ1TgVW3vuQNfVGRob1R1Aa+aTSOSvJfuImtXvX0LX+DuoYd176PNPkB+H8yxh+sRpdI28jRwDj6Bo62qybVpjZQ13ouaittc1I/rfiDW9mftFD3L2/BBlee0g/XNrkdyDKJL7FPkVuVlGdGfIfyDNnF4mAk+gqPsTKMJ7KxKEaTtCSkhgbkeOtRfIvmGhCd4p5Eg7zPxmn5XpylbCYKJ3hmVcv0mF7m50YDaR7AKdQel6b6Jo7jnya5xUYwpdjKfQJlqLcVJCx3cToa3239E6HU9XErNoEXoa1VStiR61YI2pbkLGxAHy5WnMmqw21yIL3DhzaD18jtCl8Q4kEtKkhNJpN0aPRvYNyLJON0/G22LMIGPiLbQHHEMG753IsZGV2O1FEcKHkYg6ksHrVKOL0P35QbIx2CqZRgb9CSRsX0Rp+l/R+LFaSZgl1FKeQWl+h5CD4HZkQ2TZsK4DnZ8HkI1zAaVWN0MoZCV0m+0UtWZs30W1kfUEi5YiXpP7PBK5eUpXriQroQvNT1lOgonAU8hBdwYJwSfR2rmV9B0ibciGvRk51o7TmFFT5oRoaF+VWjeeEkp3+zryMiatzY3P7zpE8RozzaLF/xi6GDZQW1OiEkoNug612j6GItrN8JY6y2MCGU/PofPZR+2NqbpQFO1eZGza7NdWJ8vNNY8bd73MoWviJRTl24DESdrRr3bkdMk6SlRJVlHdLKMCWTGOUpkvoz3lb9C6sJ5s0tQ6CHO99yERmHU9cwldw18jRK6yHqVlte+foR4gr0ffX6L5DZZqZRaJks+RnXEWNfB6CO051sAlC2yPehjVwp2JXr+R62wrrelxutD5ewIJ3f1kMzt6jDBC6NdI5Da7u/JiZOl8KNKeYFgTrQPIGWrOORs7leaebVFda0z1PM2dqZwptQrddnSjPkTyGjKL5r6FUiiGyO+NtxgjKPXsItq4axW6HSh97E40kskK0S+gzdkpBmW0ELyIFgZrTFXLgmoOjxuRgPmUlSF0IdtIXhHXkYWYQevLK0j4bEGiNE3akYNmgMYJ3azPVV4azCRhGgnO30Y/t6FznkVqrzUd2oFqwN9C4i9LVqGU5UeRuM6yFtlGTpxE0dvfonpcy74qSlQnzjShaedZguC9Ea0JWUTGrXnm9UiQHUc2jzvll0cHEhPfRGV/t5JNnbrNFX0ZZZ7lMV05TtZCtNkR/OUwjUoZXibM5b4fZXakee9bVPdGpE+OUrwgZE3UetB60UZsHROTMIpScd5AN2JRvKuV2EJyGKVw91GbUWKekx0oRaAX1Uu9ijzO58l/LYEjZpC37ZfoGhig9vSjLrTh3YnqtYdp/fOe1002r4yjaMoBZHCmLXRtwH0PjRtjAtlcB2bI5KHBTD2UkbPzecIs5TtI3wi2Wt2NKJK0mWyFbieqA38YrXVryM4RMUfoLPsqYWThOfI/cmwpZpFT/R0UHBhFDaP2o+sli/vX5mzehaLxX6Faz0Ydy0Y0L2yk+OlADstvoiZ0WdWpx0XuU8jhM0QxnDxe2lQdy/J6BWXkWOOmNaR3/cSjuvejIM7ZlJ47V9SyWLYjkfYQ8igkOchTyLP6MsG4LypTaEN9B9VbrqF20R8Xu33R39+IUkteRYvUFbSZ5bWOyBGjqFnHQ6iYfwe1LarWgfl2VBdxjOwjK3kgS5FT9M2sklmUMvgmqpfbSfqip4PQ0r8RZN1YpsjXQRkZFs8hY8M6tKc9S9G6616H9p5DZOdkW42c4lZbVmt5R1Ks6c5nKIr7BzRG6TLFMPBrwaLVn6NIyyTwp2hExxqycSB0o/P2IOoUfobG9hVpxHrRCNpR1tc3gD8hZIGlve5abecryAH/NsW4B7Jeu4u6J8SZJWR2bGX+/Nm0MLt0D7o+V6zQ7UKb7y1os0zCVYI4PE2xvazmYXkbHYtN6NjUupGb2N2ELqyNSPDchETPQRTNOUdIGSpq9LuVKaPN/znk+d5A7VHdXrSg3IVmlw1T7HtiKcoVX9N+7laMGI+htcCujTSFggndToqX7rsQRRa6oH3lMPAb5DhbR/ojR6wR2VbkaHsZRZPTxhpQPYhshizqEA2rZX0GHbsv0L7ZimvCFLpGniZkZdyM7Ii0j28Jnbf9yJl7EAneRtki1m01Cxq1VrQhu+CbqMN6Vs3YrGv/62iE0Js0bmxcGmR9LlphLZhG9+AfgRsIuiMtB7iVLGxB9/wXtKBNutSN14bC5g+RvPPXDIpYvY9OVCvUo06iz/IKoTlE0o6ZVifXhTaqbYTW4geRB+cooSHEKKrntOHXTvOZQo2pPkKR+VojMJYadkf0d8codpZDLfg1mwxLWRxBm1zaQjfe0r9RZOm1L7LINSYJje5sfuIq0o3mt6Mo4PVoDcpC6K5F0dy70L6YVdaAjeR5BjXd+QLtk6281kwju+AZtN90E5wJaWMNzO5C1+VxGpN9FB8rkjaNyv5oQ/fXI8CfoxT+TaTfE2EaBV7eQpHc19A9XSSRm+X5aCVH+DiyNT9CgbFB0mvuFy9tuQuNpBpK6blzw1IbUQeaGXoPEmVJGCPMhTpPa6TkltGC/yYSqOvRppM00g06th0oyrcOefNvRinMJ1Bd83GUlnISid4r6KJ34dtcyoQuufej89dNbYt2P6rvvRUtXFdpjXujknLFI6vXaDUsXfEq2XlWGykQG2HQtAIjKDJzN8r66CP9tPVeFGnahvaXNI9dF3rfd6P1rSfF544ziaJYLwK/Y2WIXGOaIPBXo2O8h/RHD5XQPnU9sv0+JTjesibLPSPrda+EnD2PoJnEX0M18WmLXBtV9iG6Fl5B9khRRK6RtchtlTXBsgjfQXXe1oE5zajuKrRur2aFCV378HcisduX4Hmt0/K7aENthWiuMYMicc8ho2GAMHKpHtqQSOqOnmsjqvu8BV1w5wli91z0OBH9fBkd20laUyzlmSmUrXAEGQRd1LZwd6Dr5ja0sJxETqFWJOvNphWieZWUkUE5RfqGi0VMskwPrEap4mtaZJka32jKaC14C60N1mEzrWNm6ctr0X7eRbodNlejTBVLqc0iNd7sirdRXe7nrByRa0whm+oZ5GDtRfZC2mKqMqrbiJ4SRRYnljHxCPAXaBRnFiLXRlB9hlL2rYFQ0UQuZOsAb/QelzUT6Jx/gcpPVpFetpeVVa4heUCzECwmdNuR59fqC5JsXONIjBW90/JCTKLOqL9FoqUHHavlphO0Rw+LEm9GInocRXhGUFrjcSSwTiIRfAaJXktznqT1O/rmgdPoGr8ZpZPUKnb70Hm9Ho0BaFWhC8U2XppFGRk0WQjdGbQ2tFwdTgtwleA8u43aO/vXgtVnD6ImZz2kJ3Tb0V51KxLRWczMLaP97QuUrvwu+Z4RmiVjqPHWb1Gq+ypkh6SN7VN3oehhI3pKZLVf2HNm4Ry1xlOWrvx1FHVLuxHbHLoHDqB74FlkZ3uQ41pa0e4w2/8i0htpC91BtI63UUzHyYIsJnS70Ka1n2SpufFuiK3cIGIMpZ7+BqX5PIAWt1pTWJfCRK9FetcTjNRb0CZ/EYnc0yjyex5FfC3VeYTQrXGaFrt4c8AoShv6JjIeO6l91NBmQnOBS7Sm8GjF+75RzJL+8ZtD68AEjb3esrwOWu0as/X7CorYpRkZtZTUjUjoXknpefsIM8KziOaWCRMcXkCTCi7Qeue+Vsoo2+tNVPK0BdkIaaeLW1T3DmRzHCfbtMYiipM2lCXxdeAvCenKaYtcK2k5ghwcz6A09lYLIqVBPHOplbAxUmeQLkuSZbsUnSiiewOyaVspC3dBoVtC6v5mknumplEqxTtow25FAx50E11ETQA60UZ8H0oj6iV9A6VS+K6LXusmZLiOIWF7Hm1IX6HzYGL4HNqk4hHfVlsImsFRQjpJrTMw25FBuB8Zh8dp3aZUPicvOVl9tjnCWuEZH/nEjNkLJHOe1UIJiaHVpGskrUb70C7SH40EWkOsJvEPKIq10veuWeTg/j2y0zYi0Zu2k8EmBdyKoujDZHvsi5ZyOogaTllNrtVPpkkZrdlfofP9FEpfT7P0oJVoxdRlkLayDM4J9PnSWGtNXwwSBPSKELo29P1+pPKTMIxSKz6Ovm+1iy3OLLrw/oBE5GngUbThryHdNuBxrFNaJ7oo7caeRumwt6No4wiKFp5Ci+QpZECdit7rJYLw9fSX+hhC6csPozSydmpbfHpQje6NqAaqVTMfoJiNRZqF1VL2kL7ROo3W5CEaGwnIqjlIK94vkyhScwY5I9JsNGTpy72kJ3Q70Lp3PdmMUCkTIhm/R42R3MAX0+i4vIAirutIvwtzJ4rq3oa6PJ8gu1KbojURGkSlfX+B0pazErkTKGj0CvAzlC05kfLrtBJFu45qxYJrF9GaOEt66605QbeQXSPBprHQQepDHtobUapTrcygyOGbKEq1EtIqZpFR8hLadE6jNObbUR59P+mnsVQSj/hafe9GdPwtVXEUGbmX0Lk5jBpMWKOruOidxD3mtTKD6uoOEUYN1SJQOlFdzw3IUDTDtpXIqi6qUaMimkEbun9Xk67RVEb39Vm0RreCWGjFa2AOnZ+L6Byl5bU34mI3DVahaN9etNel6ZyxWvWLyHH+CorstpoBuxxGkKP0PULT0LSd633o/O5HmXqt3FOiFizj8SHgnwLfIDuRO4lsgzeQyH0fZX0UnayFqAV/Wo2ryFE9Rvqfrx1plbSv46ZTTei2E2YqJZ3/ZZ7X9ynW4OrlMoc2nE/Q5/4EefjuR5uD1Wyk2UVzMSqFr9X4WjdXq/G9hM7XYSR2LyKhfhwZFDbGqFXTz9PiHErrupswV3mp89yGDMNdKOXsM1pP6DrJaUdG01rSdZDNoU3SnHGt4IRsVcEzHj2mSfczmmPASmDSYC1Knc2qCZXZFC+jaGKrnvN6mUV790to/9lIumnpMN8pux2VR2VhEzQiErfc57ZpJA8Df436c2wnm5rcKXSs30Ui91UUsGg1sih1aMWILmhPGELXQdpZmLY/ZJGF2lSqCd1OFGGywcS1Yk2oDqKFtxUMqaTYpnMR1VB8irx+DxLm7sYFb6NnWcbHGG1A6en7UAR6GInfY6jm9AiK/pxB6c5jhKZWnuY8nzHkcf0OSuFL0pRqCzoPA7RenW6W13crbmQldB3cjNIQ0/SsThNmdDd6rnmrnaesmSO7MVBtaP/pQ9fbcp6/Exn5O1EGQpppyxbNvYCiuS8h+8K5lnHUvf994Dp0PtJce9vQurQvenxJeo3M4hQhCmci95+jSO5WsonkTiM78iPg5yg9vZXsg1bMxmkEM2gdzCKia+djxQjdTdEjSa62bUrvoQjXShZDE6hJ0QW0UH2JHAe3IcNgLYrmdSDvei0RwDQxr35v9FhLiPbuBe5F0d4z6HMcQ1GgM9HjEtpcPdobOIEi43dTe52Uecp3R1/PsrLvmyS0otDtQJGxh9H6m+aaMIXW5VM0NvWt1c5RI8iymUo8ortcoduHHHvb0T6StoE0hrKLXkX7j1Mdc7D/ATXE3Ez60fVuJOp2owBIFkIX8r2uD6AsvUqRm+Y6bQ6eyygz8JdossflFF8jT2Rl9+b5OloO1otnhuz2h5YXuhZR2IOij0k8VWPIiDqA13DA/DFLp5Hhei+qZ7oO1XNuQB7CZopeCA1wupgf7d0fvWdranUS1aIeQgaIRa/jorcVF5dauIxSjGxQfC3nsR2d/+tQB+2DtN69k9W13Gre4BJqYHcvuu+SZNMshaUtH0aOq1ZrZNJqa44ZM1lGt9KIqPSgdWs96TehmkMpegdQbW6rrYtpM4UysY4h2yJtoWszmLej9OiTZHN9ZpV2utxrvRe4B/gbtMdnJXLnUOT2C+Bp4BcoYNKKeFQ3OeYIyWL8YKt2q75mc+pAKbZ3kDxteRgZUlnVbxQVG41wGYnDF9FmcT+Knm5DxsImJDL70SYVF72NXgyso3M/SqG0+b23otqA82hDPUCo77W25yOszEjvJEqxO4mESq01O+Ypt6hIKxl0WYrcVtokzcF4N/B9wliZtLBar89RveNKLCspEiZ0szBm4q+x3OfuRSJ3kPSbUE0T0pbPpvjcrYpFdd9G/VXWkW5kxnpKbI8en1O8varePcNE7j8le5E7ijIAnwF+Smtf+1kGdVo1omtk8fnmkN3eclmFlUK3Cy1ie0jWlXEG3ZBvofS4vNdZNItxgig8gFKGNxM6XG9CwncrIb25hzCmyB7Q2PreSuG7BTWmeIjQwOogqkk+hETwWbQRZh2ZyBOWGjpM7R1IO5GHfAeK6F2itRborMVuK9CL6nK/j5yMq1N87jJhNutBWqvOy6kPE9LLWWdK6DpdQ/pz48vI4D+OmlC1QpfZRnAVRb9/gJxlSSZm1EIXsk12kI1TNmtxUs9+0Y0ckP8c+BY6rpb2nyY2K/cwSlX+z8iWaCVboJG0qtC11OIsnASztGiQqlLo9iIRs4FkqS8m4L5AG1QrXmBpUkZpWUPIAP0QpbCuQWms1sVyIxK/65FR0Y/OUTc6d/GIb6MM/3hTqzXR+9uDDPRHUcToE0Jt8nH0OVdCavMQOpcPonOXJH15FzIiTtAao1+ypFUErkVy9wM/Bp5ARmSaaaDTKJr7Plqfm5G2nHVTslbDxjzUOpM7CRY5Wu5a3I3WLJsokOb7tPTNk7R2RCttZgmZVvsJDcfSohM5urche+Riis9tZCVQ6rGROlBflb8BHkcityfhc9TCHBK5R4DngL9H53GlBAiyohWPXzu6r3tIN2OjjPaESVpc6LYjgXszWsSSeGhtbEWju3m2CmPR4yzy6L1NmKW5GS2wlt68HomojShlrJcQ9a0mfrPGor2rCLW9tyOj/QBqTmYpkyeQAdOqUd4xlNXwXRTxriV9uYTOoUXyu3ChuxTliq9FpANlbdyFIrnfQtdMmsPa49HcD1FpQbPW51ZxTjSCTnQdpJ0eCcGgWW76+gC6XreQ7mgVG6tyFl2zQyk+d6tjZVLWlGoL6UbazSm7Fe31x0nXKM4yCpd0dEo7ur7/BbJldpGNyC2jQNFh4LfA30bft6J9FKcR+0GR7YOF6CQEvdJuGjVDGGvXUsSFrqUt30TtXWMhNF36FHn4XOguj2mUvnop+tk6ZPagi3stEkV7UfR3AxK/9nUQeXws6htf3LNeXOwmHESifDeqRT6DRO+bKNJ7GH2+VhS8Z1HK0Qi6j2pNX96EDIheWmuMRlabTRE3MRvt0ovW2geBHyKxu530Re4kuvfeQutzM1JAW8Ep0WhsrU87UmqNTKaQQbMc+tA1u450xZTVKX6FRrYVrQ602cygzA1rEtlPetdQG8Epux3VT6cpdLO2T+JBgMXWozbUTfxfAN9GdlZvBu+vjDJsjgPPI5H7OSvL0Z2X6H1R6CII3bTWXcvyGUd9EVru+uuo+H4dihQmHSt0Ee+2nBWzhIjvJWQAfAq8hoTUADpvO1Gd73XoHG6Jvg4Qor6WCpf1ImBdnNeh9OZt0Xt7AKUzv46i1odRXWsrpUtcQRG0ISRea1mMrKPl5ujruazeXBPI4jqzzTFeq1L5Olk2uaikVPF9qeL7NuR4MifVfjSe4g7kDEqaQVML8REVb6BsimY5IbMwZuICutVE9NrokUUt4DRyeCzX6dFPdvW542iva9Vus1kyh/aPS2hfTbtO15yy25BNkUWdbtoksXlKyHb6b4A/QYI3K5E7iZziL6Ga3C9oQZHRJFpV6PYRMjnTjOhOExoKL9cJmjsqI7omTJLUiE2gDekSHs1tFDNIUMVn2b2NRO1qtBHtRQb1dUhAbUUR31XoJrFob9wozwLzAvcS6ntuAb6DBPvrhCZWYxT/GppEn+Us2iRrST9sR+fOUsKO0jrCPwva0IJvUfBuQvSrMnW/vMD3RF/nKp7XvJvVjKPZ2HNAMMpK6BxabWVH9OiM3ttqlPq2HwnbrdFjkHTTPo05tDZ8CjwFvEPzGvpkWXPXivSgtXsT2XR2nULr7HKMmRKhb0TaYtzS7U+Q3azWVsaihKfReV6X8vN3IBtiHelmoEC2Mzwr1+yFfmcDajz1I7IVuVPoHP0R+EcUHW+1sW9L0YpOyizpQCWMm0nXwWjX4xWUXdCyqcvWFGU7MiCTMIbS467gF20zmSaMMTqKamNXocjAViQubyHU+m5HhrY1ljLhC9kZkVaXOBi9p/2oXf8B4FfIID+MRF5RU5qnUXfb82jxqKV7eQkZDZuiRxetIXSzuo460fX7Q3RNjxFEqY3lihtMcQFcrUFY/LqPRworxa6NZGmL/Y79vpUY2L3UHn3tRmuqjQ6zTupZGXQ26s1GVPwRZdw0e21u9usXhX5UG2hCN03mkCNuhOVF4qwp0SDz9400sEZUR2jByEKDGEP1zd9F+32atKH1axXpR4vt+bPYNyqdn9X+fxD4V6j51A1kK3LPAK+iSO5bKF1/pZF1h+2s9thm0YMav+4i/WtzEmUhnk/xOXNDR+yr1VQm8dJZfe5n6CAVVZy0IvFa30MofdFGA+1BDaNuJAjhrYQ053i3zyw2HWtqMUBIg9qLUneeQh1iv6S4w6st2j6OPmet6ctr0TlK28BtBllG3NrQsboTGSSz0b/FRe5iqcxL1WixxP9Xiw60VXm0xx6NwETuQeDXwLOk3zCmHrI2aFopursWOXFWk76hZvWv51he9KgTvc9aexDUis3PHUEGl9sT9TGDnMfD6BimeY7akCPWhG4ta2kSsiytWiiCWEJlXn+JRO4+shG5oL3qAhK3v0IlaEMZvE7eybJ3Q9zp3Ur0I1t5I8mm4iyFlYucpUXHD5rQNQ+tpQHWio3JOYjSjYooSlYK0+hCtm6Wz6EbZj1a2L+OPEVbokd8hm9WotdqeTchw24ncCtKufx7FJU+QfGuqzF0X4xRu6FhQncTOu5FT9uLn7MsDIYOdM2szuC5i4g5HY+i7p2/QI6jqSa+p6xpJYEL2nvvQg7ntEWkRXMvoNElyxG6PcjgGiQbMT6B2xPLYQ6tA0PonKcp2mzPtqaTHaSb6pjlOZ/jWudJCdm+/y3w12QrckGC4lO0Pv+e0HR0JWLnIu1z3orNqKx2fDe6XtMcQThHKBdpyZnl8YjuOpLPz51Ci+kFil9buZIoI8/+KEoRexsNKb8edd2+M/reuiuuZn5db9pY6q417NmMaos/Af4t8k4XqTHJBGGUUq1RCatlXkPy8oG84h13G4c1kvgN8A+oe2deRG6rGR1ZsRN1eb2Z9Osfrfb1FMrwWc610Y1shbRHXMRnOXracv2Y0+sUQeimiTkZ15Cu0M1ynShzbUTXas3/JYrk3kj6x6qSGZS2/CWtNV0hKdXOR5q02p6zHo25uhvdd2nW51ra8le06LprQtcaUa0i2cY1jQ7MVTzNqOhcQvW976KUmm3Iu/kQqqXdgTxK1vEtixmPpeh5NyMjagsS2i+jhg3HKEah/DjayM6jja0W55HVPlnH1eO0hvPIRW62zCCR+wXwc+CXSPDmReT6+a+NtcBfoXFsG0k/7W6O0Ln/4jKfy5qsZbEHzBFGIDn1M0V2jWWs9GgN6UaWIFuhG4/omsj9nxBqctN2LlWjGzmyHkHn53QDXjOPZCl0F6vFLiLthADUdtJNW7Zo7klatOMyhAidDQHvI9lCM42igmO0zkW1krGF5woyng8ALxLmK98Xfd2DUmwHCWOL0sIWKesOvQoJ3x7gZ0hA5v1mLKMN7CLBo76UA8lG0KxDRm8HxRe6viZkxwRady+iCN1/QTW5p2l+TW4lWdXa19JJtQh0otKRr6G1NYv5uTPIkfkVy298Y43W4r0c0sDEiAvd5TOLHK1Z7CHtyBFda/+JZlMpckG2yz9BziWL5DYiAtiNsuUeJdhXKzWym2UPllaK6G4BvgXci5wzafdFuIIc5WdSfN5cYZvVIEpFStoEZxoZWi50W48y2iSHCF1c/4jqeG9HovdmlG63jhDlTSuVrYSuzw2E8US9KLJ7GN2ceTPo41xFBuU0td8bViu/Hhm7PlPPqcYIEiwn0ViKp1H5wTC+DheRvWhm5x1IPGQRJbUylfdZ/uzTeOO3tJlD63oRMnfyTlbTC6xO18qZ0nzerLrlmgNlGtm7P0YdlhspckGfbQBF536A1vFPyLctkwW+T9XGRuAvkNDdSbrj3Cxt+Ty6BluyERUEoduPai6SCt1J1NxoHL9wW5k55GE/i2plP0WeyH3Ag0j47kKidBBdT2mktdnGN4CMwT9DkeTfIYPtGPmN7lqd2RS13xvt6D5cT2t0XgZfF9LGUlA/QGnK76MUuOWKlyJT5GtsD/A/RRHdzaQfJQVdMyPIMZJG+Ue8s3gW73WGlWf4Z0FWe4h1tO0i/Yhult2WZ1HW4pPAf4Mc9QMZveZCWHnWJpS+fA457Y828D3khaynahR5XwA5PZ8AHke2tnU5TwvbF75CgayWFrpWG9hHcu/cBDKy8io2nHSxFDiLDpwCPkKephvQzXgjEqVbkWhLI7W5hDxZu9HGtBvV7b6IPFF5bIY2wfzOy7XQRpi52ipCt1XSh/KC3Qt7gAfQvTiJRMxKTPfMsqFJ1qwH/jUyZK4jm5pX0Jp9Ea3Zl1N4vizHv8yi6zlv63nRyHqOaFaj07K4tiw7bBdylt+BnPNZjPCq9f2YPfM4Kjf5B5ZfO18ksl63q3XYLhKrkPPz++h6XUv6Kcu2L3yE0paLuo8uSVzo2vzUJEyQXR2Ik29mkYibQJHej9HGsQ3dmPcBtxBSm/tY3o0ab1TVF329BfgD8CZyuFwiP5GASbR4JO283IU+nwtdpxpWv34Dyp7Yj2p0/4jGhll0t2U3rRhWe1fEz7oFGd3fQc7BNFPS4swSHJPvsfz63EZQxPOZN7IUuvbcZj+m/dxZCN0+5IjfimyHdTRH5BptKEK3H3VaP4ac9y053qUKWQpdS1MvqtBdA3wD7Q8Po2s2bSeolbOcQPtCGg7Q3GILVS9aCJLc+DaXz9OWVzaW1jyFbpyzyPB+C83EvQ/Vo1hq83JEnKVMWefnDcgr+gAy9N9CG8YwzXe+TKFIc5L3YkK3l3Q76zmtRRvy8FoTwVvQ/NXPUUOqt8lXhLdRnVSLQDdy/v1LlJZ2K9nWCM4gB+ARtDamsS5maaSagHIH2fKwKGaWz2/p62k/ZxZCtwelC29E92AemmjZOn4PSmE+hwIGeVm3syardaTIDtDNwI+A76LrwrosZ1EichE5xw+wvLnqucdqdPuQdynJzV8mCBzHgSB6L6K6kyPAO8iYuxsJ3r0omtGPRF09mCBcj9J8txLqhV9CI5JOovqDZkV4Z9ExSJK6bFHrPhoz6sApJubwaUPXyypkwF0XPW4HnkNG0xWa6/TJwnA1imTQdKBzdA+K4j6Buq9mKXLNa/8lynpJy2uf1WgQi+a50F0+ZbK79+JNo9J8fmtylnZU195rT+znPGD7/RYUwbuC7tfD5CczLSsasWY3M2KflAEUDPoRajx1CwrkZCVyLcvnBeRgKcIeWjcmdHvRIlDrhWEGxhTeHdG5ljK6Lq4QZnS9h1L07kKzIq2OdxXJrj3DjP1eJHpXo0jJftQs621Ue3AMRVYbfZ3OIqE9STDGl1qwbOMbIGRYFClaVUmR6yeLQNzgbEeb4ipUPrAbRXdfQ80mmtXBu1TxSIv4dZUXw7USq8Vbi9alJ5DQvRWtfWmPEapkGpUWfYrWwrQallkUPSuxWyQDNa+YEyyrWmojbaGb5uSGOHleI3rQev0E6uvxt8hmanWytA0sMyTPtKH9eh1ydDwQPa5HWYtZOPzMNj+HGll+xgpIlzehW8/w91nkdZrFjVlnYWaQ4LuK0po/Rcb3LWgu2K1okbfGVfUI3g4UIe5B9Q27UEfFL5GB9zpK60xSL7tcyigdZIpkETUT71bTXGSh6zQOuw9Wo+tnDUrV2wX8HjVtG6Hxa3WWo2ggdH/NC/FSoHVo7vjdaJ27HQlcG8uSpfE9g4zmz9H6d4b01pIsU8bjHZ2d+mknO9EI2Tg6VmraujWh3IuaU50BnkIlB05yLAiSxz4n3Wj9t6aqd6HGkncggbuZYAdncR/MEcaFvowcoS1vY8YjukkvCjs4rZ5i4aTDHBK7Y+jm+gJFeW9D3qzbgB0sL8LbgTaMbkKE1zpB/xFFeb+icV3CJ5H3LElE1+5HKyUoesaER3UbSxu6/rcQ6ti3IcPpdSR+Gn0+smow044E/V60eU/HXqfaa9YaWS5X+b68xAPCHOzro8cO1DRsD3I6pDV2bSnm0Bp3HPUteJ90vfbxiG7aWFQvT86LImIOoCy7Y0+TbllEVjW6RaANRfFuQfWZZ5HNUoTmcfVg63cW57oD7Qt7kA1oDh97rcqMFFvH4u+l0m6p9jvVsN8zfRQvx1gVvR8r39uN9ocNhIatWe4PNjf3JPAqKm1aEWMJO5BRNEDyML9dCC3vDXBSxSKdk8joPoqK4W9HNba3ocVgFfXVJ1g0YBXzI1vXoejK6+gGP0f2BfjTyBGU5B6xdMd6xH7ecIHbHGxjXYNSZgcI99Kr6L5rFFnVc0IY0fGnKOUrXpdYTdTGI4VxQbyQsLWv1R6VhpLdt+uQY2ETSlkepL6MqeUwjQzlt1FX+pOku0/brNss0pYtEtOX8nOvNMzpm1VUaAo5U9J0xK5EgWtY2dJG1MBzCAUEWrU5VbyZWdrnvROtwd9HGU0dzO8QHu/KHF/Hbb9YTNsstG8Y8b+dI2QpdKF9eFP03tagYIw5P7OoTa98X9aY8GNUm7siormgE9BJfV3ozKvnqctOPZh36RyKxhxE6ZV3AF9DBvp2QoS3HsHbha5xS2nejqK776JIhzXrySpqOkfy6IcZe1mMbmgGWa0N1dafpNfIQr8f38Qqv4+/fuX7qfac8Y2x0SmZ1mhwJ8qasNd/hcaK3SwwMb8NeBStJfEoQTWhu5ghUSlwF/q+2u/F3083WnfMeGl0d9dp1AzwA+B5VIOVtqE8Q3DipXl/x7NyVqX4vCsNE02DZHP9mbP6Kuln9K3kDCBzlu1AM1SHosfh5r2lTMmqUVobcjg+iEpGKstn4kK00lkZ3++TXIeVe0U8omv7QCeyRc0uzbKkpxJrQHUYZQocpHGZjU3HjOl6vM3WdXlFeASczCijG24CLeoHUR2vjSayCG88KpUES+W01JANKLp7I/AGinocIpsUoXq6wsYXxqJ7uLMyWmw26HnkqJhhftpbtahd/OtS/wcLC1n7fiFBZM9jotI2NEsl7CFE+Bpxnq3bqBlPFo17FTmYsiZLL7V9ts3oc9WSmmz/V8t1mfS+zfKz1kJ8Zu7rqON9Fud4GomcKdK/v7uQU3Irune8NKo+7Dhm0ZDHUuPTFroLRclWElbjvwd13z0D/D3a61qNLBvaWUbTIPPX5oX27SQk2TsqS2masUfMIfv6JArwvIY68K+Ye82Ebj3GloXCV8zBcjLFBO9JdBN+hhpJ3QU8Qojw1lPnFm9YZan6ls58PSrKfyd63WYbVvZeuyh+RDcroTuFaq1fZn6dSTwVaqEHi/xb/H1Xw7pgV0bpK4VuPAWzK/Z1EIkyS2sdIKTYd5NdFN8E4TY0gH4YOZXeJ/v0/aw39rhjKOnftRJmzJxCxszL0fdZOKKn0PWTZHRaLdi5XIvq2Hpo3RrFLCkhp249JWm1MIcyKMZJ9/yv1PrcSkrIVrke+B4qQ/gtjXFMNoqsI/e271Tupyvt+iojx+QFZE+/gAI7zbZzG4otgs32RDtOnLHoMYS6w32OorsmeDcgcZB0Y4yndPWiGolNSEBvRlGuE6Sb0lHPfWWpy35PVmcKGfFvAC8iA6CWqG2t30P1VGT7fiGRG/+9uACzyG4/unY3oBnQm1C2wk50DW6g/lT9pbBIwXbgIeQoOB19zTIrx6/h7LEykFPIW/8bsklZNqZQyckI6c9otv4Ku5BYc6GbnBJhj8ui8+wMiuZeJd3zXwTH7iy6/qfQ2t5D+s4EE2hrUEbbD5HYfYPmjYnLAg+SZYsFIy8jG/ol0m9MWAjMmC7CAuOsPOKC9yASvd9E4zp2o828nqZNls4cr9/dhDrhvYSMxKHlvXWgvlSVPKRA5h2rEbuEUrqKsnBb+nIXIbNgIxK616OGbLdGPw+SfufZNiS29yCx+yU6hiMpv04cv46zxe6FU8hR9yuUnZLlOZ1ChvcQIasrrfNcQgJtPXJEnkvpeVcSHYSRfWmXRliEaJhQNpIWed/3ptE1/xXKPFuFHO8byUbsdiIH/APoPjiFaixboVywnjpYp3ZM5A6hff4llOVzlhV4zF3oOkVgFKVbXEAjMz5G0d1bCPW79YxSsGY9W5EAWIfE7h+Q9/QCy/NY17NxZ9mltlWIr1lFOk6zKFvAMgbOIaPpU+RseQM5cb6OmrJtI/3OqVa7ZHOsv0COglYwnlYacZH7GhK5byDnRZb3xSxqeHWZUKebptC1uuvdyLBPO2rcypSQALsLpYCnLRyt4/IQOv9ZNHLMo9idQk7VT5B98BnKjvkz4H5CLWiaWHOqXcAT6D7/O1S3W3TyeI5bBXNGXUb7+++BZ5ENXfRxlXVhQrfeiy7vHjindZhDN+47yPg5iDaYh9Gs3M1IFCT1YMebVdn83c1IDLyINpV604XiqatJ3lMrCd2s1od486kiMx09RpDwPYg2pMdRp+QbkOGaZvfULuTQuQOlMp0h2whgq1zLecIaAp1GIveXKKJ7gcY4LYbRNZN2Q6oSuj63onKVN1Hk0KmNEtq79iHxlTZ23V1B538l3NuWqv8u8AyqczyFnOz7UD25zUBNG6u33gv8CXIwPc0Kaybk1Iw5oi4hB/ofgN9F32fdjyO3dFD/LFyrQfNosNNIJpFhPkwQvA+h6NRe5MXuJvl12Y4iw7uRsFiNDIUXkfCoZ7B25fw2Jx1aNb3b6ixHoq8ngR+hOq21pCfsLYX5BpQu/SHKmsjScHKjLD2sRvIU6q78FBK752hc9HMcXTOTyH5I0xHThq73W1AKswvd2jGhux5FxtNmGp2P89S3JxaN+LiuXwLPoXV5htDFdj+yFdaQ3TixQeAmNDP8HBoR4/XrTpxZdE+eR3v671HK8gFW0CihanQQhicnNUTiozNazeBsJcyoNfFn80cnKU5tYzXGCOnMh9DN/E0kCrYTOk7WE93djKLFq1Ck91mUspS062E9QtdqK6ZpHXGQxfpQrdFUK1BG19kHaMMaAf4CuId0UxG7UF36zchBdIbsN8NWuZ6bhY30GyKMEHoWjUi7SGNTfCdRVGk0et00o1ltKIq1GaVttuoc0SzoQNME1lNfOc9STKN16TStHyGyRj6fIGfSc6hZpd1nI+ge3Iiu19uRgzxtx7adw3UoJf3HSOx+RuufA2dpLFV5BJVCvYtmqL+BrtdWamBWFx3opq1H6Nr4jCzSNZzl04U2u63IWFiLztUcEokXgaOEpiJFpIze+wfohj6G0j0fRClFlo6cZOOx8T7rkGjuQ5vXOrSpJal/6yR5eYDV3Y3jtWlL0WoiN84M2rR+ga6JPuBO5LRK43O3I0fOXhQ5+4jshG4rpeI3i1m0LpxDPQpeITTOG6HxNdYT6Po8h7Jg0o4edqP0+tvQuusG/dJYNPdB1ME97eiijbA6y/JKehYjL2v6DIpcf4FE7m+YL3JBx+MrlBpqI+JuRk72LMRuJ7LnvoEcDcPIhiuqndCKWVmNxu7JiyjY8yZKrX8f2aoraozQQlhEt555uG2E7qF+seaHdiTMbkTpvLcgobsGLZSW3nAJRQXeRp6fkxTXGJ1GG+9LaAM4SYjubkUbT9LMAzuO+6K/34LSh15AUeRaivrrmY2a1YzCZpLFdbUSNskyMip/BexAjqu9pDd6yqK6u6PnPk/rXHOtwhwhXfQU6lHwAjJojtM8ATiOykZOovVqgHQbUrWja/JW5GQ8ldJztzLt6F6+g2zqc2dRFtg5tAdmYUTnoe/CLHIefYHW3qeQoK0mKKdQxsHvCCOdrGY3i+ZUvegc/xCJm5+idbuItlseznVRMd02htbGD5Dz83XUZTnLnhuFwyK6MyQzcGwjslEZTj7oQMLuYdTM5i4k0Cq7Elvq8j0o3WYnWsyPUFzvICiN7hO0CZvYvZ9Qu9tJMtHZRqjbHUDOgn7Upv0Ii8+p7Ij+pptknnVLQ5mg2Ocia1qlGVUtnEGG1N0Eh0uaUd0daA04QvpR3SIaYHnA9uUxtJ59gUTuSyii22xv/SxyKp5CkaXVpJ++PIBGYd2MxJVHJxbGZufeg0p3+jJ4DRtXcgY5XrJwijWzn4X1qxlBYuGXSEgeY/G9eALNKf0dshH60HrandH7HECBjD9FUeaXKF5jsJWwb2eBlbZZVuZxtC+8iKK45/BU5WvoQAdtguSLlkd080U7ErnfBn6C0hzXIw9gZQTIFsQ1yFu+ERm8f4tSYYoc1ZlC3tchZISdQoL3VvQ5e0kmPK0L6GbUBXQVSgt7Cm2GCy0qnUhc95FMkJnQnaLY5yFrVorINT5HTpyb0TWYBm0o5XQbEhRvkU36cpEMsDwwhcTjZbQef4oiuO8iw2acfBzTEfR+LqC9J+0yJotePYbE/dmUn7+VsNm5j6FzkUVTJGsE+RXF7u9RjWoi97+wtMg1riLBYY0se5HNkPZ8XdC6bdlm30NO/U8ppsDJwzpWJKaRmD2E1sQP0L5wkOaUsBQCE7pjJPeWWs1ALyvL4Mwrq1Ftzo9RFHMDC3sU7Xx1E6LyHehG+Q/IU1T0BWgEebguos352wRv9yDJxW4HOqadKKrbB/wcCZBq3Se7kBMhaU2lNZyxbqZOdVZC6nKcSygt6WEUgU1L6Hcgh8wmtJY7zWUWpSJ+hNaWD9A6dhxF0fKU5TFGSF/eh5wmaacvb0IlOHeilO3FsmhWKhbNvZdsxwqNoevwGK3lELMGnSPoev4FcvofJ9kePIy6IQ8iwfsACjZk4XToRPfGQ+g9X0BO/TytD066zKLz/EcUwX0bZWFdYYXOx60VS12eJHR5TbJRudDNBx3A9cC3UHrjRmpPKS+hVJi9aJTJJ+gmagWDYhLVz1xGC8Ql5Ay4ARn39czcXYPShv4JOm7/gIzRSg93Z/S7vdSejmVe5Vaq0c2yPnclid0ySl89hzJw0uqo2k5ouOZCt/lMoijur9E6fBwZ0Hk0YCdRdO8EEglpz3suRc95A/AkimCcpvhO2LTpQiNuvoOcYFk0CJ1G++hx5DjOyrBudOM6E7nDqJnP08jZf7zO93EZpRKvRk7uO6PvsxC7/SgT53tI5D5PsmaZzaQI7zFvTCNh+2vUAfwsXs5RE9aMapIw+D1pwx5PXW4+a1ADijtIJnINMyiuR0LwLVpD6IKu74uorvYK2ggeJTQ56STZ9WuOgT1ogLulQr/H/Ll23ei8JI1ymOPJ7kdnYVbaunOJML8y7fTltQTnTxaiyq/l2rD0yYPIqEk60qzRXEIRvgso4yVNZ0k8qns7WrMvUswUzaxoQ8f9a2jO6uoMXsMmAZxC57pVGt2YyL3CfJF7bBnPOYccAS8S6nWz6sQMih7fimyRS6ixaNYz0dPCu/EnYw7tBxfQNesit0YsojtOSJVMcjN2IoMrizoEp3bWopSl7dQ/5sHqPm5Gi+cwrRFRNEZRLcNQ9JhEjoENyDGQdARRD2o48X1CDbSJ3XL0/6ujr0kiutaIquipKFlvYitN5EKoibxKesfWhO561Oiqi+IMl7drbI7qxyOr6y+eTZB2rXg3imDeT2gOlue1YARFWg+h7v5ppi8b/cgJ+12Uzn0CN5ANixp+B+1HWURz59C+9mX0qFaqkxaNOq9xkfsF8AzLF7nGFHJUPY9sMwsiZBEUspKq+5Aj4jy6R1a6M8j2Bdsb6jnupQW+X+w1jTbSd2x0ouvoAZT1c4B87w25wYTuGGFuZxLRarNae0jX+HJqp4QEVXyMTr30oEVzFfUvDCXyK5An0EY9ga7XMdSZeivJUowh1KjvQKlDJbSRvY2cBAMkF7rxtOVWWcCyFLwrTexOonSlMdI7plZ/brOiO0lf6GaRYm7XlY1Ki9e0lyt+xwwde1T+GxV/V/l9pcHTho7TAHIKppVGTvS8G5Gos+62Z8jvmjqOZvl+gsa5rSLdSQwW1d2Ieizcic53qzVDqodO5Fz4DiqnyaI2FyTczqK98yTFz/YykTuEROGvUU1uGiLXmEAOoDUoI2EQORKzCAr1IBvmEXSOLqIU/zyWOzQCO7+XUCq5BfDia7+t75UOy0onZvxRia359nyz0dc2FMkfJN3z3YkaR34HRXVHkNMvr3tDbogL3VGS3xjdKIo4QHFqA1qNduTVHWD5N5VF6JM2UDLi0Y28pqVMow3tKeT9HEIbxC7qm33Xjhaf76ENZxWq2d2DvLlJvLiWXTFK6wjdrFhpIhdCWmva10YbcvQULTvHZrn+EYnBGYIxY4ZJOfaz1eTHvfzxNavSYKhcw8wAsjV3D+qJsJ10xe4Aiur+ANVkv4DWqTyup6C9/yMUVdpC8nKQpbDSmv1opMp5tMY2a4ZwHighEfUASlveTDapsWV0nx1BEaQrGbxGnKyN9rjI/Qz4DfCfSFfkGqOoDGw9CiD0IGdiFuepH/VZ+S4SuVfJ95oRz8RJ+z2W0XX6DkrlNsdMCZ37udjPHbFHe/Roq/jZvo/btvHXss8yjfagHrR+344cdGnWZ/dFz/0tJHKH0Xl2FsGE7lVkQCXN+bbOb/VGAJ3lYwZcWovnclIuzHCEhdMJ88AswYAcQ+/1cRSdrSf1rg0ZeI+jTe09JH6vI5kBbE0xhmgNoZvX81900jYO2pBDxkZhFYE5tG99ilIODzF//zLBG/+52n1YS0TXfrbfNY/93ciI+ia6/9MyaNqRiLmTENl9n/ymlI8iofsxEqP9pJ9CW0Jr60OoAZbVMRc9ulgvq1B6+5+jdMYs5uaC9qGLSBQeJNtrcLEyhLSe39KVP0ci9+9QGmhWDCFH3DoU4budbOxla5R5L1ovLhGaZOZ1H84qGFJGx/1V4P9GdWEK10Zu447MuF0d///488S/tzTpWXR+H0O2n5WwpXW+26PnuxWthYfQOW4FezEzTOiOoMUs6aZho1bqjQA6y2c5EflKptG1MEzyBaiIXXBH0HzKDrQgfR0J1HpqadqRV/1BVC/diQyzJAbfTPSertA6C1deN9kiYpHXpN3Cl8LS8JPWqjcbW/vOoChfoxlHRs06QrQmrfPShaIBjyBhdx5F1fKajngWCd37kPO7cnZ7GlhfhHjq3kocqdIL3AL8BPWZWJvR61hpwGF0bs+RfcQ1K6FrYsS6Kz+PJiYcyej1jFl0jb6E7udBlGaeRb1uF7JBHkdCdwIJ+jym+Wfp1LDmaedQ6nKj6UJ79HZCaWe9vXMWev7NaG84jPaGM7ittSAmdIfQgUpawN6FFt1+imUgtRJWi3AKLWhrqW8BLUd/fwQ5PVZK3v8I8BrhmH0Nid16xma1o+M/SPImNdaI6hI6/q0idLNiJS7qPWiDy2KkWzy1tyhY1LZZQucS8ApKS1yHzkuakbU+VFLxPZSOOIyMmjxe+yOo2d8dwG50LLJojDSIugt/D9ksf0DGbB6PSRZ0owjuk8gxu4lsRteA9qALKAX0I+ZPFciKrNafMnr/B4HfA/+IREIjrpspVDv7PKqlHUQiKIsykV7kDPoRIcr4OfnLBslK6FrEdobmfWY7379E9uAaZFOmqZH6UKr64yiF+RWyLysoLDZeaJTQtj9JhzKbFbodLcArNY2o2VxAHtdjyIOU1BA2kXUOeR6Xk+5SRINjGC0UM+g4PIiMtXozFeoxPOaQ9/wsOg+tcC/ltU67qPSgDbOeWvJmk8X7jdfbNoN4tOZ6lL6cthNiFcoQeRI1AnqVxgiOpMyh9/chSunegAz6LM776ug1RtE6+SL5H8OUBl1oX/oh8OPo+zQjRXEsmvslcmCcpLjjTMy+OYnKlf4zSsVu5OcZQyVN65HY7Yu+zyJAtAqJoB+g+2Mance8dWLO0j6wbJ9mMYls8l+h9Xs16r2QZgrzWtSgz7ptf8jK7luwIOZRmkJR3atow0piqK9GdTmDsb93GssQWkTfQk6HHSRLHSsjr/gHaN5svQtE3ElSNIEzimppRtDx/C5KMcrKkKjEFuZL0eu3UjpeVulJRbvGloulx3Zn8NxZ1ktlSbOvA/PeP4cijWtI16CxsW93oFS14ygqlUfRMYz2oFtQf4Je0u3AbNhIlYeinyfJrwMgLTrRMf0Bqsu9ETlis2IaRczfRAZ7kY+tRXM/QsLjU5pz/4yg63Q9WiduIztnkKVIfx/dl6OoBCIvdkXWqctzND9YMIayIZ5D9+4e0o3idyPH96OENO0j5HNvaCp20C1F5Rzq6JVk6HsfOomDaGF0Gs8cSsP5DUpb6UWGwFKpiGYkDqG6lX9L/aMs4t3niipAxtF4oMvo2K1HC0kjmEWb0RDulVuKZoubZtFLEA9pG0fWSKNIx9XWm2ZzFY3XeQNFdtPuWWGR/G8joXuR5tQkL0UZpVi/j4T5euovpVmKbhRBf4DQQOaPtGZktwtdVz8C/oIw6z4rykiUfY72w5M0TiBllc46ifb1szRPCJSRnf0CodRhH9n1uBlE18qP0JoxjBzpeSHrpmN5KP+6gs73g2gtXE+657oPrQ0/QOf3Z2gNdmJUCt1TyMhOInS7CUOx28iPx2ilYakxv0ILp7U2X8gotsX/IhK5/xdkKCwnvaVc8Sgis0jsv4Wiuo0SujNoEzpH/lKM8kiRr7F6WY+ie2k3MjHDID56IU2yPlfNvg7mkHHxMmrGtA5FbNJOYb4ejZX4ivxGMIfRe9tDcLomsSeSYA6Ah9A90Y/SyC+QDwdIGvSgTIGfIMFi2XNZMoGusRcJnXsbSZZlGc3uJTODSsyeQfd0D0ozziJLxzox3wH8Gbo3X0JOjDyQZfaf7WfNZgplEvwtSlm3nkZpUUIZRNejjJ8j6DznsQFZ0zChO4M2h6MoqpWki58J3TU0fxFZ6ZxHBfCXgH+C8vdtzrGlkFlKxygSdO+iFuzvs3wPWLnia9HoJsyi+3MaJ3JhflZFs1NuikJRr7N66AGeQHV5aQvdWWTc2qitNGmEwM3DdTCGRMFTSOjeRrqNqcxovQuNrjiHUjDzuFacRz0PdiPHTFaNd0D3wla0v/WjNfu3yODL47FJwioUsf5z1HhqD9rLs8Qcru+jtMtGNz8rWu+BerByh6fRtbs2+poFlpl2HxK4E8DrNF8IZTmhw7KT8pLCO4acRneh4NN1pLsetqM95w50nj9H658TYQd7DoXYT5K8PtNupD0odcujUc1lCHgW+AIZxk8gz/oAOldThLEMzyOj4DTppHnkweCsh05CrcOP0YKxBRkajcCaZYxGj7ws0GlQ1NrPvHEdmqu6kfQdija/+RLZpHtlmaKW5fMn5RIyaPYiUXo96ddk7UBR3VNI7J5K8fnTYhI1+3kVHYMB0h29FMdGY21CDaoG0b3yc+TEbbZBXw9t6Hg9DvxTZCBbNChrITiCbIfnUDp+o8toslwr8pQFNIHukWcJ6ctZReo7kT3zMLLzx8jHXO6sRlJa1+U8ObqGUKfvG5EDNE3HRgk5wrejme4n0H18IcXXKDTxTXgS1S9cRBtFraMBSkgQ3IYE7witkzZUVCbRIOmTwO/QzbUZndNhFMk9is71BPlZ/BtNOxIO96CmDQ+ja3812cyBXIhZZJCdQ8ZyKwldyOY4Fr0ePAmDKHXxJtIdPg86flOEtPm0he5KOD/GNOqV8EsUzVyDBFhaxNPUvo3W9+fJZ13qZVSbZnW63WQXjSwRxGEXYRLEPyCxfY7irKm9aL9+MnrcQpjHnvV+NIbsgqdRGVOrjSvJ234xiqLmv0Xr+l7SzQIxTAjtRNkgV6LX/oLmicGsM33ykrpszCDHxn+P1sNetE6lRQk5S/YhR+g55HTNY3lLw4kL3SkkgI6gAvYkM/B6kad5HartcKHbfGxczTFkELUjY8AM26I1nkmTbmSA3oZSwh5AxsUmQlpoI1OoTGjYrExncfLkmc+ablSD+CiaqZp0PvNSzBIatZwnG8NnJZ0v65Xwn1BDwB7SjdSUCE1mHkbr+yfkL5OqjOyJF1A0qR8J9Ky62NuavQqJ3QEkEG9B0cmPyffa2oHe79eR0/Ve5CyxrrxZ70eWsvwOysw7S3OEQlYRvjw6RmeQvfwMOs8/QvdIFvW6JoT2ICfZeYLN0Wr2eh7PNWiNfgPtDeuBW0m3DKmTELg5hSK7n5KPplxNJS50Z9DGdACFv5Nszj3oAG9Fm+6KP7A5Ig9t1vNAGzJ+tqIU0IeQEXQ9inb3k76IqJVxVLNzgGKm2jUa28jyuJmlSSdKx/yX0de0o7mgYziOUqsu03pGTzMYRhGxvYTxe2kar5aK+HXUhfk0+ey0aRHuP6JMmQFUIpJVvS7o/rCOzH1obb8FRc1eQ8IiT2tsG4rw3IpSlb+BMjfWI8HeiL4n1mX5IHJMfE5z0lqzTC/O614xgSKrP0c2twWNsrhHzEm2H4nqESSyL2bwWrW8l6ww+yBPEV1jDDne7keBwetI71hY5H4HYeTQJbTmrWjiN9McMnYOoY16C7UvshYhux815DiBG0xOPuhGKSK7UaH+3Sh6uwMZEyZwm9VIbQalEh1C2RTNrpvJgizGRWTVITgPlNAmeBfwz5Dxu4lsHDEzaL0/T74EQC3krUbXsDEizyKxuw6l0qZp0FiE9DG03/6efEYsR1AH+/Uowj1AdiOHDBs3tJowam87cm7+EY3MOYWOV7PWkM7o/d2A7u+HUZR+G413uo4iR+svUaOiZqUs232cValL3tYJ4yrKOPgZuk96kIMmq5r21cix8j1k879I49eOrFOXzRGeNyzT5f9LsD83kv7ecAPqz3MGlSJcTun5C0ml12gSbQCnUD5/rW2wS2jzuhGdtFZMh3CKQwcyqDYg4+Ee5CW/ARkSg2gzaabANeJCY4T8bsZ5olWPkaVebkedvx9B8/e2sfRM7HqYQxGF4yiKU6R6nrymIxpThEjNbrQerUnx+duQwXoLEkqnkJM5jynMVi+2BtkUt5L9iBzQMTIn/CoUPbkZRTveRWm6x5ER2CgHYz8ycG9CgYE7ou/jDSMbuSddRQ7W54A/oDKnZs6ZLfLzLwdzCK1C94aNKUsba962ATlSLyCx+w7JG9Hmkfi+kFcNMoMcG/8DOscPEEoU0sCyF29Ge8MhtN7lbW9oGJVCdwp5hz9BG2iSeU+WJrQhel5PX3YaiY2XWIVqGW9DN/qNyMDZSBiz1KwU5WrMoE3uCq25EGVhXFjEJk/nsV5KKOq0CmXRfA0Zv/cggbSO7FI9rdvyUbQZZtFhNctISt46qVZyFXgTjRzagFIG05opW0Jr2VbkDLFeDKfI3/GYQ+/t90icd6GmKY3oam9rRT9ybq5FkfBbUer3EeBDZPNcRudsHNlCaRjKPdFrW6fVewlO192E+uVOGu90HUf23osohfUgje+yHKdMdut5np1ioPc1hDIO1hPW/SwcQvEuvV+LXvcKcnY2sswtq/W7CBlfkyh74gV0HnpJt+Gc1evei9a4k8iptyKpNKBmUWTpI+A76EC11/hcPWghvxttGqfI94WWFzrRxt8TPeyCj3t2rV16/DGNbhZ75LEeIUs6CQJhNRK3N0ZfdyBxuwUZNn1kJ4zMc1hvxG0SpZecwWupa6Ud3Sur0DUwSnYGTCn21a6fpTZo+5ty7O/ao0dn7NGDBNBNhHT6mwhp9V3Uvv7WwxSKth1DtTxZC9KiPG9alJF4ehH1BdiAovNpCRrz3F+P+mqcROnSeUxhnkLOlN+g67qExN4qGuOsKiF7p1LwXkKOgq+Q7XMG3Q+nkPFvjRttz50hNHKMN06y7KCO6GGpyfsImW5bkbjdiKLbPdHvNiOraAJ9xteBX6EIU7NLF7ISo/GazTzbpHMowvoCsl3+jDCOJm3ia8dj6D4YQmtI0W3JvEd0jVHkYNqP1oMtpJvC3IvWm28gkfsMOscrjmqRgnG06J9GoqHWm8za+9+KUoXOkv8LrRn0ETy8NjttA7rINxHSrOKRR2soNRE9xpAxc4EwEmoY3TgT0e/OIBFV9Mi6pdp0oRs3Hv3aHT02R4+t6BpchY6tCYW0DQlbSG326BzhvCV5rVkUzT2C7rlmetOzJG1Dtgud/zuRkDiDrvdK4RMXp9XeU7zG076337dH3EFi31drhBX/e/ud2eh7c2BZfeLa6Pt+dL3uRNfvarQedJJtwx7Q8RpC6bUfoeswC+xeyfK58yx2Z1Ck7B+RA6OPdFMS25GRdBOqQT2GIpR5dJqNoS6gJXRvfI8Q2W2U2DPB247SmlejiMptSOiNoDXlMlrfJwkid5ywx5aZvzZ0EJzV5rgeRI4Nq00eIIjbLEoRaiUucv8LSlsdpvn3UZZi1JwVeRdx00iUPEXIQNtDNt3K29E1ehMqlRlG9ZznMnitRlKEfQH0/j4DfkqY+LGW9NZCK4W6CTlCj6F7vVXtzAWpZkxZ+vKXKH05iTepj7Cwe/pywFJq16DmJPuQKBuM/fva2M/dhDmuJnTjXuVpdLFeZf6mPIy80JeQCD4VfR0j3PjxzWQm9u/xqHGjNoO4eKgW8epGx2QjWgi2oUjXGoJzYAM6Zn2E42ZiI6umFuPIuXAItYvvIBhtScYjTBO8qCaYnaXpQtfCN1Hk/irXClz7Wqr4uZLKhkalikel8LXfrdxIK1/TfsccNd0E51Z/9HM3EsBx4zfLCK5RRuvHaSRyD5Dd5pd1J9UiGDRXUcffrYSZsn2ksz5Zl+FtqObzKNq/z6Tw3GlTRkLxE7THTKN1cz9hdnmjsL3H9pw+dG7ie221LCpzIpeZvzZYNLcz9rWb+c7WZopbYwzd928C/xmJ3SHycQ/ZsU97vTDbxqLzeWccpRH/gjBz1UbLpY05yszRcxVlhTSieVGW+0Leo/fGJNobdhLOQz/p7Q2dyHa+FzlQTiDBu6KotrHYmKH3UA3L+gV+rxq9hPEtnyEDvggXWxZ0EaK1e1GkexcyzLeji7qXIM7M0I03SIob1vY1/jCxahvwBMHrPIRE7kUkfqdjvxv/vXHCxj4aPa4SPNmVRr4RTw+Ji4T4ZtVW8ag0CMzYt4jXIIp2rCGI13602K9D1+JA9O/xYxeP2mZZ4zOFjs8x1DjiZVTkbzMtk9YYTaB75ATNTxsrEm0ET+Uewhqz0LW60DmpFLmVv19NIC+1MVfes/Zvdg/Eo0BZOmQWYxqtDe8Cr6KUzSzTlrPYA6wOa5p8GOmLYfV3fwBuR+vZLoIzc7m0EWZkPkLowpxVlH45mNj9FO0xw2hu7O1ofc9ihuhSxMVqnMr9tprTplTxfbVHXhhF18brKIr0KvkaKZaV48qE7gTFCb6MocyMX6JMkNVo3UgbE0Mbkd0+hK6J18i2OWHWmT5mgxaBeIf+jaS7N8RTmG0c3RDN66zeFBYSsFdRRNe6L9faNKIdCbt7UVH9OVqzwc5idBLSEe9GAvc6ZISsR2LOvLzm4V3qgl7q/yvFrwnXKSRkJwiRWktpnmR+04256N8sLdpStOKebxOylSLbNo9yxc8d0Wc1L3d77N8sktUde1jku59Qq2yCuIf5tctZC1vDOtNeRYLgAPA8EroH0Wf+HjqvSTyuc4SOl8dp3fskq82mHTk8+rjWwVDLNZHW+6qsx618/jwZuhAaUH2Jojpfkm0qUxZpZPac5tzLi6G+GHNIZPwS1cUNkO5YCdt3bkE1p0eRmMyjYV9Ga98XaJ+5jNbQB5GjPK2GXcslb0J1OQyjveZVFCl8Cx33PImBLGt0zcFfhIgu6D2PoDV6Nypt6SK9TJA4lhVizdIuRI9PyLYEIiuxW6SILuh9HkH35S5kA68nveyuNsL85G8iu/Vd8lnekgkLCd0pZHx/iqJVSbojDiBRdyNKv2hVA74SG/dwHWoXfg9KQ9hJGGdjnRWzWKjsOS0VC7SQrGb+5hEXqpUF+7YhmDg2I6lyDE/872cJm0e54mcTtnFxGhe78Uh2/BGPdFWmjzYC+xyTBC/4J8D7KFPhI7QRTCHD8iG0ESURulYjeQxlUOTRIM07i6Uk1/q3aVFrFLmZmHPlGDJ430Ap81lSLfMjDaYJIqkoxuskEhh/j/YEqxFN03O/HRmrh1C2yIUUnjsrJpGd8Sy6Di+h2Y822jCP91DRsD4QB9AIod+hfSwPNbmVZCFQ4k6x8ZSfO2vKhGjfdmTX7UH3edr3RhsS0btRJ+az0WtnlZWZVdTVAj5FiuiCrs/3UNbPNnSObZ72crEeApvQSKnHkE17IoXnLgQLCV3rvvw28gBsIoinpehGtZP3Ia/BCK1vxNuYh7vQRXQf8tqvIQi6ZhCPxtZKXAAvlKJVXuD7aj9XSwNdKL2r2YbNHEHoDyMBehAJgvei74dR1LuMrvX7UcrJIMnTlk8jQ6+V5+e26ucqGlZbfgJ4BRm8h8h+ZmYWNbpz6LOcQ03cmjX3sx6GkeC4idCVt4d01r525JTei6Kjh1AKYp6dzTPIqH6V0Gfiu+gzrKIxNeutiDmth1Hk/Neo4/UXXNvTIC/Ey57SxNKWiyZ0Qe/9CJrHbXOWryO9NSOONafahxxOF9F1czHl1zGyugaL0ruhkmHUoX8/iujuInSoXy5thBTmbyI74FeskBTmxQTYGEprO468SKtrfM42JPBuQWLvKK0rdEvI62LjHR5DkdxNaCEq4iZdGb1tdWxznUIGwEVkPH+KorgHkTf8EtemeG5ANXFbST4DbQw4jO6PVhjU7uQXE7mnUX3eb9H1Pd6g1047SjOHNuiDyAgsWgrWOTTqYTfaK7eRXk1WJ1qXbkaO1wOoBCnPzKHI/HuE/hLfQ2U/a5Cx59ROGe0pF1FW3a9QpOgwjbnn68XKodIUu+a8tkZLRRO6ILvjQ9Qhuw/dD9tJTwTFaUe1wLeh+/Ai6kWSdr1uVmnq9txFi+iC3u+XKIV5A9IWm0gvUGYpzDej8bHHUBCnaPtnYhY7gDNI9X+AitSTRKz6UPrR/ciDeJBied1rwS6aW4AfIJFr3SNrjX47zSMevR0izE/8LHocQNe/1SpXYzuKPKwh2YZjacsHkfhoZUdQK9W5FZE5QqfV15GH/j0aN2t1oQyReimjCOV5JNZPU5zUZWMO7as/Q3tILzJs0vTc70A9Ir4k9FzIM9ak6rPo6yUkdu8jOI59HVkci+JeRc6ND9H9/gKKmufdoLVyIetqnQbWRHIkehRtrTCuorKHfmRjdhEyLdO8L0qx576L0Mz0PdJ1klRr2pgWWTa6ypoJJD63oR4Oveh8p5XCHG889jiyeY+n8Ny5ZjGhO4c2m3eQ+t9E7U0iOtDBvBedtJPkswNkvZSQgXIX8CfAt1DUu4+VFQ0tGha9NYF7CV2bVnd7GEWIziJxsJgA7UCLxUaSRxzGkSFyjBWSOuI0BatjPYlE7m+ir+dpnCEQH2O2XMxTP4qyLkzEFZGrqKmddZnvQamJaewfHdHz3oYceOeQsM670LHMgyPIXriI3vs3CTOI04p8txJ2X9jYu8OoGegb6H7PU2flxZhG590ir5UN/pJigmcCHYMrFFfogo7LK8j2bEdOoA2kL3atXncX6j8yhNbZA6S7hmRR1gJhzylaRNcYQef5BpTCfCPSXmnsDeYI3UWYrftzWtwOXSokPoku7gOEIvhasXEH96D0mTGKvcjE6UXh/x8DT6IUtLx0inTmE+9CPYkWkcvI+P8c+BgJ3cOE6O1S12kJGV5fIzShqnWjmUObxufIoFsJw7uLuuEUFUvFHyI0nnoWeeUv0Ph1OI10xLjRegHVn54l3/WnSzGEUkpvRAbNHtKJXJqhugPtv6cI9cxFETxnUL3a+ejxdXSc1hGa8ax0wRu/J66gDKSPgZeQyLWMpKIwja7V8+heSEPA2djEc+gYFT2z8AoqPbEJGPciGyQLsbsK3XPj0euOo/0kjWOYZepyVrXejeQE6qOxGZ2HnaR3juPjGX+I9tEXyXdZw7JYSujOog3ndbRhrqP2tNxOdJLuR1HhsxTX+x6nE110j6NI7m5c5OaJ+OijWWQEWHfWs2ihtsjtl8j4GyaZwdwDPIoiukm7pk6he+qj6LVbNW0Zir3RFBHrLjqGDLvP0dr9InLmjNB4oWPplMsR1+b1t+Y6XyGhW/QIDcigeRbtlQOo3j+NqGU7WptuRDO+T6Fjd3mZz9soyijj5j0kfI6gCMQ9qGRkADUDXIli1/Y4c9yeRSVir6B05Q+RE6UITo04U4S9eT/Lz5Cz8qQRdA01Y/3LgktI7Np6/wihljPN+6EdZZvcSuhh8mtULrLc41gm/Xpse17L2ivyubYuzOtQzXQfyh5Mq163HTlX70CllyeRk6wlbbZaDtooOuBfIQ/xmgTPP4A6uN2HDJNxim/Yr0K1T4+RPMpdjXKVh/17vMYxi7FErYItbvHZwVaTc5bQ9Omr6OtJgocy6fXYhu6DJwgt/2s9L1aHdgAZJldo0YUlRpHrZYqCbezWgOYkWrOtW/hxmpc5YDW1yzU8LA37NLp/jtEa5TATaFbmemTQ9CLjZrmNDK0eawMyZi4hB9ubFCsKfhU5JC+i6/g4siduQMa9NedZCXtjfFTOKDqfR1Ag4T1UwzlEcbOE5tD5/RjZWBvQfVDPua10BgwRpiW0AhdR9/YScrw/wPKO10K0o/XoDnQsryLH3HLHlsUz7dKsx7YgRxE7bFdyFXXN34LOwd3IeZlGk1urxd6GOvQfQ3trnsfR1U0tQncGHYQ3kWhN0vK/A3mqH0aNQ86S/dzGLOlA4uZ+FPbvr/N57AafRpvWTMW/WZTC5s52IwOom/mzeC31wx4miFeKKDav4CQSjZejxxkUKTmNxO1BFNEYZfnOll50Pd9B7Z3IjWl0D7yLxHeRDM56aIUUorxi0YoZtCFaOv5HKHr7DrrGhmluup5lVUwSroUka5M5sSxl+SAyhE9QXIO+kiuos+lGJHZvR/vscmuy2pARvAUZSSfR2niIYhmBc+jcv04oObkbZdRcT2jaYntjq2H3gEXVzqN7+z10v7+L9pUipSkvxCXkBD5MGK9Sj2FvGSBj6P4apvhBlkouoGydQXSf343WjzTFrs1g3YB60the83L0/XIwB209+8JCxO3BoqepQ+h6bSnMN5FeLwcrcdlD6ML8NC2YwlyL0C2jg/0C8A3kRV2V4DX6UPrUfWiDKvJc3V7UZfd2woJSK7ZZ2U04QmgUciX690lkvFnBfzta6NeidOntaFM3gW2/bwZvZ/R/AySb35vWItNILHp7CW3ynyID7hQyhk6gjWAUbXZpLHoWzf02ydv7l9E5P4AM9QsUy9ish3gKuVMf5sSKdy824WepqKeQcfg5Mnq/iv49Dxv9LGG0xxR6/0nXzUlCvfFHaOzXefLx+dLCalJ3oqyp6wm1qMuhDe0JO5CD9jhaM7OajZklY2j9PIvSc+9GdsVt6LhZw8y4M7ioxKNTFpU8j+6BT9Dnfx/d+62Q2WBMEsb73YRszaR16+YAH0PrxrnoayutF8YZ4PeERm23IXsxzTRmyw7ZhO65i+ha/JD6bfl4kMKiumk0HrOo85llvLc8MYecPs+h8zpAer0cIEyP2Qf8CNkRH6XwvLmiViE0jdKH3kFCr5/aPQrmDboXeSBPoRuliFGebuQd34YutFoxY826SH6ENqwjBKPUBO4U8z1cnShtYS9qgHVr9Poz0XOdR2JuMnpPO9BFa2nm1cRYvObN0go7CeJ4KY+gLSrx91mZXl3ZPj6ell1ieekX1mjnDNoQP0JpWweR02AieqQtJAeQs8eiuUlHCp1D99CXtE40aiFsI5ugNQ2MLKgUtbOESM4UOpbWGOQiYTTAUWT8niZ/nuxZwvsdJXijl7p37FjYaJCT6F7/EH3eVkv7n0Xr1/PIoOxH9brLrUO1iMxqJBweQ3vwHynmGlRGa/8Q+hyfIMfzzWhd3oWcwVbD206xorwmbm0ywBW0bxxG1//naL87i+6LVroHQJ//HHIG347O5QZqF26278THBp6iNer5q2Hp3r8jdKS/Cx23btK79kvR8+1AnZi/QrbnV9R3DZr9OY5s0Dnqe6/xPTNes36MYq5v1ZhE+95qwnzdraTTnCpe4nIzqvc+zPKj9bkiScTvPIrqPkRIFar1IPchD7XN1R2hmGmb3ehi66d2oWYL72kkcj5ERsYhFJGxiOxiKZ7HkTh6DxkrO5EBaOm5owSxugvl3D+INgrzctv7tZRH60R4FHn4V0fPux4JazMSKs/xdPSere5lEglkS62OR5LjRvsMwXPXh0R4PfOGJ9Hm9Qk6nm+hTfEsWjSz2vg70DX8JFrsk4wUikdz30XvvxU33UqmkBNnHH3eNGpLikils6fyEY/WmvNpEh23qwTD/mL0MOPtICFlcYx8ZgjY3nEYZfb0oyiNGa6lit+1r3YshgkjwD5E6+AQ+RLzaTFGqNe1ESJmsC6ny7AZqevRnvAYOqYHyOc1UytD6Jo4gmrZ7kDG2i0o6rEFRUF6qM2B2wziTmNzDA4TaqqPomv+E3QPXELXSSszgj7z++gcmqOmmtitXFst2+Uyso2OEMb4teKaAdorDqLr4izaI+5n/kiuNFNd46NphtF9WA/jhL1rmuoO0HLFV/s+vnfGsx5OolFqX9E6QhdkB7yL9oNBwgjXtMRuL9IK9yGd9/kynzNXJOngNYMExdtoE9lG7Qe4nflR3eMUU+h2ohs96YZpswF/ikTueZLlwc+iBWUUGbg96CavFrm0hkuHUGfo+9CCZxHIcUKXxg/Q+TiDxO3NKBp8L0rNXYMEXRvBA2eRpC8InU/7kEGxGt2E3bHft8ckEj4dKDr9YPQ39XQsfhN1/3slOh5ZN5koISPxcZQaNJjw76eRU+FNdNxaaQFeiDK61g6jz95L6KK50Dlf6N+rndvFznfSayHp+1nodSo3Y/saN2btYY4fi9ReJTRQG479bAbvuejrUPT/Y9HfFSGicxatMxvRZ99CKK+Ir6XxY2UNd06hTfe96OsZWvv+sfq31eheuQWtPRadSSJ47ffs2mtDa/pNSBRaJ+YiM4f2oCvo87yOOvbegj7nPsJeVhnlbVZ6s50Pu86teeIVVNLyFbIXDhMmAwxRTJupHqaQSH0X2SWzyAm/ipChVuJawWOlYRbVO0gYHXiZYqyV9TKJ7GrLADiJOpTfiNbdHuavH/XseXace5DT/1vR67xNfdfmKLKFL6J702xrszfnKh6V4jae+TBMmDTw++j7VnNsXET6wYKMd6G9wcRurXtDNbtlhtCJ2bRCy9wvSYTuHFp8fouiuutJliduQ4rvRKk3w+R/gH0ldpPZDVfrZ59FhulRQvS1HuYI9W4LMRW9zkW0QZ4kdKnsIKTP/hEtUBfQQtEZ/Xwd2hjuR8J3CzrPU2hR+gR1c30Nbcbj6AbpRRvRKmRM2KZj9cPWdXUTamd+OzI+aj2G09HrfwT8CqX4naExEYleJP5/gBw8SSLRZXStf4qizyslmmsG6JvomhpGjo24oVni2gV6sShf5aPav1f+3VKUqnxf+X4qf28hQTvHtRt0XNjGo7WTaE2w0Vfn0bVxAR23q+jeGoseVo9fRC6h6MwcWjMsc6SfEG2LRx1m0Wc/j5xpNnP6JDouLbMBL8ApZKxZ9souZHzYsaoUaPGSELhWEMeFgM2zt6yaogvdOHavnEdO3O3IMXkDOoa7CD1GBtC6bs6Wak6EekTwYmuTrQlxx9YVdI6/IvSXOEQQt5YNs5Kw9PRP0Hm5iM6hjeCKl2PFI7kmei4TMhY+Q2tHKzTqWopZtNa+hq6ld1BA4TaU6tpPuN4rnT217MPxjsZT6F66izDJIinDyPG/jRAwiUft7ZyaDVnpILY9dBhdIycIzqGWSr2NmEXrwu8IfSv2oXXcGtWa47hyHYvvmbYOxX+20sqT6P5ZsUIXdIF9iATRLnTzJI3qWlMqi1AUKXXKuvglbarSizxg/wx5Kd9BN+UI2XidTFy9Q5hBeE/0Pr5EA+UtGmtMRr97gbDZPoKaDwxGv/sZiqK+zcKpIfH62/gGD1rItpK8kdMMoRvjc6hpS6OunXa0EH+H0HU8CdPovb6GNt6iipWklNE18z46BjcjcdNHWJDjBmbl18pNNm4oxn+u/Gq/X+39wPwFvPL6q9zw7f1U+91Kg9beg0Vq4puybdZWY2sCdiT2sBSueBf2VtloZtB6dxkJkDVoTelHa1IPoeyhHR2/cUKats1gXylG/yzBOP8MZeSsZ74wq+a9j1+z8Yf9nx3Xc7R2n4BptF9cQlG9QbTv7ENz7zci0bQN2SQD6BrsJDSyquaMW4j4fWr3vRnpU8x3cNkIsLOEa/swOt8XkYE+wcq4zhdjGq0ZY4TRlpsIGWMmiOJC17pSX0bX+Nno6zDFsjOXywSylY6iwMBNKAvT1hATRZ0EJ+Ni4heCyLUMBHPQXiQ435Ie4xHkCL+I7sV+5p/X+D0Uf9j9NEFwbsX31VbOfJhGkftfo/viZnRfWLaK3RuVWadxx1C8QaidN8v0fA+tmy11v5TK5cS2VBsy+v+PyJuTpFZ3ltAl7t8Dr1Isz8t64E+B/wVKjaq1IZWlcQ6hhfd1wnB3S8WzC3CK+iJTC1FC4mwPWkiORe9hqahyPzIMHkHGwFkknA9FnyPpRtyGInt/iQT/PiR6FsMWOzNYngb+E7rRG2EIWMrynwL/M9QILMnc5DJaxF8E/u9I7LZc6/YlKKHzHBcyCxmRlUJ3oRS1hX6G5d071YRuZWTMvi70fbWobjxVsTLSuxINWhMS7cw3suwBoe7KnAQr8TiB7pl4k8BKR1ClMbPYI26smuhaSce1mzCVYAtyPu8idDPtiz16CI4Fc8wtVENo97xds1ZbH3di2WMYRdtOEBq0WY2icy0lQqNMe8QdPZVORhNHdp23WvpqUtrQdd/D/LWks+JrXPBWXue2btixtWt8ipCVVK+4LBFsg2qZKpXO7WpObvt+pdFLyIqqPK+Vddnx+8XuEwtKWQaVZYsWeTJOVZJGdEEX1ftIrG2PHrXWrFoO+F0olfMw8joVZbOdQJvUceQZrrUjphn7fejzbyWkcH+MhOcEEnQn0KLRTmiVvpzh1xbd/YBwUdfCVUJRfz+hVX+9N8AgOu9fR9dM9xK/b4vYJXStvAr8A40TuaCF5C4kdPeQTORCSD18HQn1VvY0LkSZpdPtnZVFvPGWszgWxXCWj0VVL6F97QNkJFpGgZXfrEb7lYlf+/94WmVlFNciilYvOISyo0YIUSY7l9Zl1lmaMn4PLAeL1iVxsC/k3M2CMqHXjJOMes4rrECnQD1CFxTd+zvCPNkk6ZzdKE3h64SubRcoxsGfQBHND1BEt5+lBVsl3YThz3uBJ5BHbByJogPRz21oQ7a5q2kU1ycViHPo3FxY5ut2oHP+AGoUMkhtad+j6Hg/jRp5HaVxBkIbOk+Po/e8JuHfz6HUHqvNPY8bN47jOHmgTBC+Vyr+zzIOOqiexmxU1kBXRr58vXeKSGW/C6c1WLHntF6hCzLgX0O1E3upvc11CXlK9wOPItH8JsVoFjBLGBN0J2GuVT1jUyzCa51IrYbvfsImOYyE7rPAH5AXuohRkNUoMnovEo+1RHNnkcB+E/jP1D+vrR5KyIHzddRZcAvJznEZefGPoTT9A6zMaK7jOE7R8IwDx3GcFmE5Qvcq8Peoffk6JPhqbdffjdJ3H0aRynOoeL4IG8tVVLD9GySGrHtwvXPKSoSc+t7oucxDPIWEoTXx+TUST5XHKZ6bD6GAvxaPsuXpV3ZiS4sOwmzfG1AUfCks3fpL4D/SWJELOg/3oHrivcgxk4RZFMF9l9Bswb37juM4juM4jtMgliN0QULkKdRkqJ/aRIzRh+pcv41SPEdR/WXeBcEcEuYvI4HfjaLTq0lnKHe8c2YHoZGANcd4BqXzTiOB2odqXm1e4Gz0/5+ieuLF0pVNWG8mzKysnMu7ECau25jfaTIuSEsoZflR1NRqM0tfcxbNPYdSfj+lsSLXUpYfRR3tVif8+3jK8lMomuv1RY7jOI7jOI7TQJYrdKdRSu2DqMX1dcxv2LAYlsK8D/geat7wNIqE5T2XfAo1F/oVIRq6H0V4lxpFkJR2JKjvQMdrPZqjdREJ29uj/9uDRNkMioA+D/wCjS2oFK5t0e9+Gx37TdHzvYBSbU+zeHS9E33eu6LnOY+OxzFU72TieiPwDeBJ1OFyqS7LoHN/GY2g+hmNT2lfgzINHkWCN+nMXBuH8AoaDeJNFhzHcRzHcRynwSxX6IJE0f8P1TEOIiFWaz2jCa7bUeOlC8AfkdApgtg9gJokjQM/RBHAddRer5yEQYKY/gYSUAPR660hDFKfA3ZG/96F6luPEI5nF3JI/GskdHehyO44arB1O0pJ/4DqdaW9aBbyX6J64lWEaPwb6Px9iY7Bo8A/IcziXQrrsHgmep7PavibNBlADbP+EqXkJ52Za/N+P0R11UtF1B3HcRzHcRzHyYA0hC7AJygauxOlL/dRu9DrQOL4IUIU8RUkdvPOFIo8XgROAn+FhnNvJaS8pil4e5Ew3YZEoXWEjL9GO8F5YPPn/gMSov1IfP4rJD43E9LNbaTCGuS0+DdojFR8FEIvas70T9H52hr92xyKzN8OfBMdk05U53oztTs/bJzQ5yhFu5HR0C4kbr8H3IbecxKsrvgA8I9IpHvKsuM4juM4juM0gbSE7hTwIhI/W1An5iQpn12EWs5ZlMb8JsXofDiHOkc/hcTk91Hq600oqtoT/V5agtcGqC/1O31IZHYhcXwIOSLuQ7XRa7n2/FvddA8SwE+j0UZfIFH9XRTttO7JXdHftUXPtxqd+wei9zAYPVctn72MospHUdOtQzX8TVp0oM/9fSTUt5L83rAZy28ix493WXYcx3Ecx3GcJpGW0AWJvP8RRQQfR/WZSUaydKJa0YeR0G1DNZ+XkIiYId+NqkbRuKVPUc3yX6II5x6UAtuz8J9mQglFW29E52ISNbYajP59IfFpM29NKF8ADqPzcz1Ke64mkkHnrLIpWa0CfxZF8Q8AL9FYobgezTP+EbXXEscpo+v0PRTNPYGnLDuO4ziO4zhO00hT6JZRbeJ/QULoQZJ1Ii4hIbYT+A4SS2+i+tITKGo6Tr4HsZeRSH8eHYvHgD9BYncHOh5J0rqXix3TzdF7K8X+fbG/aUPR6LUoGnxT9H+9KIq72Bipej5bGTkzjqJo7vE6nqNeVhNqifchR0DSz3ARpSr/e+ToyHsWguM4juM4juO0NGkKXVAU7hXgVpTCvI/Fo4eVWBTyBiRAbkIRvk+Aj1Dt5hUkivIsJmZQQ6WfoQZNd6OI4c2EiGg/2TStqqQWcbvY33RHj2r/nxbxaO7b6Pg1ggHkkPkT1OhrDck/2xByxjyDnBuesuw4juM4juM4TSZtoQtKdf0piiL2IWFXKZQWo4SihluR8NiF0m9vQnNV30aNn0ZRbXBeU0TjUcqTaHTPnUjw3oFSZDcgwdvDtU2l8kLW72kOjeQ5hMYmncz49Ywe1GX6SdQ0axPJ5iCXgRGUXv8bNGrqIvnvFu44juM4juM4LU8WQreMInP/DkVnu1HNZ5LmVJY+2x/9/WpUZ7oBpdS+iYTRBZTOnGfmUJTPot3vosjuvWg8zw4U/V6HjlcPOlZJRFdRKaPo7TmU+vsujYmI9qCsgx+h9PKdyLlSq6g3J8YJ4Dng75DgzavTxXEcx3Ecx3FWFFkIXZB4+RT4BUrT7UYiNcnrmejoQHWT3UigWPSzjETRDPlOY44zFT3eQd2Mf4dE1i0o2rsLid4NKK22G33+VhS98Zm57wA/B75qwOv2IEfDj9EooRtIll5v4vw88Crwt8BBinMNOo7jOI7jOE7Lk5XQBaWjvoGiuYPAXUj0JunEDBIgJSRQtkZ/P4migJej1ylRrJmlJnjHUOOlt5G43YVE2J1o3M0mdMz6UMSxE33+DpIfxzwxiyLxZ9Fn/zt0rWQdze1mvsjdhxwKtToSyoRZv5+iFP1PKda15ziO4ziO4zgtT5ZCt4xSi59FQrcPRS4HqS9CaY2RNqG00yOo/vU0xY2mzUaPCVTveRKNqHke2I66Ne9GzoKNKL15NTqGVtvbhc5jHut741gkdAq4ShC5P0NjmYYzfv0OFD3/FhK5+9HYp6R1udY066coEu3NpxzHcRzHcRwnZ2QpdEHRr+PA0wSxewMSafUIMxO7W1FDp4+R2B1HAqrIjYDmkOA10XsM1awOoKZca1HUdxsSbDtRfe9GJH5t9I9FfRcbAdQI5pCIn44e4+hzXULOiQ+QE+Qj1Fgsy3NnIve7wA9QY7OkDpc51PH7MPAU8FvkyHEcx3Ecx3EcJ2dkLXRBUbyDqAbT6k53o2hkvWK3D6WgfhOJjdHo0SrEI71XUB1rOxKx3chRsB6Jtz3RYxuKdq9DIs4aW3VFf2tdndMUv2VCOq89ZpHTYQJFbq8QxO0hJBS/Qo2cTka/lyVd6Pg8CfwEuJ36RO4ocj48C/wSvfe8znN2HMdxHMdxnBVNI4QuSPh8gtI9LfK4g2Rjh+K0oejmg0joHo6+tmLXW0v5nUFpsqNojM0p4Esk+lchgbsJCd4dKOptXapXR7/XTYj2WsQ3HvktVTwsylpmvqi1rxaxnSREo8dRGvIFJGaPR+/1XPQYJcxBzvp8dQF7ge+jDsu3k7xO3CLtp4GX0TV8iNa81hzHcRzHcRynJWiU0AWJhQ8I9aXfRB2Gu+p4Lpu1uxON6DmMooZnWTlRthlCJPs8EpRdKIrbh47x2uixBgneNQRHQxehk3Vc/NoDgqidqfKYQgJ3DKUkD0dfx6L3dAXVs46gcz8V/V2j6EKzir+Pmk/djj5/EpFrnb3PoJFWP0fNp4paE+44juM4juM4K4JGCl2Q6HmLkFr7ILCZZDN2jVL0HDeitNQzaE7tRVaO2DXKhFrYq+gYlND57YweJmwtjbmDkNq8UIQ3npY8W+UxjQSsfZ2J/n0m9jvNqJvuQx2Vv4uE7nJE7jnUNOsf0LWbdaq14ziO4ziO4zjLpNFCt4xE2CtIfLUBD6CU23rEbhuqVb0bRRHbgd+jaOJKTy2Ni9+FqKzbjX+NC9TKFOb493lrALYajbL6NvAY6vS9muQidwqJ3HdRuvIfkaPGcRzHcRzHcZyc02ihCxKgp4EXkFBtR+nHG0gudi2FeTPwtei5JoD3USrzWBpvuIVpJWdAO0rTfgz4Drqm9qD65aQidxqlg78H/AJ4CV1PjuM4juM4juMUgGYIXVBq6ykkdruRWL0b1ZHWK3a3Ag8hobIbRY2PoNrRvEUdnXTpQfXa30Tjg+5Czo9eks/JnUKNtN5DY4T+gCK7fg05juM4juM4TkFoltAFid2TKNW4O3ovt1O/2LX5uo8Cu5DweRrNab3CyqvbXSmsAu5EIvcJ4FbCNZRklFI8XdlE7vPoGm2lyLfjOI7jOI7jtDzNFLqgFFGbTWrppcsRu12o3ncQpbH2R9+/hwSMd8ttHdrQdfJN4FuosdkektfjwvzGU+8SRO4JXOQ6juM4juM4TuFottAFRdEOAb8hdPm9AzWZqkfsdgADSPR0oW67W1AzoVOooZBHd4tNL7Ad+Aaaj3snSlXuI1mqMgSRexZ1V34KpdSfoLHjkBzHcRzHcRzHSYk8CF2YL3ZtTutdqEFVPXN2QWJoZ/R1A0prfgP4DNVgTuB1l0WjhBwXd6LmY99CXZXrSVUGnf8JwpzcX6LGU2dwkes4juM4juM4hSUvQheC2J0lzGS9B6Ui1yt2u1GkrwdFda9HYvct4Ciq3XVBUwx60Tl8BHVWvgc1HRskeaoyKKo/iaL8rwM/B15G3ZY94u84juM4juM4BSZPQhdUQ3sUeA4J0Gk0JmYzErtJI3agz7gOCaX1wHVIIL0JfICid+O4uMkrnUjM3oWuhe8ih8UGkndVNubQOT8NvAr8DIncS3iU33Ecx3Ecx3EKT96ELkjcHkdidzJ6PIhSj3uoT9iUUP1mF6rf3YrE0m4keL8ELqNIspMP2lFH5b0oVflxVLu9C53DdupzfMwCo6ib8htoTq6JXMdxHMdxHMdxWoA8Cl0Ic3b/gETJCPB1JHL6qC9VFfR516BI4BokdPcgwfM+akA0iqczN5MSiuDuQh24HwQeQJH4NcjZUY/ALaPzegU4jNKVfxN9vbzcN+04juM4juM4Tn7Iq9AFiZIzKNp2BQnQx5HgGSR5R+Y43aj2dwClwN4A3Ii67n6GOvCO44K3kbShcVAbgHuBh5DI3Q5sjP6vHoELSlWeBi6i8/sSGh/0IXKiOI7jOI7jOI7TQuRZ6IIEymUkQEeBYTRS5mYkiDqpL5WZ6O8GCCOIdgO3Rq/1EXAACd4xXPBmSQcSsZuB/Sg9+TEUaTeBW28EH0I97jlUk/088CJKV59YxvM6juM4juM4jpNTSuVyYXrvdCIx+k3gCeA+YBtKQ16OEDImUXTvDBK57wIfI0F0GriK1/CmSScSsdvQiKA7USflPdG/9bN8R8wMcpCcQg6M36LmUydRhNdxHMdxHMdxnBakSEIXFIVdh0TRT5Do3QGsZnmpzEYZidkRNGbmMEpv/RD4Avgq+r8pvDtvPZQIDcG2oQju/agGdzuK0vdT/zgpYw6J3MvoHL6B6nHfjv7NO2w7juM4juM4TgtTNKFr9AI3IbH7IEo53oIigPWmMscpo4jfVSSMvkLR3XdRjecRYAhFgQt5ABtMOxKvq1GN9T5Uh3snajq1ATUZW24Et0xIVT4LfIIiuL+Pvh9b5vM7juM4juM4jlMAiip0QeJpJ2pa9B3UlTnNVGZjBgmnYZTW/AlhBq8J3qno9wp7MDOgDYnbHlRrex0hgrsfnau1LK+LdhyL4g6h7tnvoBFVb6DUc087dxzHcRzHcZwVQpGFLigVdgBFd/8a+Bqq8VyLooMl6u/UW8kcIa35NPA5ElOfAkeRCB5FkeDZlF6zaLShFPIeYD1KK78eNZi6Kfp5E+qaXe+YoEosijuGGk59AbwFPIuaio3gDgjHcRzHcRzHWVEUXegabUjgfit63I0ElXXsTUvsgkTTJIrwXiKI3veAQ9HP5wnjiaZp7ZpQE7fdqH56G7AXpSXfTBgPNEA66clx4rW4x1EN7gvR1xN4FNdxHMdxHMdxViStInSNPuA24J+hFNnrkMjqJN3oLkjwzqARNcNoRqs1sDqAhNYp4BhKpx1Hkd45ih9hbEeCtQvV1+5Ac4hvRzOJTdwOolTyNBqFxSmjY3kVRXE/QSnKzyKnw9WUX89xHMdxHMdxnALRakLXuB74NkplvgdFGQcJ0d00BS+E5lUT0WMselxATazeR2OKTgBXCLN548I3ryeihI5bW/R1AM283UhoKrUXRdDXoih6H8vvnFwNO07jKJr+FUoffxY1CjuLzzx2HMdxHMdxnBVPqwpdkNC6Efgz1LBqH/PTmSF9wRtnDoneqyjaexGlNh9HNb0nUDTyEkqFHo++mvhtdOS3FHu0RY8OJF43AmtQpPY21ExqI0pVXovEbw/ppiXHMYE7jRwFp1H97ZsoVflLdKwdx3Ecx3Ecx3FaWuga64BbgB+gea3Xo3TbeHfmLAUvSKRZpHcSid8xJHJPRI9jKCJ5Jfp6nhD5LTM/6luueO7Kk2g/V/tcJa4VtR3IAdCPBO0gcgrsQ+nIm9DxWhX9nwnbbtIZ57QQ9tlmCOnhn6Ma3N8igTtMa9dAO47jOI7jOI6TkJUgdI2tSOg+BtyH6nfXM1+sZS1441gX5wnC+KLR6N/GkNg9haKX56L/nyA0uJomRILHCE2vyrGv7YS0Y2saNYAEaxcStruQoN0d+/cuJGxXR7/fHf2bdbLOGrsoZ1HXZIuGfwA8jWYZX8QFruM4juM4juM4VVhJQhck9rajmbuPEup3VyNh1wzBG8e6CE8hERsXtPGaXhOys0j8jhK6PJvQBYncLiRwKx8d0aMHCV77PfuadrfqWrH3P4yaeB1Gdc6/BT5Ewt8FruM4juM4juM4C7LShK7RDuwEvgs8TBiDM0j6I3DqodaTYmLXZvdWpjTD/JrbeIR3IRHbLJEPErCjaFzQMYLAfRc4SX4bdjmO4ziO4ziOkyNWqtAFCboelNL8MGpYdRMSwOsIKbtZ1qCmQT0nsJlithJrMnUV1Sd/hcYFPYs6Kn+FR3Adx3Ecx3Ecx0nASha6hgneDahW9X7gTtSxeSuqW+2l+VHeVuP/3969vEQZxWEc/5qmVuYlS7KUgpKgXbXo7xfaB22SiiA3I6XdtPAyWovnHN5XC00dc6jvBw7zjjoIunr4XU6dUa4V3GUye7tIE3APVqklSZIk6UgG3cYACbMTwBxZXPUUuE/meCdJW/MIzbZmHd8umSfeIFun6wzuIvCKLODaxiquJEmSpBMy6P5qgATZMeAmaWd+QuZ458rX6vKqYfqrDbhf1fbk76Q9eYXcJfwCeEbmcTtkmZYBV5IkSdKpGHQPV7cWT5Gq7gL7Q+8NsrF4tPycGnVJ1hap3q6SduQlslxqiVyhtIoBV5IkSVIPGXT/TG1rrrO886TS+xi4R0LwdZoFVv9ra/MuaTtutyZ3gNfAcxJuV0i43aK5LkmSJEmSesage3wXyD2zY2RZ1RzwAHhE5nln2d/a/C8vsaotydvkPt91YI1cBfSGzN6+JeH2A9msbPVWkiRJ0pky6J5cneUdJpuZbwN3gYflzJMq7ySpBF8iAbnfrys6Sm1J3iQzt59JuO2QcPuSZuZ2lVR3u+VzkiRJknTmDLq90W5tniQLq2ZJ2F0A7pAgPEFmemu1twbfevplsVVtKd4rz11Std0iwfUTma9dJm3J74D3pHK7RkLwDoZbSZIkSefAoNt7AyTADpNAOw3MkMVVMyT03irPU6QFepTM9g7RBOBB9gfg9vnd7zzMD5r7aPda7+vzLk2g7dIE2nq+ktbjFdKWXKu1dZnUNxKErdxKkiRJOncG3bM3SILrRRJop1pnAhgvz9fK63Q5V0n4raG3noPBt/2e8lr/qe3Xg4G2ztbWs0kC7Tq5AqhDqrT1ax9JJfcLaVneKad76r+QJEmSJPWQQffvq4F1qJwagK+QcDteTr22aIRmxvdyea6fG2m9H6QJu+1qbft5k1Rfv5azQSq2W+V7G+X7dfZ2gybM1mPFVpIkSVJf+wkUqOvc3e1zOAAAAABJRU5ErkJggg=="
    id_prop  = p.get("public_id") or p.get("id") or ""
    titulo   = p.get("title") or p.get("property_type") or "Propiedad"
    ops      = p.get("operations") or []
    op       = ops[0] if ops else {}
    tipo_op  = "EN VENTA" if op.get("type") == "sale" else "EN RENTA" if op.get("type") == "rental" else "EN VENTA"
    monto    = op.get("amount", 0)
    moneda   = op.get("currency", "MXN")
    precio   = "${:,.0f} {}".format(monto, moneda) if monto else "—"
    loc      = p.get("location") or {}
    colonia  = loc.get("name") or ""
    ciudad   = loc.get("city") or ""
    ubicacion= ", ".join(filter(None, [colonia, ciudad])) or p.get("address") or "—"
    rec      = p.get("bedrooms")
    ban      = p.get("bathrooms")
    m2c      = p.get("construction_size")
    m2t      = p.get("lot_size")
    parking  = p.get("parking_spaces")
    desc     = (p.get("description") or "").replace("<br>", " ").replace("<br/>", " ")
    desc     = _re.sub(r"<[^>]+>", "", desc).strip()
    fotos    = p.get("property_images") or []
    amenids  = p.get("amenities") or []

    def fmt_m2(n):
        if not n: return "—"
        s = "{:,.2f}".format(n).rstrip("0").rstrip(".")
        return s + " m²"

    SVG_BED  = '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#2a9db5" stroke-width="1.8"><path d="M2 20v-8a2 2 0 012-2h16a2 2 0 012 2v8"/><path d="M2 14h20"/><rect x="6" y="4" width="4" height="6" rx="1"/></svg>'
    SVG_BATH = '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#2a9db5" stroke-width="1.8"><rect x="2" y="11" width="20" height="4" rx="1"/><path d="M4 15v3a2 2 0 002 2h12a2 2 0 002-2v-3"/><line x1="6" y1="5" x2="6" y2="11"/></svg>'
    SVG_AREA = '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#4caf7d" stroke-width="1.8"><rect x="3" y="3" width="18" height="18" rx="1"/><path d="M3 9h18M9 3v18"/></svg>'
    SVG_LAND = '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#4caf7d" stroke-width="1.8"><path d="M2 20l5-8 4 5 3-4 8 7"/><circle cx="18" cy="5" r="2"/></svg>'
    SVG_CAR  = '<svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="#2a9db5" stroke-width="1.8"><rect x="2" y="10" width="20" height="8" rx="2"/><path d="M5 10l2-4h10l2 4"/><circle cx="7" cy="18" r="1.5"/><circle cx="17" cy="18" r="1.5"/></svg>'
    SVG_PIN  = '<svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#6b7a99" stroke-width="2"><path d="M12 2a7 7 0 017 7c0 5-7 13-7 13S5 14 5 9a7 7 0 017-7z"/><circle cx="12" cy="9" r="2.5"/></svg>'

    specs = []
    if rec:     specs.append((SVG_BED,  str(rec),      "Recámaras"))
    if ban:     specs.append((SVG_BATH, str(ban),      "Baños"))
    if m2c:     specs.append((SVG_AREA, fmt_m2(m2c),   "Construcción"))
    if m2t:     specs.append((SVG_LAND, fmt_m2(m2t),   "Terreno"))
    if parking and len(specs)<4: specs.append((SVG_CAR, str(parking), "Estacion."))

    specs_items = "".join(
        '<div class="spec-item"><div class="spec-ico">{}</div><div class="spec-val">{}</div><div class="spec-lbl">{}</div></div>'.format(s[0],s[1],s[2])
        for s in specs[:4]
    )
    specs_html = '<div class="cover-specs">{}</div>'.format(specs_items) if specs_items else ""

    foto_urls = [f.get("url") or f.get("original") or "" for f in fotos if f]
    hero_src  = images_b64.get(foto_urls[0], foto_urls[0]) if foto_urls else ""
    hero_html = '<img class="cover-hero" src="{}" alt="portada"/>'.format(hero_src) if hero_src else '<div class="cover-hero-placeholder"></div>'

    def footer():
        return '<div class="ficha-footer"><img src="{}" class="ft-logo" alt="Brokr"/><div class="ft-id">{}</div></div>'.format(LOGO, id_prop)

    gallery_fotos = foto_urls[:]  # include hero photo as first in gallery
    gallery_pages = ""
    total = len(gallery_fotos)
    full_pages = total // 6
    remainder  = total % 6

    for i in range(full_pages):
        batch = gallery_fotos[i*6:(i+1)*6]
        imgs  = "".join('<img src="{}" alt="foto"/>'.format(images_b64.get(u,u)) for u in batch)
        gallery_pages += '<div class="ficha-page"><div class="section-header"><h2>Galería fotográfica</h2></div><div class="photo-grid-6">{}</div>{}</div>'.format(imgs, footer())

    rows = []
    if p.get("property_type"): rows.append(("Tipo de inmueble", p["property_type"]))
    rows.append(("Operación", tipo_op))
    rows.append(("Precio", precio))
    if rec:  rows.append(("Recámaras", str(rec)))
    if ban:  rows.append(("Baños completos", str(ban)))
    if p.get("half_bathrooms"): rows.append(("Medios baños", str(p["half_bathrooms"])))
    if m2c:  rows.append(("Superficie construida", fmt_m2(m2c)))
    if m2t:  rows.append(("Superficie de terreno", fmt_m2(m2t)))
    if parking: rows.append(("Estacionamientos", str(parking)))
    if p.get("floors"): rows.append(("Niveles", str(p["floors"])))
    if colonia: rows.append(("Colonia", colonia))
    if ciudad:  rows.append(("Ciudad", ciudad))
    if id_prop: rows.append(("Clave", id_prop))

    rows_html = "".join('<tr><td class="char-lbl">{}</td><td class="char-val">{}</td></tr>'.format(k,v) for k,v in rows)

    amen_html = ""
    if amenids:
        items = "".join('<div class="amen-item">{}</div>'.format(a.get("name") or a) for a in amenids)
        amen_html = '<div class="amen-section"><div class="amen-ttl">Amenidades</div><div class="amen-grid">{}</div></div>'.format(items)

    chars_section = (
        '<div class="section-header chars-hdr"><h2>Características del inmueble</h2></div>'
        '<div class="chars-body"><table class="char-table"><tbody>{}</tbody></table>{}</div>'
    ).format(rows_html, amen_html)

    if remainder > 0:
        batch = gallery_fotos[full_pages*6:]
        imgs  = "".join('<img src="{}" alt="foto"/>'.format(images_b64.get(u,u)) for u in batch)
        rows_r = (remainder + 1) // 2
        gallery_pages += (
            '<div class="ficha-page">'
            '<div class="section-header"><h2>Galería fotográfica</h2></div>'
            '<div class="photo-grid-auto" style="grid-template-rows:repeat({},82mm);height:{}mm">{}</div>'
            '<div class="chars-inline">{}</div>'
            '{}</div>'
        ).format(rows_r, rows_r*82, imgs, chars_section, footer())
    else:
        gallery_pages += '<div class="ficha-page">{}{}</div>'.format(chars_section, footer())

    CSS = """
    <link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Poppins:ital,wght@0,100;0,200;0,300;0,400;0,500;0,600;0,700;0,800;0,900;1,100;1,200;1,300;1,400;1,500;1,600;1,700;1,800;1,900&display=swap" rel="stylesheet">
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Poppins',sans-serif;background:white;color:#0f1829}
.ficha-page{width:210mm;height:297mm;background:white;display:flex;flex-direction:column;overflow:hidden;page-break-after:always}
.ficha-page:last-child{page-break-after:avoid}
.cover-accent{height:4px;background:linear-gradient(90deg,#2a9db5 0%,#4caf7d 100%);flex-shrink:0}
.cover-hero{width:100%;height:120mm;object-fit:cover;display:block;flex-shrink:0}
.cover-hero-placeholder{width:100%;height:120mm;background:linear-gradient(135deg,#0f1829,#1a2744);flex-shrink:0}
.cover-info{padding:14px 24px 10px;border-bottom:1px solid #e8ecf2}
.cover-badge{display:inline-block;background:linear-gradient(135deg,#2a9db5,#1f8ba0);color:white;font-size:9px;font-weight:700;letter-spacing:1px;text-transform:uppercase;padding:4px 12px;border-radius:20px;margin-bottom:7px}
.cover-precio{font-family:'Poppins',sans-serif;font-size:30px;font-weight:700;color:#0f1829;line-height:1;margin-bottom:4px}
.cover-titulo{font-size:13px;font-weight:600;color:#1a2744;margin-bottom:3px}
.cover-ubicacion{font-size:11px;color:#6b7a99;display:flex;align-items:center;gap:4px}
.cover-specs{display:grid;grid-template-columns:repeat(4,1fr);border-bottom:2px solid #eef2f7;background:#fafbfc}
.spec-item{padding:10px 8px;text-align:center;border-right:1px solid #e8ecf2;display:flex;flex-direction:column;align-items:center;gap:3px}
.spec-item:last-child{border-right:none}
.spec-ico{display:flex;align-items:center;justify-content:center;height:26px}
.spec-val{font-size:14px;font-weight:700;color:#0f1829;line-height:1}
.spec-lbl{font-size:8px;text-transform:uppercase;letter-spacing:.6px;color:#6b7a99}
.cover-desc-wrap{padding:14px 24px 8px;flex:1;overflow:hidden}
.cover-desc-ttl{font-family:'Poppins',sans-serif;font-size:12px;font-weight:600;color:#0f1829;margin-bottom:7px;padding-bottom:5px;border-bottom:2px solid #4caf7d;display:inline-block}
.cover-desc{font-size:10.5px;color:#3a4a5c;line-height:1.65}
.section-header{padding:11px 24px 9px;border-bottom:1px solid #e8ecf2;border-left:4px solid #2a9db5;flex-shrink:0;background:#fafbfc}
.section-header h2{font-family:'Poppins',sans-serif;font-size:14px;font-weight:600;color:#0f1829}
.chars-hdr{border-left-color:#4caf7d}
.photo-grid-6{display:grid;grid-template-columns:1fr 1fr;grid-template-rows:82mm 82mm 82mm;gap:2px;padding:2px;height:246mm;flex-shrink:0;overflow:hidden}
.photo-grid-auto{display:grid;grid-template-columns:1fr 1fr;gap:2px;padding:2px;flex-shrink:0;overflow:hidden}
.photo-grid-6 img,.photo-grid-auto img{width:100%;height:100%;object-fit:cover;display:block}
.chars-inline{flex:1;overflow:hidden;display:flex;flex-direction:column;min-height:0}
.chars-body{padding:10px 24px 8px;overflow:hidden}
.char-table{width:100%;border-collapse:collapse}
.char-table tr{border-bottom:1px solid #eef2f7}
.char-table tr:nth-child(even) td{background:#f7f9fb}
.char-lbl{padding:6px 10px;font-size:11px;color:#6b7a99;width:48%}
.char-val{padding:6px 10px;font-size:11px;color:#0f1829;font-weight:600;text-align:right}
.amen-section{margin-top:10px}
.amen-ttl{font-size:9px;font-weight:700;text-transform:uppercase;letter-spacing:.7px;color:#4caf7d;margin-bottom:6px}
.amen-grid{display:flex;flex-wrap:wrap;gap:5px}
.amen-item{font-size:10px;padding:3px 9px;background:#f0faf5;border-radius:20px;color:#1a5c38;border:1px solid #b2dfcc;font-weight:500}
.ficha-footer{width:100%;height:44px;background:linear-gradient(90deg,#0f1829 0%,#1a2744 100%);display:flex;align-items:center;justify-content:space-between;padding:0 20px;flex-shrink:0;margin-top:auto}
.ft-logo{height:16px;filter:brightness(10)}
.ft-id{font-size:9px;color:rgba(255,255,255,.45);letter-spacing:.6px}
@page{size:A4 portrait;margin:0}
"""
    cover_desc_html = (
        '<div class="cover-desc-wrap"><div class="cover-desc-ttl">Descripción</div>'
        '<div class="cover-desc">{}</div></div>'.format(desc)
    ) if desc else '<div style="flex:1"></div>'

    return (
        "<!DOCTYPE html><html lang='es'><head><meta charset='UTF-8'/>"
        "<style>{}</style></head><body>"
        "<div class='ficha-page'>"
        "<div class='cover-accent'></div>"
        "{}"
        "<div class='cover-info'>"
        "<div class='cover-badge'>{}</div>"
        "<div class='cover-precio'>{}</div>"
        "<div class='cover-titulo'>{}</div>"
        "<div class='cover-ubicacion'>{} {}</div>"
        "</div>"
        "{}"
        "{}"
        "{}"
        "</div>"
        "{}"
        "</body></html>"
    ).format(CSS, hero_html, tipo_op, precio, titulo, SVG_PIN, ubicacion,
             specs_html, cover_desc_html, footer(), gallery_pages)


# ────────────────────────────────────────────
# NOTICIAS INMOBILIARIAS — RSS REAL
# ────────────────────────────────────────────
import xml.etree.ElementTree as ET

@app.get("/noticias")
async def get_noticias():
    """Fetch real estate news from Google News RSS — parsed server-side to avoid CORS."""
    FEEDS = [
        "https://news.google.com/rss/search?q=bienes+raices+Mexico&hl=es-419&gl=MX&ceid=MX:es-419",
        "https://news.google.com/rss/search?q=mercado+inmobiliario+Mexico&hl=es-419&gl=MX&ceid=MX:es-419",
    ]

    cached = cache_get("noticias_rss")
    if cached is not None:
        return cached

    items = []
    seen = set()

    async with httpx.AsyncClient(timeout=10, follow_redirects=True) as client:
        for feed_url in FEEDS:
            try:
                r = await client.get(feed_url, headers={"User-Agent": "Mozilla/5.0"})
                if r.status_code != 200:
                    continue
                root = ET.fromstring(r.text)
                channel = root.find("channel")
                if channel is None:
                    continue
                for item in channel.findall("item")[:8]:
                    title_el = item.find("title")
                    link_el  = item.find("link")
                    source_el = item.find("source")
                    if title_el is None or link_el is None:
                        continue
                    title = title_el.text or ""
                    # Strip trailing source name like "- El Universal"
                    title = re.sub(r"\s*[-–]\s*[^-–]+$", "", title).strip()
                    link  = link_el.text or ""
                    source = source_el.text if source_el is not None else "Google News"
                    if title in seen or not title or not link:
                        continue
                    seen.add(title)
                    items.append({"title": title, "url": link, "source": source})
                    if len(items) >= 12:
                        break
            except Exception:
                continue
            if len(items) >= 12:
                break

    if not items:
        # Fallback vacío — el front usará sus estáticos
        return {"items": []}

    result = {"items": items}
    cache_set("noticias_rss", result, ttl=1800)  # Cache 30 minutos
    return result


@app.post("/ficha-pdf")
async def generar_ficha_pdf(p: dict):
    """Generate PDF from property data dict using Playwright."""
    import httpx
    
    # Collect all image URLs
    fotos = p.get("property_images") or []
    urls = list(set(filter(None, [f.get("url") or f.get("original") for f in fotos])))
    
    # Download all images concurrently and convert to base64
    images_b64 = {}
    async with httpx.AsyncClient(timeout=30) as client:
        async def fetch_img(url):
            try:
                r = await client.get(url, follow_redirects=True, timeout=10.0)
                if r.status_code == 200:
                    ext = url.split(".")[-1].split("?")[0].lower()
                    mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
                            "webp": "image/webp", "gif": "image/gif"}.get(ext, "image/jpeg")
                    b64 = base64.b64encode(r.content).decode()
                    images_b64[url] = f"data:{mime};base64,{b64}"
            except Exception:
                pass  # skip failed images, show blank

        # Limit to 19 gallery images (1 hero + 18 gallery = 3 full pages max)
        await asyncio.gather(*[fetch_img(u) for u in urls[:19]])
    
    # Build HTML
    html = build_ficha_html(p, images_b64)
    
    # Render to PDF with Playwright
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        page = await browser.new_page()
        # Use domcontentloaded instead of networkidle — images are already base64
        await page.set_content(html, wait_until="domcontentloaded")
        await page.wait_for_timeout(500)  # small wait for fonts
        pdf_bytes = await page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "0", "right": "0", "bottom": "0", "left": "0"}
        )
        await browser.close()
    
    from fastapi.responses import JSONResponse
    import re as _re2
    id_prop   = p.get("public_id") or p.get("id") or ""
    loc       = p.get("location") or {}
    colonia   = (loc.get("name") or "").strip()
    tipo_raw  = (p.get("property_type") or "Propiedad").strip()
    # Sanitize: remove accents and special chars for filename
    def _slug(s):
        for a, b in [('á','a'),('é','e'),('í','i'),('ó','o'),('ú','u'),('ü','u'),('ñ','n'),
                     ('Á','A'),('É','E'),('Í','I'),('Ó','O'),('Ú','U'),('Ñ','N')]:
            s = s.replace(a, b)
        return _re2.sub(r'[^A-Za-z0-9_]', '_', s).strip('_')
    parts = ["Ficha_Brokr"]
    if colonia:  parts.append(_slug(colonia))
    if tipo_raw: parts.append(_slug(tipo_raw))
    if id_prop:  parts.append(_slug(id_prop))
    filename = "_".join(parts) + ".pdf"
    token = str(_uuid.uuid4()).replace("-","")[:16]
    _pdf_store[token] = (pdf_bytes, filename)
    # Clean old entries if too many
    if len(_pdf_store) > 50:
        oldest = list(_pdf_store.keys())[0]
        del _pdf_store[oldest]
    return JSONResponse({"token": token, "filename": filename})

@app.get("/ficha-pdf/{token}")
async def descargar_ficha_pdf(token: str):
    """Serve generated PDF by token — opens natively in Safari."""
    from fastapi.responses import Response
    if token not in _pdf_store:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="PDF no encontrado o expirado")
    pdf_bytes, filename = _pdf_store[token]
    # Use attachment for direct download on all devices including PWA
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "application/pdf",
            "Cache-Control": "no-store",
        }
    )
# ────────────────────────────────────────────
# AVM — COMPARABLES VÍA APIFY + INMUEBLES24
# ────────────────────────────────────────────

APIFY_ACTOR   = "azzouzana~inmuebles24-scraper-pro-by-search-url"

# Mapeo de tipo de inmueble a término de búsqueda en Inmuebles24
TIPO_URL = {
    "casa":         "casas",
    "departamento": "departamentos",
    "terreno":      "terrenos",
    "local":        "locales-comerciales",
    "oficina":      "oficinas",
    "bodega":       "bodegas",
    "edificio":     "edificios",
}

class ComparablesRequest(BaseModel):
    colonia: str
    ciudad: str = "morelia"
    estado: str = "michoacan-de-ocampo"
    tipo: str = "casa"          # casa | departamento | terreno | local | oficina | bodega | edificio
    max_resultados: int = 10    # cuántos comparables traer


def construir_url_inmuebles24(tipo: str, colonia: str, ciudad: str, estado: str) -> str:
    segmento = TIPO_URL.get(tipo, "casas")
    ciudad = ciudad.lower().strip().replace(" ", "-")
    col = colonia.lower().strip().replace(" ", "-")
    return f"https://www.inmuebles24.com/{segmento}-en-{ciudad}-o-{col}.html"


def normalizar_listing(item: dict) -> dict:
    """Convierte un resultado de Apify (scraper Azzouzana) al formato que espera el AVM."""
    
    # Precio
    precio = item.get("price_amount") or 0
    moneda = item.get("price_currency", "MN")
    # Ignorar propiedades en USD (fuera de mercado local)
    if moneda == "USD":
        return None

    # m² de construcción — viene en generatedTitle: "Casa · 120m² · 3 Recámaras"
    m2c = 0
    titulo_gen = item.get("generatedTitle", "")
    match_m2 = re.search(r'(\d+)m²', titulo_gen)
    if match_m2:
        m2c = float(match_m2.group(1))

    # Recámaras
    recamaras = 0
    match_rec = re.search(r'(\d+)\s+Rec[áa]maras?', titulo_gen, re.IGNORECASE)
    if match_rec:
        recamaras = int(match_rec.group(1))

    # Estacionamientos
    estac = 0
    match_estac = re.search(r'(\d+)\s+Estacionamientos?', titulo_gen, re.IGNORECASE)
    if match_estac:
        estac = int(match_estac.group(1))

    # m² terreno — intentar extraer de descripción
    m2t = 0
    desc = item.get("descriptionNormalized", "")
    patrones_terreno = [
        r'[Tt]erreno[:\s/]+(\d+[\.,]?\d*)\s*(?:m²|m2|metros cuadrados|metros)',
        r'(\d+[\.,]?\d*)\s*(?:m²|m2)\s*de\s+terreno',
        r'[Ss]uperficie\s+de\s+terreno[:\s]+[\d,\s]*(\d+)\s*(?:m²|m2)',
        r'[Tt]erreno\s+de\s+(\d+[\.,]?\d*)\s*(?:m²|m2)',
    ]
    for patron in patrones_terreno:
        match_t = re.search(patron, desc)
        if match_t:
            val = match_t.group(1).replace(',', '').replace('.', '')
            try:
                m2t = float(val)
                if m2t < 10 or m2t > 50000:
                    m2t = 0
            except:
                m2t = 0
            if m2t > 0:
                break

    titulo = item.get("title") or ""
    url = item.get("url") or ""
    imagenes = item.get("images", [])
    imagen = imagenes[0].split("?")[0] if imagenes else ""

    return {
        "precio": int(precio),
        "m2Construccion": m2c,
        "m2Terreno": m2t,
        "recamaras": recamaras,
        "banos": 0,
        "estacionamiento": estac,
        "edad": 0,
        "conservacion": "bueno",
        "calidad": "medio",
        "mismaZona": "si",
        "titulo": titulo,
        "url": url,
        "imagen": imagen,
    }

@app.post("/api/comparables")
async def buscar_comparables(req: ComparablesRequest):
    """
    Llama a Apify (actor de Inmuebles24) y regresa comparables normalizados
    listos para el AVM.
    """
    if not APIFY_API_KEY:
        raise HTTPException(status_code=500, detail="APIFY_API_KEY no configurada en el servidor")

    url_busqueda = construir_url_inmuebles24(req.tipo, req.colonia, req.ciudad, req.estado)

    # Cache key para no re-scrapear la misma búsqueda en 2 horas
    cache_key = f"comparables_{req.tipo}_{req.colonia}_{req.ciudad}"
    cached = cache_get(cache_key)
    if cached is not None:
        return cached

    # Llamada a Apify — run-sync-get-dataset-items (espera hasta que termina)
    apify_url = (
        f"https://api.apify.com/v2/acts/{APIFY_ACTOR}"
        f"/run-sync-get-dataset-items?token={APIFY_API_KEY}"
        f"&timeout=60&memory=256"
    )

    payload = {
        "startUrl": url_busqueda,
        "maxItems": req.max_resultados,
    }

    async with httpx.AsyncClient(timeout=90) as client:
        try:
            r = await client.post(apify_url, json=payload)
        except httpx.TimeoutException:
            raise HTTPException(status_code=504, detail="Apify tardó demasiado. Intenta de nuevo.")

        if r.status_code not in (200, 201):
            raise HTTPException(
                status_code=502,
                detail=f"Error de Apify: {r.status_code} — {r.text[:300]}"
            )

        items = r.json()

    if not isinstance(items, list):
        raise HTTPException(status_code=502, detail="Respuesta inesperada de Apify")

    # Filtrar items con precio y m2 válidos, normalizar
    comparables = []
    for item in items:
        n = normalizar_listing(item)
        if n["precio"] > 0 and n["m2Construccion"] > 0:
            comparables.append(n)

    resultado = {
        "url_busqueda": url_busqueda,
        "total": len(comparables),
        "comparables": comparables,
    }

    cache_set(cache_key, resultado, ttl=7200)  # cache 2 horas
    return resultado

# ────────────────────────────────────────────
# AVM — COLONIAS (Nominatim) Y COMPARABLES CERCANOS (Supabase)
# ────────────────────────────────────────────

class ColoniasRequest(BaseModel):
    texto: str
    ciudad: str = "Morelia"

@app.get("/api/colonias")
async def buscar_colonias(texto: str, ciudad: str = "Morelia"):
    if len(texto) < 3:
        return {"colonias": []}

    cache_key = f"colonias_g3_{ciudad}_{texto}".lower()
    cached = cache_get(cache_key)
    if cached:
        return cached

    if not GOOGLE_PLACES_KEY:
        return {"colonias": [], "error": "GOOGLE_PLACES_KEY no configurada"}

    async with httpx.AsyncClient(timeout=15) as client:
        try:
            r = await client.get(
                "https://maps.googleapis.com/maps/api/place/autocomplete/json",
                params={
                    "input": texto,
                    "types": "geocode",
                    "language": "es",
                    "components": "country:mx",
                    "locationbias": "circle:50000@19.7059504,-101.1949825",
                    "key": GOOGLE_PLACES_KEY,
                }
            )
            data = r.json()
        except Exception as e:
            return {"colonias": [], "error": str(e)}

    colonias = []
    for pred in data.get("predictions", []):
        descripcion = pred.get("description", "")
        tipos = pred.get("types", [])

        if not any(t in tipos for t in ["sublocality", "sublocality_level_1", "neighborhood"]):
            continue

        nombre = pred.get("structured_formatting", {}).get("main_text", "").strip()
        place_id = pred.get("place_id", "")

        lat, lon = 0.0, 0.0
        if place_id:
            try:
                async with httpx.AsyncClient(timeout=10) as client2:
                    r2 = await client2.get(
                        "https://maps.googleapis.com/maps/api/place/details/json",
                        params={
                            "place_id": place_id,
                            "fields": "geometry",
                            "key": GOOGLE_PLACES_KEY,
                        }
                    )
                    details_data = r2.json()
                    loc = details_data.get("result", {}).get("geometry", {}).get("location", {})
                    lat = loc.get("lat", 0.0)
                    lon = loc.get("lng", 0.0)
            except Exception:
                pass

        if nombre:
            colonias.append({
                "nombre":    nombre,
                "display":   descripcion,
                "latitud":   lat,
                "longitud":  lon,
                "place_id":  place_id,
            })

    resultado = {"colonias": colonias[:6]}
    cache_set(cache_key, resultado, ttl=86400)
    return resultado


# ────────────────────────────────────────────
# AVM — COMPARABLES CERCANOS (PostGIS + Supabase)
# ────────────────────────────────────────────

# CercanosRequest — única definición
class CercanosRequest(BaseModel):
    latitud: float
    longitud: float
    tipo: str = "casa"
    radio_km: float = 2.0
    max_resultados: int = 15

# TIPO_MAP_DB — mapeo hacia tipos de Supabase/PostGIS (distinto del TIPO_MAP de EasyBroker arriba)
TIPO_MAP_DB = {
    "casa":         ["Casas", "Desarrollos horizontales", "Desarrollos Horizontal/Vertical"],
    "departamento": ["Departamentos", "Desarrollos verticales"],
    "terreno":      ["Terrenos"],
    "local":        ["Locales comerciales", "Locales Comerciales"],
    "oficina":      ["Oficinas"],
    "bodega":       ["Bodegas"],
    "edificio":     ["Edificios"],
}

@app.post("/api/comparables-cercanos")
async def comparables_cercanos(req: CercanosRequest):
    """Busca propiedades cercanas en Supabase usando PostGIS."""
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise HTTPException(status_code=500, detail="SUPABASE_URL o SUPABASE_ANON_KEY no configuradas")

    cache_key = f"cercanos_{req.tipo}_{req.latitud:.4f}_{req.longitud:.4f}_{req.radio_km}"
    cached = cache_get(cache_key)
    if cached:
        return cached

    tipos_db = TIPO_MAP_DB.get(req.tipo, ["Casas"])
    radio_metros = int(req.radio_km * 1000)

    # Llamar a función RPC en Supabase que ejecuta la query PostGIS
    payload = {
        "lat": req.latitud,
        "lon": req.longitud,
        "radio": radio_metros,
        "tipos": tipos_db,
        "limite": req.max_resultados,
    }

    headers = {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
    }

    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.post(
            f"{SUPABASE_URL}/rest/v1/rpc/buscar_cercanos",
            headers=headers,
            json=payload,
        )

    if r.status_code not in (200, 201):
        # Fallback: buscar por ciudad sin PostGIS
        async with httpx.AsyncClient(timeout=15) as client:
            r2 = await client.get(
                f"{SUPABASE_URL}/rest/v1/propiedades_avm",
                headers=headers,
                params={
                    "ciudad": "eq.Morelia",
                    "precio": "gt.0",
                    "metros_construccion": "not.is.null",
                    "select": "id,titulo,precio,tipo_propiedad,metros_construccion,metros_terreno,recamaras,estacionamientos,colonia,ciudad,url,latitud,longitud",
                    "limit": req.max_resultados,
                    "order": "precio.asc",
                }
            )
        items = r2.json() if r2.status_code == 200 else []
    else:
        items = r.json() or []

    comparables = []
    for item in items:
        precio = item.get("precio") or 0
        m2c    = item.get("metros_construccion") or 0
        if precio <= 0 or m2c <= 0:
            continue
        comparables.append({
            "precio":           int(precio),
            "m2Construccion":   float(m2c),
            "m2Terreno":        float(item.get("metros_terreno") or 0),
            "recamaras":        int(item.get("recamaras") or 0),
            "estacionamiento":  int(item.get("estacionamientos") or 0),
            "banos":            0,
            "edad":             0,
            "conservacion":     "bueno",
            "calidad":          "medio",
            "mismaZona":        "si",
            "titulo":           item.get("titulo") or "",
            "url":              item.get("url") or "",
            "imagen":           "",
            "colonia":          item.get("colonia") or "",
            "distancia_metros": int(item.get("distancia_metros") or 0),
        })

    resultado = {
        "total":       len(comparables),
        "comparables": comparables,
        "latitud":     req.latitud,
        "longitud":    req.longitud,
        "radio_km":    req.radio_km,
    }
    cache_set(cache_key, resultado, ttl=3600)
    return resultado


# ─── LIMPIEZA DE IMÁGENES ─────────────────────────────────────────────────────

FB_APP_ID     = os.environ.get("FB_APP_ID", "")
FB_APP_SECRET = os.environ.get("FB_APP_SECRET", "")
FRONTEND_URL  = os.environ.get("FRONTEND_URL", "https://brokr.app")

def _process_image_sync(file_bytes: bytes, content_type: str) -> bytes:
    """Pipeline de mejora automática (sin IA generativa): denoising, CLAHE, WB, unsharp."""
    if not PIL_AVAILABLE:
        return file_bytes
    img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    if CV2_AVAILABLE:
        arr = np.array(img)
        arr_bgr = cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)

        # 1. Denoising adaptativo
        gray = cv2.cvtColor(arr_bgr, cv2.COLOR_BGR2GRAY)
        noise_est = np.std(cv2.Laplacian(gray.astype(np.float64), cv2.CV_64F))
        if noise_est > 12:
            arr_bgr = cv2.fastNlMeansDenoisingColored(arr_bgr, None, 7, 7, 7, 21)

        # 2. Espacio LAB
        lab = cv2.cvtColor(arr_bgr, cv2.COLOR_BGR2LAB)
        l_ch, a_ch, b_ch = cv2.split(lab)

        # 3. CLAHE en L
        clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(8, 8))
        l_ch = clahe.apply(l_ch)

        # 4. LUT sombras/altas luces
        lut = np.arange(256, dtype=np.float32)
        lut = np.where(lut < 80,  lut * 1.12, lut)
        lut = np.where(lut > 210, 210 + (lut - 210) * 0.55, lut)
        lut = np.clip(lut, 0, 255).astype(np.uint8)
        l_ch = cv2.LUT(l_ch, lut)

        # 5. Vibrance A/B
        a_ch = np.clip((a_ch.astype(np.int16) - 128) * 1.1 + 128, 0, 255).astype(np.uint8)
        b_ch = np.clip((b_ch.astype(np.int16) - 128) * 1.1 + 128, 0, 255).astype(np.uint8)
        arr_bgr = cv2.cvtColor(cv2.merge([l_ch, a_ch, b_ch]), cv2.COLOR_LAB2BGR)

        # 6. Balance de blancos parcial (70% gray-world)
        bc, gc, rc = cv2.split(arr_bgr.astype(np.float32))
        mb, mg, mr = bc.mean(), gc.mean(), rc.mean()
        mg_all = (mb + mg + mr) / 3
        s = 0.7
        bc = np.clip(bc * (1 + s * (mg_all / max(mb, 1) - 1)), 0, 255)
        gc = np.clip(gc * (1 + s * (mg_all / max(mg, 1) - 1)), 0, 255)
        rc = np.clip(rc * (1 + s * (mg_all / max(mr, 1) - 1)), 0, 255)
        arr_bgr = cv2.merge([bc.astype(np.uint8), gc.astype(np.uint8), rc.astype(np.uint8)])

        # 7. Unsharp masking
        blur = cv2.GaussianBlur(arr_bgr, (0, 0), 1.5)
        arr_bgr = np.clip(cv2.addWeighted(arr_bgr, 1.45, blur, -0.45, 0), 0, 255).astype(np.uint8)

        img = Image.fromarray(cv2.cvtColor(arr_bgr, cv2.COLOR_BGR2RGB))
    else:
        img = ImageEnhance.Contrast(img).enhance(1.2)
        img = ImageEnhance.Brightness(img).enhance(1.05)
        img = ImageEnhance.Color(img).enhance(1.15)
        img = ImageEnhance.Sharpness(img).enhance(1.6)

    out = io.BytesIO()
    fmt = "JPEG" if (content_type or "").lower() in ("image/jpeg", "image/jpg") else "PNG"
    img.save(out, format=fmt, quality=92, optimize=True)
    return out.getvalue()


async def _process_with_gemini(img_bytes: bytes, content_type: str, prompt: str) -> bytes:
    """Edita la imagen con Gemini Flash imagen-generation."""
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY no configurada")

    # Resize a máx 1024px para reducir payload y tiempo de proceso
    if PIL_AVAILABLE:
        pil = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        w, h = pil.size
        if max(w, h) > 1024:
            scale = 1024 / max(w, h)
            pil = pil.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        pil.save(buf, format="JPEG", quality=85)
        img_bytes = buf.getvalue()

    img_b64 = base64.b64encode(img_bytes).decode()
    full_prompt = (
        "You are a professional real estate photo editor. "
        "Edit this photo: " + prompt + ". "
        "Output only the edited image."
    )

    # Solo v1beta — los modelos Nano Banana no están en v1
    # Solo 2 payloads: con imagen (preferido) y solo texto (fallback)
    _payloads = [
        {"contents": [{"parts": [
            {"text": full_prompt},
            {"inline_data": {"mime_type": "image/jpeg", "data": img_b64}},
        ]}]},
        {"contents": [{"parts": [{"text": full_prompt}]}]},
    ]

    # Modelos en orden de preferencia — solo v1beta
    _model_names = [m for m in [
        os.environ.get("GEMINI_IMAGE_MODEL", ""),
        "gemini-3.1-flash-image-preview",   # Nano Banana 2
        "gemini-2.5-flash-image",            # Nano Banana
        "gemini-3-pro-image-preview",        # Nano Banana Pro
    ] if m]

    GEMINI_BASE_URL = "https://generativelanguage.googleapis.com/v1beta"
    last_err = "Sin modelos disponibles"

    # Timeout 25s por petición — Railway corta a ~60s total, necesitamos margen
    async with httpx.AsyncClient(timeout=25) as client:
        for model_name in _model_names:
            url = f"{GEMINI_BASE_URL}/models/{model_name}:generateContent?key={GEMINI_API_KEY}"
            for payload in _payloads:
                try:
                    r = await client.post(url, json=payload, headers={"Content-Type": "application/json"})
                except Exception as e:
                    last_err = f"Timeout ({model_name}): {e}"
                    break  # red fallida, pasar al siguiente modelo

                if r.status_code == 404:
                    last_err = f"Modelo no encontrado: {model_name}"
                    break  # este modelo no existe, probar siguiente

                if r.status_code == 429:
                    # Cuota agotada — no tiene sentido probar otros modelos
                    raise RuntimeError(
                        "Cuota de Gemini agotada. Espera a que se reinicie tu límite gratuito "
                        "(~24h) o activa billing en https://aistudio.google.com/apikey."
                    )

                if r.status_code == 200:
                    try:
                        data = r.json()
                        parts = data["candidates"][0]["content"]["parts"]
                    except Exception as e:
                        last_err = f"JSON inválido ({model_name}): {e}"
                        continue

                    for part in parts:
                        if "inlineData" in part:
                            raw = base64.b64decode(part["inlineData"]["data"])
                            if PIL_AVAILABLE:
                                pil2 = Image.open(io.BytesIO(raw)).convert("RGB")
                                out = io.BytesIO()
                                pil2.save(out, format="JPEG", quality=92)
                                return out.getvalue()
                            return raw

                    text_parts = [p.get("text", "") for p in parts if "text" in p]
                    last_err = f"Sin imagen en respuesta ({model_name}): {' '.join(text_parts)[:150]}"
                    continue

                last_err = f"Error {r.status_code} ({model_name}): {r.text[:200]}"
                continue

    raise RuntimeError(last_err)


from fastapi import Form as _Form

@app.post("/images/clean")
async def clean_images(
    files: List[UploadFile] = File(...),
    prompt: str = _Form(""),
    # legacy field kept for backward compat
    remove_furniture: str = _Form("false"),
):
    use_gemini = bool(prompt.strip()) and bool(GEMINI_API_KEY)

    async def process_one(uf: UploadFile):
        raw = await uf.read()
        ct = uf.content_type or "image/jpeg"
        try:
            if use_gemini:
                processed = await _process_with_gemini(raw, ct, prompt.strip())
                return {
                    "name": uf.filename,
                    "cleaned_b64": base64.b64encode(processed).decode(),
                    "content_type": "image/jpeg",
                    "used_gemini": True,
                    "error": None,
                }
            else:
                loop = asyncio.get_event_loop()
                processed = await loop.run_in_executor(
                    _thread_pool, _process_image_sync, raw, ct
                )
                return {
                    "name": uf.filename,
                    "cleaned_b64": base64.b64encode(processed).decode(),
                    "content_type": ct,
                    "used_gemini": False,
                    "error": None,
                }
        except Exception as exc:
            return {
                "name": uf.filename,
                "cleaned_b64": None,
                "content_type": ct,
                "used_gemini": False,
                "error": str(exc),
            }

    results = await asyncio.gather(*[process_one(f) for f in files])
    return {"images": list(results)}


# ─── FACEBOOK OAUTH ───────────────────────────────────────────────────────────

@app.get("/facebook/callback")
async def facebook_callback(code: str = Query(...), state: str = Query(None), redirect_uri: str = Query(None)):
    """Intercambia el code de OAuth por un token de página de Facebook."""
    redirect_uri = redirect_uri or (FRONTEND_URL + "/facebook/callback")
    async with httpx.AsyncClient(timeout=15) as client:
        # 1. Token de usuario (corta duración)
        r = await client.get(
            "https://graph.facebook.com/v21.0/oauth/access_token",
            params={
                "client_id": FB_APP_ID,
                "client_secret": FB_APP_SECRET,
                "redirect_uri": redirect_uri,
                "code": code,
            },
        )
        if r.status_code != 200:
            return {"error": r.text}
        short_token = r.json().get("access_token", "")

        # 2. Token de larga duración
        r2 = await client.get(
            "https://graph.facebook.com/v21.0/oauth/access_token",
            params={
                "grant_type": "fb_exchange_token",
                "client_id": FB_APP_ID,
                "client_secret": FB_APP_SECRET,
                "fb_exchange_token": short_token,
            },
        )
        long_token = r2.json().get("access_token", short_token)

        # 3. Lista de páginas administradas
        r3 = await client.get(
            "https://graph.facebook.com/v21.0/me/accounts",
            params={"access_token": long_token},
        )
        pages = r3.json().get("data", [])

    if not pages:
        return {"error": "No se encontraron páginas administradas en esta cuenta de Facebook."}

    # Usar la primera página
    page = pages[0]
    page_token = page.get("access_token", "")
    page_id    = page.get("id", "")
    page_name  = page.get("name", "")

    # Devolver datos para que el frontend los guarde en Supabase
    return {
        "ok": True,
        "page_id": page_id,
        "page_name": page_name,
        "page_token": page_token,
    }


class FbPublishRequest(BaseModel):
    page_id: str
    page_token: str
    message: str
    photo_urls: list[str] = []

@app.post("/facebook/publish")
async def facebook_publish(req: FbPublishRequest):
    """Publica una propiedad en la página de Facebook."""
    photo_ids = []
    async with httpx.AsyncClient(timeout=30) as client:
        # Subir fotos como no publicadas
        for url in req.photo_urls[:10]:
            r = await client.post(
                f"https://graph.facebook.com/v21.0/{req.page_id}/photos",
                params={"access_token": req.page_token},
                json={"url": url, "published": False},
            )
            if r.status_code == 200:
                pid = r.json().get("id")
                if pid:
                    photo_ids.append({"media_fbid": pid})

        # Crear el post
        payload: dict = {
            "message": req.message,
            "access_token": req.page_token,
        }
        if photo_ids:
            payload["attached_media"] = photo_ids

        r_post = await client.post(
            f"https://graph.facebook.com/v18.0/{req.page_id}/feed",
            params={"access_token": req.page_token},
            json=payload,
        )

    if r_post.status_code not in (200, 201):
        raise HTTPException(status_code=502, detail=r_post.text)

    return {"ok": True, "post_id": r_post.json().get("id")}
